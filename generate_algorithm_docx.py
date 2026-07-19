# -*- coding: utf-8 -*-
"""
生成"核心算法表对照与修正"docx文档
公式使用Word原生OMML可编辑格式，表格使用三线表格式
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import copy

# ============================================================
# OMML 辅助函数
# ============================================================

M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def _m(tag):
    return f'{{{M_NS}}}{tag}'

def _w(tag):
    return f'{{{W_NS}}}{tag}'

def make_run(text, italic=True):
    """创建数学运行元素 m:r"""
    r = etree.SubElement(etree.Element(_m('tmp')), _m('r'))
    # 字体属性
    rpr = etree.SubElement(r, _w('rPr'))
    rfonts = etree.SubElement(rpr, _w('rFonts'))
    rfonts.set(_w('ascii'), 'Cambria Math')
    rfonts.set(_w('hAnsi'), 'Cambria Math')
    if italic:
        i = etree.SubElement(rpr, _w('i'))
    t = etree.SubElement(r, _m('t'))
    t.text = text
    return r

def make_text(text):
    """创建普通数学文本"""
    return make_run(text, italic=False)

def make_italic(text):
    """创建斜体数学文本"""
    return make_run(text, italic=True)

def make_ssub(base, sub):
    """创建下标: base_sub"""
    ssub = etree.Element(_m('sSub'))
    e = etree.SubElement(ssub, _m('e'))
    if isinstance(base, str):
        e.append(make_italic(base))
    else:
        e.append(base)
    sub_elem = etree.SubElement(ssub, _m('sub'))
    if isinstance(sub, str):
        sub_elem.append(make_italic(sub))
    else:
        sub_elem.append(sub)
    return ssub

def make_ssup(base, sup):
    """创建上标: base^sup"""
    ssup = etree.Element(_m('sSup'))
    e = etree.SubElement(ssup, _m('e'))
    if isinstance(base, str):
        e.append(make_italic(base))
    else:
        e.append(base)
    sup_elem = etree.SubElement(ssup, _m('sup'))
    if isinstance(sup, str):
        sup_elem.append(make_italic(sup))
    else:
        sup_elem.append(sup)
    return ssup

def make_ssubsup(base, sub, sup):
    """创建上下标: base_sub^sup"""
    ssubsup = etree.Element(_m('sSubSup'))
    e = etree.SubElement(ssubsup, _m('e'))
    if isinstance(base, str):
        e.append(make_italic(base))
    else:
        e.append(base)
    sub_elem = etree.SubElement(ssubsup, _m('sub'))
    if isinstance(sub, str):
        sub_elem.append(make_italic(sub))
    else:
        sub_elem.append(sub)
    sup_elem = etree.SubElement(ssubsup, _m('sup'))
    if isinstance(sup, str):
        sup_elem.append(make_italic(sup))
    else:
        sup_elem.append(sup)
    return ssubsup

def make_func(name, arg_elements):
    """创建函数: name(arg)"""
    func = etree.Element(_m('func'))
    fname = etree.SubElement(func, _m('fName'))
    fname.append(make_text(name))
    e = etree.SubElement(func, _m('e'))
    for elem in arg_elements:
        e.append(elem)
    return func

def make_d(elements, beg='(', end=')'):
    """创建定界符: (elements)"""
    d = etree.Element(_m('d'))
    dpr = etree.SubElement(d, _m('dPr'))
    begchr = etree.SubElement(dpr, _m('begChr'))
    begchr.set(_m('val'), beg)
    endchr = etree.SubElement(dpr, _m('endChr'))
    endchr.set(_m('val'), end)
    e = etree.SubElement(d, _m('e'))
    for elem in elements:
        e.append(elem)
    return d

def make_frac(num_elements, den_elements):
    """创建分式: num/den"""
    f = etree.Element(_m('f'))
    num = etree.SubElement(f, _m('num'))
    for elem in num_elements:
        num.append(elem)
    den = etree.SubElement(f, _m('den'))
    for elem in den_elements:
        den.append(elem)
    return f

def make_nary(op, sub_elements, sup_elements, body_elements):
    """创建n元运算: Σ"""
    nary = etree.Element(_m('nary'))
    narypr = etree.SubElement(nary, _m('naryPr'))
    chr_elem = etree.SubElement(narypr, _m('chr'))
    chr_elem.set(_m('val'), op)
    sub = etree.SubElement(nary, _m('sub'))
    for elem in sub_elements:
        sub.append(elem)
    sup = etree.SubElement(nary, _m('sup'))
    for elem in sup_elements:
        sup.append(elem)
    e = etree.SubElement(nary, _m('e'))
    for elem in body_elements:
        e.append(elem)
    return nary

def make_abs(elements):
    """创建绝对值: |elements|"""
    return make_d(elements, '|', '|')

def make_omath(*elements):
    """创建oMath容器"""
    omath = etree.Element(_m('oMath'))
    for elem in elements:
        omath.append(elem)
    return omath

def insert_omml(paragraph, omath):
    """将OMML插入段落"""
    # 添加空格run作为占位
    paragraph._p.append(omath)

def add_text_run(paragraph, text, bold=False, size=10.5, font_name='宋体'):
    """添加普通文本run"""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    # 设置中文字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(_w('rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(_w('eastAsia'), font_name)
    return run

def add_code_run(paragraph, text, size=9):
    """添加代码格式文本run"""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Consolas'
    return run

# ============================================================
# 公式构建器
# ============================================================

def build_formula_1():
    """情绪演化方程: E_t^k = tanh(α·E_{t-1}^k + γ·Φ_t)"""
    elements = [
        make_ssubsup('E', 't', 'k'),
        make_text('='),
        make_func('tanh', [
            make_italic('α'),
            make_text('·'),
            make_ssubsup('E', make_text('t-1'), 'k'),
            make_text('+'),
            make_italic('γ'),
            make_text('·'),
            make_ssub('Φ', 't'),
        ])
    ]
    return make_omath(*elements)

def build_formula_2():
    """反馈信号: Φ_t = -w_s·stockout_rate + w_m·match_factor - w_e·excess_rate"""
    elements = [
        make_ssub('Φ', 't'),
        make_text('='),
        make_text('-'),
        make_ssub('w', 's'),
        make_text('·'),
        make_text('stockout_rate'),
        make_text('+'),
        make_ssub('w', 'm'),
        make_text('·'),
        make_text('match_factor'),
        make_text('-'),
        make_ssub('w', 'e'),
        make_text('·'),
        make_text('excess_rate'),
    ]
    return make_omath(*elements)

def build_formula_3():
    """情绪传染: E_{k+1} = tanh(E_{k+1} + (-s_c))"""
    elements = [
        make_ssub('E', make_text('k+1')),
        make_text('←'),
        make_func('tanh', [
            make_ssub('E', make_text('k+1')),
            make_text('+'),
            make_d([
                make_text('-'),
                make_ssub('s', 'c'),
            ]),
        ]),
    ]
    return make_omath(*elements)

def build_formula_4():
    """恐慌放大: q_t^k = q_base × (1.0 + 0.5|E_t^k|)"""
    elements = [
        make_ssubsup('q', 't', 'k'),
        make_text('='),
        make_ssub('q', make_text('base')),
        make_text('×'),
        make_d([
            make_text('1.0+0.5'),
            make_abs([
                make_ssubsup('E', 't', 'k'),
            ]),
        ]),
    ]
    return make_omath(*elements)

def build_formula_5():
    """乐观缩减: q_t^k = D_t^k + (q_base - D_t^k)×(1.0 - 0.3·E_t^k)"""
    elements = [
        make_ssubsup('q', 't', 'k'),
        make_text('='),
        make_ssubsup('D', 't', 'k'),
        make_text('+'),
        make_d([
            make_ssub('q', make_text('base')),
            make_text('-'),
            make_ssubsup('D', 't', 'k'),
        ]),
        make_text('×'),
        make_d([
            make_text('1.0-0.3·'),
            make_ssubsup('E', 't', 'k'),
        ]),
    ]
    return make_omath(*elements)

def build_formula_6():
    """情绪调节权重: w_s_eff = w_s^0·(1+|E_t^k|)"""
    elements = [
        make_ssub('w', make_text('s_eff')),
        make_text('='),
        make_ssubsup('w', 's', '0'),
        make_text('·'),
        make_d([
            make_text('1+'),
            make_abs([
                make_ssubsup('E', 't', 'k'),
            ]),
        ]),
    ]
    return make_omath(*elements)

def build_formula_7():
    """正向激励权重: w_b_eff = w_b^0·(1+E_t^k)"""
    elements = [
        make_ssub('w', make_text('b_eff')),
        make_text('='),
        make_ssubsup('w', 'b', '0'),
        make_text('·'),
        make_d([
            make_text('1+'),
            make_ssubsup('E', 't', 'k'),
        ]),
    ]
    return make_omath(*elements)

def build_formula_8():
    """综合奖励: r_t^k = fill_rate + bonus - w_s_eff·stockout_rate - w_h·max(0,NS)"""
    elements = [
        make_ssubsup('r', 't', 'k'),
        make_text('='),
        make_text('fill_rate'),
        make_text('+'),
        make_text('bonus'),
        make_text('-'),
        make_ssub('w', make_text('s_eff')),
        make_text('·'),
        make_text('stockout_rate'),
        make_text('-'),
        make_ssub('w', 'h'),
        make_text('·'),
        make_func('max', [make_text('0,NS')]),
    ]
    return make_omath(*elements)

def build_formula_9():
    """需求生成: D_t = d + ρ·D_{t-1} + ε_t"""
    elements = [
        make_ssub('D', 't'),
        make_text('='),
        make_italic('d'),
        make_text('+'),
        make_italic('ρ'),
        make_text('·'),
        make_ssub('D', make_text('t-1')),
        make_text('+'),
        make_ssub('ε', 't'),
    ]
    return make_omath(*elements)

def build_formula_10():
    """EWC正则损失: L_ewc = ½·λ·Σ F_i·(θ_i - θ*_i)²"""
    elements = [
        make_ssub('L', make_text('ewc')),
        make_text('='),
        make_text('½·'),
        make_italic('λ'),
        make_text('·'),
        make_nary('∑', [make_italic('i')], [],
            [make_ssub('F', 'i'),
             make_text('·'),
             make_d([
                 make_ssub('θ', 'i'),
                 make_text('-'),
                 make_ssup(make_ssub('θ', 'i'), '*'),
             ]),
             make_ssup('', '2'),
            ]),
    ]
    return make_omath(*elements)

def build_formula_11():
    """情绪感知噪声: E_perceived = clip(E_true + N(0, σ_noise), -1, 1)"""
    elements = [
        make_ssub('E', make_text('perceived')),
        make_text('='),
        make_func('clip', [
            make_ssub('E', make_text('true')),
            make_text('+'),
            make_func('N', [
                make_text('0,'),
                make_ssub('σ', make_text('noise')),
            ]),
            make_text(',-1,1'),
        ]),
    ]
    return make_omath(*elements)

def build_formula_12():
    """覆盖率: c_r = NS_t^k / forecast"""
    elements = [
        make_ssub('c', 'r'),
        make_text('='),
        make_frac(
            [make_ssubsup('NS', 't', 'k')],
            [make_text('forecast')]
        ),
    ]
    return make_omath(*elements)

# ============================================================
# 文档格式辅助
# ============================================================

def set_three_line_table(table):
    """设置三线表格式"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # 移除所有边框
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(_w('val'), 'nil')
        borders.append(border)
    tblPr.append(borders)

    # 设置表格宽度
    tblW = OxmlElement('w:tblW')
    tblW.set(_w('type'), 'pct')
    tblW.set(_w('w'), '5000')  # 100%
    tblPr.append(tblW)

    # 为第一行和最后一行添加边框
    for i, row in enumerate(table.rows):
        tc = row._tr
        tcPr = tc.find(_w('tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)

        borders_tc = OxmlElement('w:tcBorders')
        if i == 0:
            # 第一行：顶部粗线 + 底部细线
            top = OxmlElement('w:top')
            top.set(_w('val'), 'single')
            top.set(_w('sz'), '12')
            top.set(_w('color'), '000000')
            borders_tc.append(top)
            bottom = OxmlElement('w:bottom')
            bottom.set(_w('val'), 'single')
            bottom.set(_w('sz'), '6')
            bottom.set(_w('color'), '000000')
            borders_tc.append(bottom)
        elif i == len(table.rows) - 1:
            # 最后一行：底部粗线
            bottom = OxmlElement('w:bottom')
            bottom.set(_w('val'), 'single')
            bottom.set(_w('sz'), '12')
            bottom.set(_w('color'), '000000')
            borders_tc.append(bottom)
        tcPr.append(borders_tc)

def add_three_line_table_header(table, headers):
    """添加三线表表头行"""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.bold = True
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(_w('rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(_w('eastAsia'), '宋体')

def add_three_line_table_row(table, row_data, align='left'):
    """添加三线表数据行"""
    row = table.add_row()
    for i, data in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(data)
        run.font.size = Pt(9)
        run.font.name = '宋体'
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(_w('rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(_w('eastAsia'), '宋体')

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(_w('rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(_w('eastAsia'), '黑体')
    return heading

def add_algorithm_line(doc, line_num, text_parts, indent=0):
    """
    添加算法伪代码行
    text_parts: list of (type, content) tuples
        type: 'text', 'code', 'omml'
        content: str for text/code, omath element for omml
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)

    # 行号
    if line_num:
        num_str = f'{line_num:2d}. '
        run = p.add_run(num_str)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'

    # 缩进
    if indent > 0:
        indent_str = '    ' * indent
        run = p.add_run(indent_str)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'

    # 内容
    for part_type, content in text_parts:
        if part_type == 'text':
            run = p.add_run(content)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
        elif part_type == 'code':
            run = p.add_run(content)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0, 0, 128)
        elif part_type == 'omml':
            p._p.append(content)

    return p

# ============================================================
# 主文档生成
# ============================================================

def generate_docx():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(_w('rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(_w('eastAsia'), '宋体')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ========== 标题 ==========
    title = doc.add_heading('核心算法表对照与修正', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.font.size = Pt(16)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(_w('rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(_w('eastAsia'), '黑体')

    doc.add_paragraph()

    # ========== 第一节：判定 ==========
    add_heading(doc, '一、判定：不完全一致，存在7处关键差异', level=1)

    p = doc.add_paragraph()
    add_text_run(p, '用户提供的算法表与实验代码（')
    add_text_run(p, 'marl_supply_chain_env.py', font_name='Consolas', size=9)
    add_text_run(p, '、')
    add_text_run(p, 'emotion_module.py', font_name='Consolas', size=9)
    add_text_run(p, '、')
    add_text_run(p, 'dynamic_events.py', font_name='Consolas', size=9)
    add_text_run(p, '、')
    add_text_run(p, 'run_exp2_20k.py', font_name='Consolas', size=9)
    add_text_run(p, '）逐行对照后，发现以下7处关键差异。')

    # ========== 差异对照表（三线表）==========
    add_heading(doc, '表1 算法差异对照表', level=2)

    table1 = doc.add_table(rows=1, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers1 = ['行号', '用户伪代码', '实际代码', '差异类型']
    add_three_line_table_header(table1, headers1)

    rows1 = [
        ['8', 'E_t^k ← E_t^k + η(E_t^{k-1} - E_t^k) 线性插值传染', 'E_{k+1} = tanh(E_{k+1} + (-s_c)) 固定冲击注入，需 stockout_rate > θ 触发', '公式错误'],
        ['10-12', 'DQN ε-greedy动作选择', 'Exp_2实际使用 RationalAgent.decide() + 情绪调节，非DQN', '机制不符'],
        ['17', 'If NS > 5×NS_avg: a=0 惩罚机制', 'MARLSupplyChainEnv 中未实现，由 holding_penalty 替代', '机制缺失'],
        ['21', 'r = fill_rate + bonus - holding_penalty', 'r = fill_rate + bonus - stockout_penalty - holding_penalty', '公式不完整'],
        ['25', 'PER采样+DQN更新', 'Exp_2中不使用DQN/PER，无Q网络更新', '机制不符'],
        ['27', '传递 q_t 与 E_t', '传递 forecast、inventory、emotion_label', '内容不符'],
        ['—', '无动态事件触发', '需求突变53次+供应中断23次+情绪传染2830次', '机制遗漏'],
    ]
    for row_data in rows1:
        add_three_line_table_row(table1, row_data, align='left')

    set_three_line_table(table1)

    doc.add_paragraph()

    # ========== 第二节：修正后的核心算法表 ==========
    add_heading(doc, '二、修正后的核心算法表', level=1)

    p = doc.add_paragraph()
    add_text_run(p, '以下为基于实验代码修正后的多智能体情绪感知人智协同决策系统（MAEA-HCDS）核心算法表，公式采用Word原生OMML可编辑格式。')

    doc.add_paragraph()

    # 算法标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('算法1  多智能体情绪感知人智协同决策系统（MAEA-HCDS）')
    run.font.size = Pt(10)
    run.font.name = '黑体'
    run.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(_w('rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(_w('eastAsia'), '黑体')

    # 算法输入输出
    p = doc.add_paragraph()
    add_code_run(p, '输入：仿真周期T=20000，需求参数(d=10, ρ=0.5, σ_ε=5)，提前期L=2')

    p = doc.add_paragraph()
    add_code_run(p, '      情绪参数(α=0.7, γ=2.0, w_s=1.0, w_m=0.5, w_e=0.3)')

    p = doc.add_paragraph()
    add_code_run(p, '      传染参数(θ=0.3, p_c=0.3, s_c=0.4)')

    p = doc.add_paragraph()
    add_code_run(p, '      动态事件概率(p_d=0.00265, p_s=0.00115)')

    p = doc.add_paragraph()
    add_code_run(p, '      成本参数(h=1.0, b=2.0)，观测维度=8')

    p = doc.add_paragraph()
    add_code_run(p, '输出：各节点订货序列{q_t^k}，情绪序列{E_t^k}，绩效指标(BWE, SL, Cost)')

    # 分隔线
    p = doc.add_paragraph()
    add_code_run(p, '─' * 70)

    # 算法主体
    # 第1行
    add_algorithm_line(doc, 1, [
        ('text', '初始化MARLSupplyChainEnv(PettingZoo AECEnv)，创建4个SupplyChainAgent'),
    ])

    # 第2行
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '    （零售商k=1 → 批发商k=2 → 分销商k=3 → 制造商k=4）')

    # 第2行
    add_algorithm_line(doc, 2, [
        ('text', '为每个Agent独立挂载EmotionState(α, γ, w'),
        ('text', 's'),
        ('text', ', w'),
        ('text', 'm'),
        ('text', ', w'),
        ('text', 'e'),
        ('text', ')'),
    ])

    # 第3行
    add_algorithm_line(doc, 3, [
        ('text', '初始化CommunicationChannel(拓扑=下游→上游，共享模式=forecast+inventory)'),
    ])

    # 第4行
    add_algorithm_line(doc, 4, [
        ('text', '初始化DynamicEventTrigger(p'),
        ('text', 'd'),
        ('text', ', p'),
        ('text', 's'),
        ('text', ', θ, p'),
        ('text', 'c'),
        ('text', ', s'),
        ('text', 'c'),
        ('text', ', seed=42)'),
    ])

    # 第5行
    add_algorithm_line(doc, 5, [
        ('text', 'For t = 1 : T do'),
    ])

    # 第6-7行
    add_algorithm_line(doc, 6, [
        ('code', '  // —— 动态事件触发 ——'),
    ], indent=0)

    add_algorithm_line(doc, 7, [
        ('text', '  触发需求突变(m∈{2.0, 0.5})与维护供应中断状态(持续3-5周期)'),
    ])

    # 第8行 - 需求生成公式 D_t = d + ρ·D_{t-1} + ε_t
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, ' 8.   生成顾客需求 ')
    p._p.append(build_formula_9())
    add_code_run(p, '；若需求突变则 ')

    # 用OMML构建 D_t ← D_t × m
    omath_dm = make_omath(
        make_ssub('D', 't'),
        make_text('←'),
        make_ssub('D', 't'),
        make_text('×'),
        make_italic('m'),
    )
    p._p.append(omath_dm)

    # 第9行
    add_algorithm_line(doc, 9, [
        ('text', '  For k = 1 : 4 do  // AECEnv顺序决策：下游先决策，订单逐级上传'),
    ])

    # 第10-11行
    add_algorithm_line(doc, 10, [
        ('code', '    // —— 库存更新与履约 ——'),
    ])

    add_algorithm_line(doc, 11, [
        ('text', '    收到上游L期前到货（k=4制造商：断供期间到货=0）'),
    ])

    # 第12行
    add_algorithm_line(doc, 12, [
        ('text', '    履约下游需求：fulfilled = min(NS, D'),
        ('text', 't'),
        ('text', '^k), stockout = max(0, D'),
        ('text', 't'),
        ('text', '^k - fulfilled)'),
    ])

    # 第13-14行 - 情绪演化公式
    add_algorithm_line(doc, 13, [
        ('code', '    // —— 行为感知：情绪演化（式1-2）——'),
    ])

    # 第14行 - Φ_t公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '14.     ')
    p._p.append(build_formula_2())

    # 第15行 - E_t^k公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '15.     ')
    p._p.append(build_formula_1())

    # 第16-17行 - 情绪传染
    add_algorithm_line(doc, 16, [
        ('code', '    // —— 情绪传染（式8，周期末由DynamicEventTrigger执行）——'),
    ])

    # 第17行
    add_algorithm_line(doc, 17, [
        ('text', '    If stockout_rate > θ AND random() < p'),
        ('text', 'c'),
        ('text', ':'),
    ])

    # 第18行 - 传染公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '18.       ')
    p._p.append(build_formula_3())
    add_code_run(p, '  // 固定恐慌冲击注入上游')

    # 第19-20行 - 传授经验
    add_algorithm_line(doc, 19, [
        ('code', '    // —— 人机协同机制1：传授经验（理性决策作为基础策略）——'),
    ])

    add_algorithm_line(doc, 20, [
        ('text', '    q'),
        ('text', 'base'),
        ('text', ' = SMA(D'),
        ('text', 't-p:t-1'),
        ('text', ') + z·σ'),
        ('text', 'L'),
        ('text', '·√(L+1)  // OUT策略'),
    ])

    # 第21行 - 情绪调节注释
    add_algorithm_line(doc, 21, [
        ('code', '    // —— 情绪调节订货量（式5-6）——'),
    ])

    # 第22行 - 恐慌放大公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '22.     If E')
    # E_t^k
    omath_e = make_omath(make_ssubsup('E', 't', 'k'))
    p._p.append(omath_e)
    add_code_run(p, ' < 0: ')
    p._p.append(build_formula_4())
    add_code_run(p, '  // 恐慌放大')

    # 第23行 - 乐观缩减公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '23.     Elif E')
    p._p.append(make_omath(make_ssubsup('E', 't', 'k')))
    add_code_run(p, ' > 0: ')
    p._p.append(build_formula_5())
    add_code_run(p, '  // 乐观缩减')

    # 第24行 - 情绪调节奖励权重注释
    add_algorithm_line(doc, 24, [
        ('code', '    // —— 情绪调节奖励权重（式3-4）——'),
    ])

    # 第25行 - w_s_eff公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '25.     ')
    p._p.append(build_formula_6())
    add_code_run(p, '  // 恐慌放大缺货惩罚')

    # 第26行 - w_b_eff公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '26.     ')
    p._p.append(build_formula_7())
    add_code_run(p, '  // 乐观放大正向激励')

    # 第27行 - 正向激励注释
    add_algorithm_line(doc, 27, [
        ('code', '    // —— 人机协同机制2：正向激励目标重塑（式7）——'),
    ])

    # 第28行 - 覆盖率公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '28.     ')
    p._p.append(build_formula_12())
    add_code_run(p, '  // 库存覆盖率')

    # 第29行
    add_algorithm_line(doc, 29, [
        ('text', '    If c'),
        ('text', 'r'),
        ('text', ' ∈ [0.8, 1.5]: match_factor = max(0, 1-|c'),
        ('text', 'r'),
        ('text', '-1|), bonus = w'),
        ('text', 'b_eff'),
        ('text', '·match_factor'),
    ])

    # 第30行 - 综合奖励注释
    add_algorithm_line(doc, 30, [
        ('code', '    // —— 综合奖励（式7）——'),
    ])

    # 第31行 - 综合奖励公式
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '31.     ')
    p._p.append(build_formula_8())

    # 第32-34行 - 协同通信
    add_algorithm_line(doc, 32, [
        ('code', '    // —— 人机协同机制3：协同通信（信息共享）——'),
    ])

    add_algorithm_line(doc, 33, [
        ('text', '    广播{forecast, inventory, emotion_label}至通信通道'),
    ])

    add_algorithm_line(doc, 34, [
        ('text', '    上游k+1从通信通道接收下游k的forecast与inventory（观测第6-7维）'),
    ])

    # 第35-36行 - 局部观测
    add_algorithm_line(doc, 35, [
        ('code', '    // —— 构建局部观测（8维）——'),
    ])

    add_algorithm_line(doc, 36, [
        ('text', '    s'),
        ('text', 't'),
        ('text', '^k = [NS, WIP, D'),
        ('text', 'down'),
        ('text', ', arrival, D'),
        ('text', 'last'),
        ('text', ', E'),
        ('text', 't'),
        ('text', ', forecast'),
        ('text', 'shared'),
        ('text', ', inventory'),
        ('text', 'shared'),
        ('text', ']'),
    ])

    # 第37-38行
    add_algorithm_line(doc, 37, [
        ('text', '  End for'),
    ])

    add_algorithm_line(doc, 38, [
        ('text', 'End for'),
    ])

    # 分隔线
    p = doc.add_paragraph()
    add_code_run(p, '─' * 70)

    # 注释
    p = doc.add_paragraph()
    add_code_run(p, '注1：在Exp_1/Exp_1b中，分销商(k=3)使用ContinualIDMRAgent(DQN+PER+EWC)替代')

    p = doc.add_paragraph()
    add_code_run(p, '    理性决策，第20-23行替换为ε-greedy动作选择→情绪调节→PER优先级采样→')

    p = doc.add_paragraph()
    add_code_run(p, '    DQN梯度更新→EWC正则约束，每target_update_freq步更新目标网络。')

    # 注2 - 情绪感知噪声公式
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '注2：情绪感知噪声 ')
    p._p.append(build_formula_11())
    add_code_run(p, '（式15），')

    p = doc.add_paragraph()
    add_code_run(p, '    σ')

    # σ_noise 下标
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = Pt(14)
    add_code_run(p2, '    noise=0.15时应用于Q网络输入第6维，真实情绪仍用于环境动力学。')

    # 注3 - EWC正则损失公式
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(14)
    add_code_run(p, '注3：EWC正则损失 ')
    p._p.append(build_formula_10())
    add_code_run(p, '（式16），λ=400，')

    p = doc.add_paragraph()
    add_code_run(p, '    任务切换前调用consolidate_knowledge()计算Fisher信息矩阵对角元。')

    doc.add_paragraph()

    # ========== 第三节：修正要点说明 ==========
    add_heading(doc, '三、修正要点说明', level=1)

    # 表2: 情绪传染公式修正
    add_heading(doc, '表2 情绪传染公式修正（第17-18行）', level=2)
    table2 = doc.add_table(rows=1, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_three_line_table_header(table2, ['项目', '用户原版', '修正版'])
    rows2 = [
        ['触发条件', '无条件概率η', 'stockout_rate > θ(=0.3) AND random() < p_c(=0.3)'],
        ['传染方式', '线性插值 E + η(E_{k-1} - E)', '固定冲击 tanh(E + (-s_c))，s_c=0.4'],
        ['传染方向', '未明确', '下游k → 上游k+1（单向）'],
        ['代码位置', '—', 'dynamic_events.py L200-278'],
    ]
    for row_data in rows2:
        add_three_line_table_row(table2, row_data)
    set_three_line_table(table2)

    doc.add_paragraph()

    # 表3: 决策机制修正
    add_heading(doc, '表3 决策机制修正（第20-23行）', level=2)
    table3 = doc.add_table(rows=1, cols=3)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_three_line_table_header(table3, ['项目', '用户原版', '修正版'])
    rows3 = [
        ['决策方式', 'DQN ε-greedy', 'RationalAgent(SMA+OUT) + 情绪调节'],
        ['适用实验', '暗示Exp_2使用DQN', 'Exp_2使用理性决策；Exp_1/1b使用DQN'],
        ['代码位置', '—', 'run_exp2_20k.py L140-150'],
    ]
    for row_data in rows3:
        add_three_line_table_row(table3, row_data)
    set_three_line_table(table3)

    doc.add_paragraph()

    # 表4: 奖励函数修正
    add_heading(doc, '表4 奖励函数修正（第31行）', level=2)
    table4 = doc.add_table(rows=1, cols=3)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_three_line_table_header(table4, ['项目', '用户原版', '修正版'])
    rows4 = [
        ['奖励公式', 'fill_rate + bonus - holding_penalty', 'fill_rate + bonus - w_s_eff·stockout_rate - w_h·max(0,NS)'],
        ['缺失项', '缺货惩罚', '补充情绪调节的缺货惩罚项'],
        ['代码位置', '—', 'marl_supply_chain_env.py L435-506'],
    ]
    for row_data in rows4:
        add_three_line_table_row(table4, row_data)
    set_three_line_table(table4)

    doc.add_paragraph()

    # 表5: 协同通信内容修正
    add_heading(doc, '表5 协同通信内容修正（第33-34行）', level=2)
    table5 = doc.add_table(rows=1, cols=3)
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_three_line_table_header(table5, ['项目', '用户原版', '修正版'])
    rows5 = [
        ['传递内容', 'q_t（订货量）与 E_t（情绪值）', 'forecast（预测需求）、inventory（库存）、emotion_label（情绪标签）'],
        ['代码位置', '—', 'marl_supply_chain_env.py L556-573'],
    ]
    for row_data in rows5:
        add_three_line_table_row(table5, row_data)
    set_three_line_table(table5)

    doc.add_paragraph()

    # 表6: 新增动态事件触发
    add_heading(doc, '表6 新增动态事件触发（第6-8行）', level=2)
    table6 = doc.add_table(rows=1, cols=3)
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_three_line_table_header(table6, ['事件类型', '次数', '代码位置'])
    rows6 = [
        ['需求突变（D_t × {2.0, 0.5}）', '53次', 'dynamic_events.py L110-165'],
        ['供应中断（持续3-5周期，制造商到货=0）', '23次', 'dynamic_events.py L130-175'],
        ['情绪传染（30%概率上游传染）', '2830次', 'dynamic_events.py L200-278'],
    ]
    for row_data in rows6:
        add_three_line_table_row(table6, row_data, align='left')
    set_three_line_table(table6)

    doc.add_paragraph()

    # 表7: 新增8维局部观测
    add_heading(doc, '表7 8维局部观测构建（第36行）', level=2)
    table7 = doc.add_table(rows=1, cols=3)
    table7.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_three_line_table_header(table7, ['维度', '符号', '含义'])
    rows7 = [
        ['1', 'NS', '当前净库存（归一化）'],
        ['2', 'WIP', '在途库存（归一化）'],
        ['3', 'D_down', '下游节点上期订单'],
        ['4', 'arrival', '上游到货量'],
        ['5', 'D_last', '上期需求'],
        ['6', 'E_t', '当前情绪状态'],
        ['7', 'forecast_shared', '下游共享的预测需求（协同通信）'],
        ['8', 'inventory_shared', '下游共享的库存水平（协同通信）'],
    ]
    for row_data in rows7:
        add_three_line_table_row(table7, row_data, align='center')
    set_three_line_table(table7)

    doc.add_paragraph()

    # 表8: 惩罚机制修正
    add_heading(doc, '表8 惩罚机制修正', level=2)
    table8 = doc.add_table(rows=1, cols=3)
    table8.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_three_line_table_header(table8, ['项目', '用户原版', '修正版'])
    rows8 = [
        ['惩罚方式', 'If NS > 5×NS_avg: a=0（强制零订货）', 'holding_penalty = w_h·max(0, NS)，w_h=0.0001（软惩罚）'],
        ['适用环境', '单智能体IDMRSupplyChainEnv', '多智能体MARLSupplyChainEnv'],
        ['代码位置', 'idmr_agent.py L380-395', 'marl_supply_chain_env.py L435-506'],
    ]
    for row_data in rows8:
        add_three_line_table_row(table8, row_data)
    set_three_line_table(table8)

    # 保存文档
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        '核心算法表_对照与修正.docx'
    )
    doc.save(output_path)
    print(f'文档已保存: {output_path}')
    return output_path

if __name__ == '__main__':
    generate_docx()
