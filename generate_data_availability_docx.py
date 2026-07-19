# -*- coding: utf-8 -*-
"""
向论文主稿追加"数据可用性声明"章节
==================================
目标文件: 情绪扰动、激励阻断与协同鲁棒：多智能体供应链牛鞭效应缓解研究_719修改后.docx
插入位置: 参考文献之后, 章节标题"数据可用性声明"
输出文件: 情绪扰动、激励阻断与协同鲁棒：多智能体供应链牛鞭效应缓解研究_719修改后.docx (原地追加)

格式:
  - A4竖版, 宋体10.5pt正文
  - 黑体三号一级标题
  - Times New Roman 英文字体
  - 论文规范段落行距
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 字体样式常量
# ============================================================
FONT_CN_BODY = '宋体'
FONT_CN_HEADING = '黑体'
FONT_EN = 'Times New Roman'
COLOR_HEADING = RGBColor(0x00, 0x00, 0x00)  # 黑色
COLOR_BODY = RGBColor(0x00, 0x00, 0x00)
COLOR_LINK = RGBColor(0x05, 0x63, 0xC1)  # 超链接蓝色


# ============================================================
# 辅助函数
# ============================================================
def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN_BODY,
                 size=10.5, bold=False, color=None, italic=False):
    """设置 run 的字体 (中英文分别设置)"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent_chars=2, line_spacing=1.5,
                         space_before=0, space_after=0):
    """设置段落格式: 对齐、首行缩进、行距、段前段后"""
    pf = para.paragraph_format
    pf.alignment = alignment
    if first_line_indent_chars > 0:
        # 首行缩进 N 字符 (按 10.5pt 字号计算)
        pf.first_line_indent = Pt(10.5 * first_line_indent_chars)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_heading(doc, text, level=1):
    """添加章节标题 (黑体, 一级三号=16pt, 二号小四=12pt)"""
    para = doc.add_paragraph()
    if level == 1:
        # 一级标题: 黑体三号(16pt), 居中, 段前段后24pt
        run = para.add_run(text)
        set_run_font(run, font_cn=FONT_CN_HEADING, size=16, bold=True,
                     color=COLOR_HEADING)
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             first_line_indent_chars=0,
                             line_spacing=1.5,
                             space_before=24, space_after=18)
    elif level == 2:
        # 二级标题: 黑体小四(12pt), 左对齐, 段前段后12pt
        run = para.add_run(text)
        set_run_font(run, font_cn=FONT_CN_HEADING, size=12, bold=True,
                     color=COLOR_HEADING)
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             first_line_indent_chars=0,
                             line_spacing=1.5,
                             space_before=12, space_after=6)
    return para


def add_body_paragraph(doc, text, indent=True):
    """添加正文段落 (宋体10.5pt, 1.5倍行距, 首行缩进2字符)"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_cn=FONT_CN_BODY, font_en=FONT_EN,
                 size=10.5, color=COLOR_BODY)
    set_paragraph_format(para,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent_chars=2 if indent else 0,
                         line_spacing=1.5,
                         space_before=0, space_after=0)
    return para


def add_body_paragraph_with_runs(doc, runs_config, indent=True):
    """添加含多种格式的正文段落
    runs_config: [(text, {'font_cn':..., 'font_en':..., 'size':..., 'bold':..., 'italic':..., 'color':...}), ...]
    """
    para = doc.add_paragraph()
    for text, fmt in runs_config:
        run = para.add_run(text)
        set_run_font(run,
                     font_en=fmt.get('font_en', FONT_EN),
                     font_cn=fmt.get('font_cn', FONT_CN_BODY),
                     size=fmt.get('size', 10.5),
                     bold=fmt.get('bold', False),
                     italic=fmt.get('italic', False),
                     color=fmt.get('color', COLOR_BODY))
    set_paragraph_format(para,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent_chars=2 if indent else 0,
                         line_spacing=1.5,
                         space_before=0, space_after=0)
    return para


def add_bullet_paragraph(doc, text):
    """添加项目符号段落 (宋体10.5pt, 1.5倍行距, 无首行缩进)"""
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    set_run_font(run, font_cn=FONT_CN_BODY, font_en=FONT_EN,
                 size=10.5, color=COLOR_BODY)
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return para


def add_page_break(doc):
    """添加分页符"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


# ============================================================
# 数据可用性声明内容生成
# ============================================================
def append_data_availability_statement(doc):
    """向论文末尾追加数据可用性声明章节"""

    # ---- 一级标题 ----
    add_heading(doc, "数据可用性声明", level=1)

    # ---- 总声明段落 ----
    add_body_paragraph_with_runs(doc, [
        ("为促进学术透明与可复现研究，本研究全部仿真代码、实验数据生成脚本、"
         "预训练模型与论文配图已在 GitHub 公开仓库发布，遵循 MIT 开源许可证。"
         "仓库地址：", {}),
        ("https://github.com/lixiangyang12/MAEA-HCDS",
         {'color': COLOR_LINK, 'italic': False}),
        ("（版本 v719，提交标识 3790212）。其他研究者可通过上述地址获取全部工程文件，"
         "实现本文实验结果的完整复现与二次开发。", {}),
    ])

    # ---- 二级标题: 数据来源 ----
    add_heading(doc, "1. 数据来源", level=2)

    add_body_paragraph(doc,
        "本研究未使用任何真实世界数据，全部实验数据均通过计算机仿真生成。"
        "仿真环境基于李勇等（2022）提出的四级供应链牛鞭效应模型构建，"
        "需求过程采用 AR(1) 自回归模型 D_t = d + ρ·D_{t-1} + ε_t，"
        "其中 d = 10、ρ = 0.5、ε ~ N(0, 5²)，订货提前期 L = 2，"
        "供应链层级 K = 4（零售商、批发商、分销商、制造商）。"
        "全部仿真参数详见仓库根目录的 config.yaml 配置文件。"
    )

    add_body_paragraph_with_runs(doc, [
        ("为确保实验结果完全可复现，本研究采用固定随机种子 ", {}),
        ("seed = 42", {'font_en': FONT_EN, 'bold': True}),
        ("，覆盖 NumPy 随机数生成器、DQN 经验回放采样、"
         "动态事件触发（需求突变、供应中断、情绪传染）等全部随机过程。"
         "在相同硬件与软件环境下，重复运行将得到数值完全一致的实验结果。", {}),
    ])

    # ---- 二级标题: 代码与模型 ----
    add_heading(doc, "2. 代码与预训练模型", level=2)

    add_body_paragraph(doc,
        "仓库包含 68 个 Python 脚本，覆盖仿真环境、智能体决策、持续学习、"
        "实验运行与论文配图生成等全部模块。核心模块包括：四级供应链仿真环境"
        "（supply_chain_env.py）、IDMR 智能体（idmr_agent.py）、情绪演化模块"
        "（emotion_module.py）、多智能体协同环境（marl_supply_chain_env.py）、"
        "动态事件触发器（dynamic_events.py）、优先级经验回放（prioritized_replay.py）、"
        "弹性权重巩固（ewc.py）、持续学习智能体（continual_idmr.py）以及"
        "三组对比实验自动化运行脚本（batch_runner.py）。"
    )

    add_body_paragraph_with_runs(doc, [
        ("仓库同时提供预训练模型 ", {}),
        ("idmr_model.pkl", {'font_en': FONT_EN, 'bold': True}),
        ("，保存了 DQN 神经网络权重、经验回放池状态与训练超参数。"
         "其他研究者可在无需重新训练的条件下，直接加载该模型完成全部评估实验。"
         "深度 Q 网络采用纯 NumPy 实现，前向传播与反向传播均不依赖 PyTorch、"
         "TensorFlow 等深度学习框架，降低了复现门槛。", {}),
    ])

    # ---- 二级标题: 复现路径 ----
    add_heading(doc, "3. 复现路径", level=2)

    add_body_paragraph(doc,
        "结合不同研究需求，本仓库提供三条复现路径："
    )

    add_body_paragraph_with_runs(doc, [
        ("（1）预训练模型快速复现（约 1 分钟）：", {'bold': True}),
        ("直接加载 idmr_model.pkl，运行 batch_runner.py 即可复现"
         "全部四组对比实验的关键指标，包括牛鞭效应方差比、系统平均成本、"
         "服务水平、情绪波动指数与协同收益。", {}),
    ])

    add_body_paragraph_with_runs(doc, [
        ("（2）完整训练复现（约 30 分钟）：", {'bold': True}),
        ("删除 idmr_model.pkl 后运行 batch_runner.py，从零开始训练 DQN 智能体。"
         "由于随机种子固定，所得结果与预训练模型评估结果完全一致，"
         "可验证训练过程的确定性。", {}),
    ])

    add_body_paragraph_with_runs(doc, [
        ("（3）图表演示（无需训练）：", {'bold': True}),
        ("仓库 svg_figures_paper_719/ 目录直接提供论文图 1（系统机制图）、"
         "图 2（多智能体决策流程图）与图 8（情绪传染网络图）的 PDF、SVG 与 PNG "
         "三种格式文件，其他研究者可不经任何计算直接查阅与引用。", {}),
    ])

    # ---- 二级标题: 依赖环境 ----
    add_heading(doc, "4. 依赖环境", level=2)

    add_body_paragraph_with_runs(doc, [
        ("本仓库代码在 Python 3.9 及以上版本运行，依赖库详见仓库根目录的 ", {}),
        ("requirements.txt", {'font_en': FONT_EN, 'bold': True}),
        (" 文件，主要包括：NumPy（1.21 及以上，数值计算与 DQN 实现）、"
         "Matplotlib（3.5 及以上，学术级矢量图表绘制）、NetworkX（2.6 及以上，"
         "情绪传染路径可视化）、Pandas（1.3 及以上，实验数据分析）、"
         "PettingZoo（1.21 及以上，多智能体环境框架）、SciPy（1.7 及以上，"
         "统计分析）、PyYAML（5.4 及以上，配置文件加载）。", {}),
    ])

    add_body_paragraph(doc,
        "推荐使用 Anaconda 或 Miniconda 创建独立虚拟环境以避免依赖冲突。"
        "安装命令为：pip install -r requirements.txt。"
        "本仓库已在 Windows 10/11、Ubuntu 20.04 与 macOS 12 平台完成兼容性测试。"
    )

    # ---- 二级标题: 许可证 ----
    add_heading(doc, "5. 许可证", level=2)

    add_body_paragraph_with_runs(doc, [
        ("本仓库代码与配套实验数据采用 ", {}),
        ("MIT 许可证", {'bold': True}),
        ("（详见 LICENSE 文件）开源。MIT 许可证允许其他研究者自由复制、修改、"
         "合并、发布、分发、再授权与销售本仓库代码，仅需在所有副本中保留版权声明"
         "与许可声明。论文正文、图表与文字内容不在 MIT 许可证范围内，"
         "其版权由作者保留，引用须遵循学术规范。", {}),
    ])

    # ---- 二级标题: 引用方式 ----
    add_heading(doc, "6. 引用方式", level=2)

    add_body_paragraph(doc,
        "若本仓库代码或实验方法对您的研究有帮助，请按以下 BibTeX 格式引用："
    )

    # BibTeX 引用块 (等宽字体)
    bibtex_para = doc.add_paragraph()
    bibtex_text = (
        "@article{yang2026maea_hcds,\n"
        "  title   = {情绪扰动、激励阻断与协同鲁棒：多智能体供应链牛鞭效应缓解研究},\n"
        "  author  = {杨理想 and 王晶},\n"
        "  journal = {中国管理科学 (投稿中)},\n"
        "  year    = {2026},\n"
        "  note    = {Doctoral Research, Management Science and Engineering,\n"
        "             Yanshan University, School of Economics and Management},\n"
        "  url     = {https://github.com/lixiangyang12/MAEA-HCDS}\n"
        "}"
    )
    run = bibtex_para.add_run(bibtex_text)
    set_run_font(run, font_en='Consolas', font_cn='宋体',
                 size=9, color=RGBColor(0x33, 0x33, 0x33))
    pf = bibtex_para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(21)  # 左缩进2字符
    pf.line_spacing = 1.2
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    add_body_paragraph_with_runs(doc, [
        ("仓库根目录同时提供 ", {}),
        ("CITATION.cff", {'font_en': FONT_EN, 'bold': True}),
        (" 元数据文件，GitHub 平台可自动识别并生成一键引用功能，"
         "支持 BibTeX、APA、Chicago 等多种引用格式。", {}),
    ])

    # ---- 二级标题: 联系方式 ----
    add_heading(doc, "7. 联系方式", level=2)

    add_body_paragraph_with_runs(doc, [
        ("通讯作者：杨理想，燕山大学经济与管理学院。", {}),
    ])

    add_body_paragraph_with_runs(doc, [
        ("邮箱：", {}),
        ("yanglixiang@stdu.ysu.edu.cn", {'color': COLOR_LINK}),
    ])

    add_body_paragraph_with_runs(doc, [
        ("若在复现过程中遇到任何问题，或对代码改进存在建议，"
         "欢迎在 GitHub 仓库的 Issues 页面提交反馈"
         "（https://github.com/lixiangyang12/MAEA-HCDS/issues），"
         "作者将及时回应并维护代码更新。", {}),
    ])

    # ---- 致谢声明 ----
    add_heading(doc, "8. 致谢", level=2)

    add_body_paragraph(doc,
        "感谢燕山大学经济与管理学院对本研究提供的学术支持。"
        "感谢李勇教授及其团队在牛鞭效应人机协同决策领域的前瞻性工作，"
        "为本研究提供了重要的理论框架与基准模型。"
        "感谢匿名审稿人对论文修改提出的宝贵意见。"
    )


# ============================================================
# 主函数
# ============================================================
def main():
    import docx.enum.text

    paper_path = r"c:\个人资料\申博材料\企业运营与科研管理数据库\情绪扰动、激励阻断与协同鲁棒：多智能体供应链牛鞭效应缓解研究_719修改后.docx"

    if not os.path.exists(paper_path):
        print(f"[错误] 论文主稿不存在: {paper_path}")
        return False

    # 备份原始文件
    backup_path = paper_path.replace('.docx', '_备份_数据可用性声明前.docx')
    if not os.path.exists(backup_path):
        shutil.copy2(paper_path, backup_path)
        print(f"[备份] 原始论文已备份至: {os.path.basename(backup_path)}")

    # 加载论文
    print(f"[加载] 正在加载论文: {os.path.basename(paper_path)}")
    doc = Document(paper_path)
    print(f"[加载] 完成, 当前段落数: {len(doc.paragraphs)}")

    # 检查是否已存在"数据可用性声明"章节
    existing_titles = [p.text.strip() for p in doc.paragraphs
                       if p.text.strip() and len(p.text.strip()) < 30]
    if "数据可用性声明" in existing_titles:
        print("[跳过] 论文中已存在'数据可用性声明'章节, 未重复添加")
        return True

    # 追加数据可用性声明
    print("[追加] 正在追加数据可用性声明章节...")
    append_data_availability_statement(doc)

    # 保存
    doc.save(paper_path)
    print(f"[保存] 完成, 新段落数: {len(doc.paragraphs)}")
    print(f"[输出] {os.path.basename(paper_path)}")
    print("")
    print("=" * 60)
    print("数据可用性声明已成功追加到论文末尾")
    print("=" * 60)
    return True


if __name__ == '__main__':
    import docx.enum.text
    main()
