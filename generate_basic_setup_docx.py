"""
实验基础设置 docx 生成脚本
=========================
生成"多智能体情绪感知供应链人智协同决策实验"基础设置 docx 文档
- 三线表格式
- OMML 原生公式格式
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# OMML 公式转换
# ============================================================
try:
    from latex2mathml.converter import convert as _latex_to_mathml
    from lxml import etree
    _XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
    _xslt_tree = etree.parse(_XSLT_PATH)
    _xslt_transform = etree.XSLT(_xslt_tree)
    _OMML_AVAILABLE = True
except Exception as e:
    _OMML_AVAILABLE = False
    print(f"[警告] OMML 不可用: {e}")


def latex_to_omml(latex_str, display='inline'):
    mathml_str = _latex_to_mathml(latex_str, display=display)
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = _xslt_transform(mathml_tree)
    return omml_tree.getroot()


# ============================================================
# 字体样式
# ============================================================
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_HEADING_CN = '黑体'
COLOR_HEADING = RGBColor(0x1F, 0x3A, 0x5F)
COLOR_FORMULA = RGBColor(0x00, 0x00, 0x80)


def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=10, bold=False, color=None):
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


# ============================================================
# 三线表样式
# ============================================================
def set_three_line_table(table):
    """三线表: 顶线粗、表头下线细、底线粗, 无竖线"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')

    # 顶线 - 粗 1.5pt
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)

    # 左线 - 无
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'nil')
    tblBorders.append(left)

    # 底线 - 粗 1.5pt
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)

    # 右线 - 无
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'nil')
    tblBorders.append(right)

    # 内部横线 - 无
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'nil')
    tblBorders.append(insideH)

    # 内部竖线 - 无
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'nil')
    tblBorders.append(insideV)

    tblPr.append(tblBorders)

    # 表头行下方细线
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '6')  # 0.75pt
        bottom_border.set(qn('w:space'), '0')
        bottom_border.set(qn('w:color'), '000000')
        existing_bottom = tcBorders.find(qn('w:bottom'))
        if existing_bottom is not None:
            tcBorders.remove(existing_bottom)
        tcBorders.append(bottom_border)


def set_cell_content(cell, text, bold=False, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格内容, 支持 $...$ 公式"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    pattern = r'(\$[^$]+\$)'
    segments = re.split(pattern, text)

    for seg in segments:
        if not seg:
            continue
        if seg.startswith('$') and seg.endswith('$'):
            formula = seg[1:-1]
            if _OMML_AVAILABLE:
                try:
                    omml = latex_to_omml(formula)
                    p._element.append(omml)
                    continue
                except Exception as e:
                    print(f"[提示] 公式转换失败: {formula} ({e})")
            run = p.add_run(formula)
            set_run_font(run, size=size, bold=True, color=COLOR_FORMULA)
            run.font.italic = True
        else:
            run = p.add_run(seg)
            set_run_font(run, size=size, bold=bold)


def set_cell_vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def add_table_caption(doc, caption_text):
    """添加表格标题（居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(caption_text)
    set_run_font(run, size=10, bold=True)


def create_table(doc, data, col_widths, first_col_left=False):
    """创建三线表"""
    n_rows = len(data)
    n_cols = len(data[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    # 填充数据
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            set_cell_vertical_center(cell)
            if i == 0:
                set_cell_content(cell, cell_text, bold=True, size=9)
            else:
                if first_col_left and j == 0:
                    set_cell_content(cell, cell_text, bold=False, size=9,
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
                elif first_col_left and j == 1:
                    set_cell_content(cell, cell_text, bold=False, size=9,
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
                else:
                    set_cell_content(cell, cell_text, bold=False, size=9)

    set_three_line_table(table)
    return table


# ============================================================
# 创建文档
# ============================================================
def create_basic_setup_docx():
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)

    # ============================================================
    # 标题
    # ============================================================
    heading = doc.add_heading(level=1)
    run = heading.add_run('多智能体情绪感知供应链人智协同决策实验')
    run.font.name = FONT_HEADING_CN
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_HEADING_CN)

    heading2 = doc.add_heading(level=2)
    run = heading2.add_run('实验基础设置')
    run.font.name = FONT_HEADING_CN
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_HEADING_CN)

    # 摘要
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('本节系统阐述多智能体情绪感知供应链人智协同决策实验的基础设置，涵盖仿真环境、供应链结构、需求生成、决策算法、情绪与协同机制、持续学习等关键配置参数，确保实验结果可量化、可复现。')
    set_run_font(run, size=11)

    # ============================================================
    # 表1: 仿真环境配置
    # ============================================================
    add_table_caption(doc, '表1  仿真环境配置')
    table1_data = [
        ['配置类别', '项目', '规格/版本', '说明'],
        ['软件环境', '操作系统', 'Windows 11', '64位'],
        ['', 'Python', '3.10+', '主要编程语言'],
        ['', '深度学习框架', 'PyTorch 2.0+', 'DQN神经网络'],
        ['', '仿真框架', 'PettingZoo AECEnv', '多智能体环境'],
        ['', '数据处理', 'NumPy 1.24+, Pandas 2.0+', '数值计算'],
        ['', '可视化', 'Matplotlib 3.7+, NetworkX', '图表绘制'],
        ['硬件环境', 'CPU', 'Intel Core i7', '8核以上'],
        ['', '内存', '16 GB+', '—'],
        ['', 'GPU', '可选（CUDA支持）', 'DQN训练加速'],
    ]
    create_table(doc, table1_data,
                col_widths=[Cm(2.0), Cm(3.0), Cm(4.5), Cm(5.0)],
                first_col_left=True)

    # ============================================================
    # 表2: 供应链结构配置
    # ============================================================
    add_table_caption(doc, '表2  供应链结构配置')
    table2_data = [
        ['参数', '符号', '取值', '说明'],
        ['供应链层级', '—', '4级', '串联结构'],
        ['节点编号', '$k$', '1, 2, 3, 4', '零售商→批发商→分销商→制造商'],
        ['节点名称', '—', 'Retailer, Wholesaler, Distributor, Manufacturer', '四级啤酒游戏模型'],
        ['特殊节点', '—', '$k=3$（分销商）', 'IDMR智慧决策机器人部署节点'],
        ['运输延迟', '$L$', '2 周期', '订货至收货提前期'],
        ['信息延迟', '—', '0 周期', '下游需求实时可见（协同模式）'],
    ]
    create_table(doc, table2_data,
                col_widths=[Cm(3.0), Cm(1.5), Cm(5.5), Cm(4.5)],
                first_col_left=True)

    # ============================================================
    # 表3: 需求生成参数
    # ============================================================
    add_table_caption(doc, '表3  需求生成参数（AR(1)模型）')
    table3_data = [
        ['参数', '符号', '取值', '说明'],
        ['需求基础水平', '$d$', '10', 'AR(1)常数项'],
        ['自相关系数', '$\\rho$', '0.5', '$|\\rho|<1$保证平稳性'],
        ['误差标准差', '$\\sigma_\\varepsilon$', '5', '正态分布$N(0, \\sigma_\\varepsilon^2)$'],
        ['初始需求', '$D_0$', '20', '$D_0 = d/(1-\\rho)$'],
        ['需求模型', '—', '$D_t = d + \\rho D_{t-1} + \\varepsilon_t$', '一阶自回归过程'],
    ]
    create_table(doc, table3_data,
                col_widths=[Cm(3.0), Cm(2.0), Cm(4.0), Cm(5.5)],
                first_col_left=True)

    # ============================================================
    # 表4: 决策算法参数
    # ============================================================
    add_table_caption(doc, '表4  决策算法参数')
    table4_data = [
        ['决策方式', '算法', '关键参数', '取值'],
        ['理性决策\n（Baseline）', 'SMA + OUT', '移动平均窗口 $p$', '5'],
        ['', '', '安全库存系数 $z$', '2（97.7%服务水平）'],
        ['', '', '订至点策略', '$S_t = \\hat{D}_t + z \\cdot \\sigma_{\\hat{D}}$'],
        ['智慧决策\n（IDMR）', 'DQN', '状态维度', '5'],
        ['', '', '动作维度', '30（离散化）'],
        ['', '', '动作范围', '$[11, 40]$'],
        ['', '', '隐藏层维度', '64'],
        ['', '', '学习率', '$10^{-4}$'],
        ['', '', '折扣因子 $\\gamma$', '0.9'],
        ['', '', '批大小', '32'],
        ['', '', '经验回放池', '20000'],
        ['', '', '探索率 $\\varepsilon$', '衰减策略'],
        ['', '', '训练步数', '40000'],
    ]
    create_table(doc, table4_data,
                col_widths=[Cm(2.5), Cm(2.0), Cm(5.0), Cm(5.0)],
                first_col_left=True)

    # ============================================================
    # 表5: 情绪与协同机制参数
    # ============================================================
    add_table_caption(doc, '表5  情绪与协同机制参数')
    table5_data = [
        ['机制类别', '参数', '符号', '取值', '说明'],
        ['情绪演化', '情绪状态范围', '$E_t$', '$[-1, 1]$', '-1=恐慌, +1=乐观'],
        ['', '缺货惩罚权重', '$w_s$', '1.0', '情绪反馈信号'],
        ['', '精准匹配奖励', '$w_m$', '0.5', '钟形奖励触发'],
        ['', '库存积压惩罚', '$w_e$', '0.3', '过量库存惩罚'],
        ['', '有效惩罚放大', '$w_{eff}$', '$w_{s0}(1+|E_t|)$', '损失厌恶放大'],
        ['情绪传染', '传染概率', '$\\eta$', '$B(0.3)$', '向上游传染'],
        ['正向激励', '覆盖率区间', '$c_r$', '$[0.8, 1.5]$', '钟形奖励触发区间'],
        ['', '奖励函数', '$bonus$', '$f(c_r)$', '钟形函数'],
        ['成本参数', '单位库存成本', '$h$', '1.0', '每周期每单位'],
        ['', '单位缺货成本', '$b$', '2.0', '每周期每单位'],
        ['', '初始净库存', '$NS_0$', '10.0', '各节点统一'],
    ]
    create_table(doc, table5_data,
                col_widths=[Cm(2.0), Cm(2.5), Cm(2.0), Cm(3.5), Cm(4.5)],
                first_col_left=True)

    # ============================================================
    # 表6: 实验运行配置
    # ============================================================
    add_table_caption(doc, '表6  实验运行配置')
    table6_data = [
        ['参数', '取值', '说明'],
        ['仿真周期 $T$', '5000', '单组实验步数'],
        ['训练步数', '40000', 'DQN预训练'],
        ['随机种子 seed', '42', '100%可复现'],
        ['动态事件总数', '76次', '需求突变53+供应中断23'],
        ['归因分析记录', '8000条', '逐周期数据'],
        ['三组对比实验', 'Baseline / Exp_1 / Exp_2', '控制变量法'],
    ]
    create_table(doc, table6_data,
                col_widths=[Cm(3.5), Cm(4.5), Cm(6.5)],
                first_col_left=True)

    # ============================================================
    # 表7: 持续学习参数
    # ============================================================
    add_table_caption(doc, '表7  持续学习参数（EWC + PER）')
    table7_data = [
        ['机制', '参数', '符号', '取值', '说明'],
        ['EWC', '正则化权重', '$\\lambda$', '自适应', 'Fisher矩阵约束'],
        ['', 'Fisher矩阵', '$F_i$', '在线估计', '参数重要性'],
        ['', '参考参数', '$\\theta_i^*$', 'Task1最优', '保护已学知识'],
        ['PER', '优先级指数', '$\\alpha$', '0.6', 'TD误差优先'],
        ['', '重要性采样修正', '$\\beta$', '0.4→1.0', '偏差修正'],
        ['', '经验回放池', '—', '20000', '优先级队列'],
        ['情绪感知噪声', '噪声标准差', '$\\sigma_{noise}$', '0.15', '$E_p = clip(E_t + N(0,\\sigma), -1, 1)$'],
    ]
    create_table(doc, table7_data,
                col_widths=[Cm(2.0), Cm(2.5), Cm(2.0), Cm(2.5), Cm(5.5)],
                first_col_left=True)

    # ============================================================
    # 基础设置总结
    # ============================================================
    doc.add_paragraph()
    heading3 = doc.add_heading(level=2)
    run = heading3.add_run('基础设置总结')
    run.font.name = FONT_HEADING_CN
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_HEADING_CN)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('本研究在四级串联供应链仿真环境下，采用控制变量法设计三组对比实验，通过AR(1)需求模型、情绪演化方程、正向激励函数、DQN决策算法、多智能体协同通道、EWC+PER持续学习机制的系统集成，构建了完整的"情绪扰动→激励机制→协同鲁棒"实验框架。所有参数严格固定（seed=42），确保结果可量化、可复现，为后续假设检验与因果分析奠定坚实基础。')
    set_run_font(run, size=11)

    # ============================================================
    # 保存
    # ============================================================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        '实验基础设置.docx'
    )

    try:
        if os.path.exists(output_path):
            os.remove(output_path)
    except PermissionError:
        output_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            '实验基础设置_v2.docx'
        )
        print(f"[提示] 原文件被占用，保存为: {output_path}")

    doc.save(output_path)
    print(f"[OK] 实验基础设置文档已生成: {output_path}")
    print(f"     大小: {os.path.getsize(output_path) / 1024:.1f} KB")

    # 统计 OMML 公式
    import zipfile
    with zipfile.ZipFile(output_path, 'r') as z:
        with z.open('word/document.xml') as f:
            content = f.read().decode('utf-8')
            omath_count = content.count('<m:oMath>')
            omathpara_count = content.count('<m:oMathPara>')
            print(f"     OMML 公式: {omath_count} 个 oMath + {omathpara_count} 个 oMathPara")


if __name__ == '__main__':
    create_basic_setup_docx()
