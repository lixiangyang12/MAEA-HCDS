# -*- coding: utf-8 -*-
"""燕山大学读博课程补充规划 -> DOCX"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 封面 =====
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('燕山大学读博 · 课程补充规划')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x2a)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('工商管理博士 + 人工智能与硬件方向创业 = 综合能力矩阵')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x09, 0x84, 0xe3)

doc.add_paragraph()

# 现状描述
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    '你的现状：工商管理博士在读，有AI产品经验（智预算PRO），\n'
    '创业方向为"AI+硬件"（AI教育硬件等）。\n'
    '燕山大学优势：机械工程全国A级、AI学院2025年新成立、\n'
    '控制科学B级、计算机科学B级、有智谱AI联合实验室。\n\n'
    '核心策略：以工商管理博士为"主学位"，利用燕山大学工科优势，\n'
    '跨学科选修/旁听，构建"管理+AI+硬件"三栖能力。'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ===== 辅助函数 =====
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # 深色背景
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1a2a3a', qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
    # Rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return table

def add_highlight_box(text, color_hex='0984e3'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*[int(color_hex[i:i+2],16) for i in range(0,6,2)])
    return p

# ================================================================
# 一、博士必修课程
# ================================================================
add_heading_styled('一、博士必修课程（学位课）', level=1)

add_para('以下为燕山大学工商管理博士必修课程，必须修读。每门课程标注了"对创业的价值"，帮助你带着目标学习。', size=10, color=RGBColor(0x55,0x55,0x55))

add_table(
    ['课程', '学分', '对AI硬件创业的价值'],
    [
        ['管理研究方法', '2', '学术研究基本功，但案例研究法、行动研究法可直接用于产品验证和商业模式迭代'],
        ['博弈论/高级微观经济学', '3', '市场竞争、定价策略、供应链谈判——AI硬件定价和渠道博弈的核心理论工具'],
        ['多元统计分析', '3', '硬件产品用户数据分析、A/B测试、市场调研数据处理——产品经理必备技能'],
        ['中国马克思主义与当代', '2', '理解政策导向——AI和硬件的政策红利（新质生产力、智能制造、人工智能+）'],
        ['博士生英语', '3', '阅读国际AI/硬件前沿论文、国际商务谈判、海外市场拓展'],
    ]
)

add_para('')
add_para('⚠ 核心提醒：这些是"必须修"的课程，但你可以把它们"用活"——管理研究方法用于产品验证，博弈论用于定价策略，统计用于用户数据分析。不要把它们当作"应付毕业"的课，而是当作"创业工具箱"的基础层。', bold=True, size=10, color=RGBColor(0xe1,0x70,0x55))

doc.add_paragraph()

# 跨学科补修
add_heading_styled('跨学科补修课程（不计学分但必须）', level=2)
add_para('燕山大学规定：跨学科报考的博士生须补修本学科硕士骨干课程2门，不计学分。如果你的硕士不是工商管理背景，建议选择以下与创业最相关的：', size=10, color=RGBColor(0x55,0x55,0x55))

add_table(
    ['推荐补修课程', '理由'],
    [
        ['运筹学', '供应链优化、生产排程、库存管理——AI硬件供应链的生命线'],
        ['战略管理', '创业公司战略制定，理解竞争格局——AI硬件赛道竞争激烈'],
        ['数智创新与创业管理', '直接对口的创业课程，燕山大学的特色课程（数智化方向）'],
        ['管理信息系统', '理解企业数字化架构，AI硬件产品需要嵌入企业IT生态'],
    ]
)

add_para('')
add_para('💡 积极心态：补修不是"惩罚"，而是"补短板"——你的竞争对手（纯商科博士）没有这些课程的"压力"，但他们也没有"动力"去学这些。你带着创业目标去学，效果完全不同。', bold=True, size=10, color=RGBColor(0x2e,0x7d,0x32))

doc.add_page_break()

# ================================================================
# 二、专业选修课程
# ================================================================
add_heading_styled('二、工商管理博士选修课程（至少选2门）', level=1)

add_table(
    ['课程', '创业价值', '推荐度'],
    [
        ['博弈论与信息经济学', '信息不对称下的激励机制设计——AI Agent之间的协作激励、供应链信息共享', '★★★★★'],
        ['产业组织理论与应用', '理解产业竞争结构，判断AI硬件赛道的市场集中度、进入壁垒', '★★★★★'],
        ['技术经济评价理论与方法', 'AI硬件项目的技术可行性+经济可行性评估——投资人最看重的技能', '★★★★★'],
        ['物流与供应链管理', '硬件创业的生命线：元器件采购、生产排期、库存管理、全球物流', '★★★★'],
        ['中外管理思想比较', '理解不同文化下的管理哲学，出海硬件产品需要本土化管理思维', '★★★'],
        ['营销管理理论与研究前沿', 'AI硬件的品牌定位、渠道策略、用户增长', '★★★★'],
        ['现代财务理论', '融资估值、现金流管理、创业财务——硬件创业烧钱快，财务管控是生死线', '★★★★'],
    ]
)

add_para('')
add_heading_styled('优先推荐：3门必选 + 2门建议', level=2)
add_para('3门必选（与创业最直接相关）：', bold=True, size=11, color=RGBColor(0x09,0x84,0xe3))
add_para('① 博弈论与信息经济学：你的智预算PRO本质上就是多Agent博弈系统，这门课是理论基础')
add_para('② 产业组织理论与应用：AI硬件赛道竞争格局分析——你做教育硬件，竞争对手是谁？市场结构是什么？')
add_para('③ 技术经济评价理论与方法：写BP（商业计划书）的核心技能——投资人问的第一个问题就是"这个项目经济上可行吗？"')
add_para('')
add_para('2门建议（根据时间灵活安排）：', bold=True, size=11, color=RGBColor(0x09,0x84,0xe3))
add_para('④ 物流与供应链管理：硬件创业90%的失败在供应链上')
add_para('⑤ 营销管理理论与研究前沿：技术出身的人最容易忽视营销')

add_para('')
add_para('选修课策略 = 博弈论(理论基础) + 产业组织(竞争分析) + 技术经济(投资评估) + 供应链(执行保障) + 营销(市场落地)', bold=True, size=10, color=RGBColor(0x09,0x84,0xe3))

doc.add_page_break()

# ================================================================
# 三、跨学科旁听
# ================================================================
add_heading_styled('三、跨学科旁听：燕山大学工科资源', level=1)
add_para('这是你选择燕山大学的最大"隐藏福利"——燕山大学以工科见长，机械工程全国A级，AI学院2025年新成立，有智谱AI联合实验室。', bold=True, size=10, color=RGBColor(0x6c,0x5c,0xe7))

add_heading_styled('四大跨学科旁听领域', level=2)
add_table(
    ['领域', '核心课程', '与AI硬件创业的关联'],
    [
        ['AI学院：人工智能', '机器学习、深度学习、大模型技术、具身智能、计算机视觉、NLP', 'AI教育硬件的"大脑"——语音交互、视觉识别、个性化推荐、大模型驱动'],
        ['机械学院：硬件工程', '智能制造工程、机器人技术、传感器、工业设计', 'AI教育硬件的"骨骼"——工业设计、结构工程、传感器选型、生产工艺'],
        ['电气学院：控制与嵌入式', '嵌入式系统、物联网技术、模式识别、智能控制', 'AI教育硬件的"神经系统"——硬件控制、嵌入式AI、端侧推理'],
        ['材料学院：新材料', '新材料科学与工程、增材制造（3D打印）', 'AI教育硬件的"皮肤"——外壳材料、触感设计、生产成本优化'],
    ]
)

add_para('')
add_heading_styled('精选旁听课程清单（按优先级）', level=2)
add_table(
    ['优先级', '课程', '所属学院', '为什么重要'],
    [
        ['★★★★★', '机器学习与深度学习', 'AI学院', '理解AI硬件的核心算法，能和技术团队深度对话'],
        ['★★★★★', '嵌入式系统', '电气学院', 'AI硬件=算法+嵌入式，不懂嵌入式做不了硬件产品经理'],
        ['★★★★★', '工业设计与产品创新', '机械学院', '硬件产品的用户体验、外观设计——教育硬件是给孩子用的'],
        ['★★★★', '大模型技术/具身智能', 'AI学院', '前沿技术趋势，未来AI硬件的方向'],
        ['★★★★', '物联网技术', '电气学院', 'AI硬件需要联网、需要数据互通'],
        ['★★★★', '传感器与检测技术', '机械学院', '教育硬件需要传感器采集数据'],
        ['★★★', '增材制造（3D打印）', '材料学院', '硬件原型快速迭代——创业初期不用开模'],
        ['★★★', '智能制造工程', '机械学院', '理解制造业全流程，将来量产时能管理工厂'],
    ]
)

add_para('')
add_para('💡 旁听策略：不要贪多。4年读博期间，每学期旁听1-2门工科课程即可。重点不是"学会写代码"，而是"能和技术团队深度对话"——你不需要成为工程师，但你需要"听得懂工程师在说什么，能判断技术方案的可行性"。', bold=True, size=10, color=RGBColor(0x2e,0x7d,0x32))

doc.add_page_break()

# ================================================================
# 四、自学补充
# ================================================================
add_heading_styled('四、课堂之外的"自学知识地图"', level=1)
add_para('以下领域课程体系难以覆盖，但对你创业至关重要，需要自学：', size=10, color=RGBColor(0x55,0x55,0x55))

add_table(
    ['领域', '核心内容', '为什么必须自学'],
    [
        ['AI产品管理', 'AI产品经理方法论、AI硬件产品定义、PRD撰写、用户研究', '商学院不教产品管理，工学院不教产品管理——这是"翻译层"能力'],
        ['硬件创业方法论', '硬件MVP、DFM（可制造性设计）、BOM成本管理、供应链管理实操', '硬件创业和软件创业完全不同——失败成本高、迭代慢'],
        ['创业融资与法律', 'BP撰写、估值模型、股权结构、专利保护、数据合规（含AI法规）', 'AI+硬件的知识产权保护、儿童数据隐私'],
        ['教育行业知识', '教育心理学、儿童发展、K12政策、家校共育模式', '你做的是教育硬件，必须理解教育场景'],
        ['AI硬件技术扫盲', '端侧AI（TinyML）、语音交互、CV、边缘计算、芯片选型', '不需要成为工程师，但需要能判断技术方案的可行性'],
    ]
)

add_para('')
add_heading_styled('自学资源推荐', level=2)
add_para('AI产品管理：', bold=True)
add_para('• 书籍：《Inspired》（Marty Cagan）、《AI产品经理手册》')
add_para('• 实践：把智预算PRO作为一个产品来打磨——写PRD、做用户研究、迭代')
add_para('')
add_para('硬件创业：', bold=True)
add_para('• 书籍：《硬件创业》（The Hardware Startup）、《从0到1》')
add_para('• 实践：去深圳华强北走一趟，了解硬件供应链——这是任何课堂都教不了的')
add_para('')
add_para('创业融资：', bold=True)
add_para('• 书籍：《创业融资》、《风险投资交易》')
add_para('• 实践：参加创业大赛（互联网+、挑战杯），撰写BP、路演')
add_para('')
add_para('教育行业：', bold=True)
add_para('• 书籍：《教育心理学》、《人是如何学习的》')
add_para('• 实践：去学校、培训机构做用户调研，观察孩子如何使用电子产品学习')
add_para('')
add_para('AI硬件技术：', bold=True)
add_para('• 课程：Coursera "TinyML"专项课程、MIT "6.S094"')
add_para('• 实践：买一块树莓派或Arduino，自己做一个AI硬件原型')

doc.add_page_break()

# ================================================================
# 五、四年路线图
# ================================================================
add_heading_styled('五、四年课程补充路线图', level=1)

add_heading_styled('第一年：打好基础（课程为主）', level=2)
add_para('• 完成所有博士必修课（管理研究方法、博弈论、统计、英语）')
add_para('• 补修2门硕士课程（推荐：运筹学 + 数智创新与创业管理）')
add_para('• 选修：博弈论与信息经济学 + 产业组织理论与应用')
add_para('• 旁听：机器学习与深度学习（AI学院，了解AI底层原理）')
add_para('• 确定博士论文研究方向（建议与"人机共生系统"相关）')
add_para('• 参加创业大赛，写第一版BP，验证创业想法')

add_para('')
add_heading_styled('第二年：跨学科拓展（旁听+实践）', level=2)
add_para('• 选修：技术经济评价 + 物流与供应链管理')
add_para('• 旁听：嵌入式系统 + 工业设计与产品创新（理解硬件设计和制造）')
add_para('• 开始撰写博士论文开题报告')
add_para('• 利用燕山大学AI学院资源，与智谱AI联合实验室建立联系')
add_para('• 去深圳华强北调研硬件供应链，做第一个硬件原型')

add_para('')
add_heading_styled('第三年：论文+创业双向推进', level=2)
add_para('• 旁听：大模型技术/具身智能 + 物联网技术（前沿追踪）')
add_para('• 博士论文中期进展，发表期刊论文')
add_para('• 硬件原型迭代，种子用户测试')
add_para('• 参加创业大赛，争取种子轮融资')
add_para('• 利用博士身份，建立行业人脉（学术会议+创业社群）')

add_para('')
add_heading_styled('第四年：毕业+落地', level=2)
add_para('• 完成博士论文答辩')
add_para('• 创业公司正式注册，争取天使轮融资')
add_para('• 全职创业或博后（利用博后站作为创业缓冲）')
add_para('• 燕山大学博士后流动站是"安全网"——如果创业需要更多时间，可以申请博后')

doc.add_page_break()

# ================================================================
# 六、全景总结
# ================================================================
add_heading_styled('六、全景总结', level=1)

add_heading_styled('五大能力模块 × 课程来源', level=2)
add_table(
    ['能力模块', '对应课程', '来源', '对创业的作用'],
    [
        ['管理理论', '管理研究方法、博弈论、产业组织、战略管理', '博士必修+选修', '商业模式设计、竞争策略'],
        ['数据分析', '多元统计、技术经济评价、运筹学', '博士必修+选修+补修', '用户分析、投资评估'],
        ['技术理解', '机器学习、嵌入式、物联网、大模型', '跨学科旁听', '和技术团队深度对话'],
        ['硬件产品', '工业设计、传感器、增材制造、智能制造', '跨学科旁听', '产品定义、供应链管理'],
        ['创业实践', '创业大赛、融资、行业知识', '自学+实践', '公司运营、市场落地'],
    ]
)

add_para('')
add_heading_styled('燕山大学可利用资源', level=2)
add_table(
    ['资源', '如何利用'],
    [
        ['AI学院（2025年新成立）', '名誉院长段广仁院士，顾问唐杰教授（清华智谱AI创始人）——主动接触，争取合作机会'],
        ['智谱AI联合实验室', '大模型技术前沿，你的智预算PRO可以对接智谱AI的底层能力'],
        ['机械学院（全国A级）', '硬件设计、制造工艺、工业设计——AI教育硬件的"硬件"来源'],
        ['创业大赛', '互联网+、挑战杯、中国创新创业大赛——获奖可以为融资加分'],
        ['博士后流动站', '12个博后站——毕业后的创业缓冲或技术合作'],
        ['河北省产业政策', '京津冀协同发展、河北制造业升级——AI教育硬件可能获得政策支持'],
    ]
)

add_para('')
add_heading_styled('三个核心原则', level=2)

add_para('原则一：博士论文 = 创业的"理论背书"', bold=True, size=11, color=RGBColor(0x2e,0x7d,0x32))
add_para('你的博士论文（人机共生系统）不是"应付毕业"的作业，而是你创业的核心壁垒——当投资人问你"为什么是你"时，你的回答是："我是这个领域的博士，我研究了4年人机共生系统，我的论文就是我的产品理论基础。"', size=10)

add_para('')
add_para('原则二：旁听 = 不是"学技术"，而是"建翻译能力"', bold=True, size=11, color=RGBColor(0x2e,0x7d,0x32))
add_para('你不需要成为AI工程师或硬件工程师。你需要的是"翻译能力"——能用商业语言和技术团队对话，能用技术语言和投资人对话。工商管理博士+工科旁听 = 天然的"翻译官"。', size=10)

add_para('')
add_para('原则三：创业 = 学术的"自然延伸"', bold=True, size=11, color=RGBColor(0x2e,0x7d,0x32))
add_para('你的智预算PRO和五个研究设想，不是"课余爱好"，而是学术研究的应用场景。你的创业公司（AI教育硬件）完全可以成为你论文的"案例研究对象"——学术和创业不是二选一，而是互相促进。', size=10)

add_para('')
add_para('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('你的独特优势 = 工商管理博士(理论深度) + 工科旁听(技术理解) + AI产品经验(实践能力) + 创业驱动(行动力) = "懂技术、懂管理、懂产品、懂商业"的复合型创业者')
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x09,0x84,0xe3)

# 保存
output_path = r'c:\个人资料\申博材料\东北财经大学\燕山大学读博_课程补充规划.docx'
doc.save(output_path)
print(f'DOCX saved: {output_path}')