"""
引用标注处理脚本
================
功能：
1. 将正文中"作者(年份)"格式的引用转换为上标[n]角标
2. 修正错误的引用序号
3. 为缺失引用的作者添加引用标注
4. 将英文双引号"替换为中文双引号""
5. 跳过参考文献部分和包含OMML公式的段落
"""
import re
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

INPUT_PATH = r'c:\个人资料\申博材料\企业运营与科研管理数据库\改进牛鞭效应新途径：人智协同决策系统\改进牛鞭效应新途径：人智协同决策系统.docx'
OUTPUT_PATH = r'c:\个人资料\申博材料\企业运营与科研管理数据库\改进牛鞭效应新途径：人智协同决策系统\改进牛鞭效应新途径：人智协同决策系统_引用版.docx'

# 字体设置
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_HEADING_CN = '黑体'

# ============================================================
# 文本替换规则（按顺序执行）
# 替换文本中的 [n] 会被转换为上标角标
# ============================================================
TEXT_REPLACEMENTS = [
    # ===== 1. 修正错误的引用序号（优先处理）=====
    ('Sterman[1]', 'Sterman[11]'),       # 段落87: Sterman应为[11]非[1]
    ('Chen等[2]', 'Chen等[13]'),         # 段落87: Chen应为[13]非[2]
    ('理性决策[6]', '理性决策[33]'),      # 段落145: 理性决策应引李勇[33]非[6]

    # ===== 2. 作者(年份)格式 -> 作者[序号] =====
    ('Lee等(1997)', 'Lee等[1]'),
    ('Lee等（1997）', 'Lee等[1]'),
    ('Kahneman & Tversky, 1979', 'Kahneman & Tversky[2]'),
    ('Cachon, 2003', 'Cachon[8]'),

    # ===== 3. 为缺失引用的作者添加[序号] =====
    ('Lee等的理论框架', 'Lee等[1]的理论框架'),
    ('Chen刻画的', 'Chen[13]刻画的'),
    ('Sterman设计的', 'Sterman[11]设计的'),
    ('李勇等智慧四级', '李勇等[33]智慧四级'),
    ('李勇等基于DQN', '李勇等[33]基于DQN'),
    ('李勇等提出的', '李勇等[33]提出的'),
    ('Chen 等', 'Chen等[13]'),
    # 行为经济学研究 -> 引用Kahneman前景理论[2]
    ('行为经济学研究表明', '行为经济学研究[2]表明'),
]


def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=None, bold=None, color=None):
    """设置run字体（中英文分别设置）"""
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def replace_quotation_marks(text):
    """替换英文双引号为中文双引号（成对替换）"""
    result = []
    count = 0
    for char in text:
        if char == '"':  # 英文双引号 U+0022
            if count % 2 == 0:
                result.append('\u201C')  # 中文左双引号 "
            else:
                result.append('\u201D')  # 中文右双引号 "
            count += 1
        else:
            result.append(char)
    return ''.join(result)


def has_omml(paragraph):
    """检查段落是否包含OMML公式"""
    p_element = paragraph._element
    ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    omaths = p_element.findall('.//m:oMath', ns)
    omathparas = p_element.findall('.//m:oMathPara', ns)
    return len(omaths) > 0 or len(omathparas) > 0


def get_paragraph_font_info(paragraph):
    """获取段落的字体信息（从第一个有文本的run继承）"""
    for run in paragraph.runs:
        if run.text and run.text.strip():
            return {
                'size': run.font.size,
                'bold': run.font.bold,
                'color': run.font.color.rgb if (run.font.color and run.font.color.rgb) else None,
            }
    return {'size': Pt(11), 'bold': None, 'color': None}


def needs_replacement(text):
    """检查文本是否需要替换"""
    for old, new in TEXT_REPLACEMENTS:
        if old in text:
            return True
    if '"' in text:
        return True
    return False


def rebuild_paragraph(paragraph, font_info):
    """重建段落：保留段落格式，重建run，将[n]转换为上标"""
    full_text = paragraph.text

    # 应用文本替换
    for old, new in TEXT_REPLACEMENTS:
        full_text = full_text.replace(old, new)

    # 替换英文双引号为中文双引号
    full_text = replace_quotation_marks(full_text)

    # 清空段落所有run（保留段落属性pPr）
    p_element = paragraph._element
    for r in p_element.findall(qn('w:r')):
        p_element.remove(r)

    # 按 [n] 模式分割文本
    parts = re.split(r'(\[\d+\])', full_text)

    for part in parts:
        if not part:
            continue
        if re.match(r'^\[\d+\]$', part):
            # 上标角标
            run = paragraph.add_run(part)
            run.font.superscript = True
            set_run_font(run, size=font_info['size'], bold=font_info['bold'])
        else:
            # 正常文本
            run = paragraph.add_run(part)
            set_run_font(run, size=font_info['size'], bold=font_info['bold'],
                        color=font_info['color'])


def replace_in_runs(paragraph):
    """对含OMML公式的段落，仅在run级别做文本替换（不重建段落）
    用于处理包含公式无法重建的段落
    """
    # 收集所有文本run（跳过包含m:oMath的run）
    text_runs = []
    for run in paragraph.runs:
        # 检查run是否包含公式
        r_element = run._element
        ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        has_math = len(r_element.findall('.//m:oMath', ns)) > 0
        if not has_math and run.text:
            text_runs.append(run)

    if not text_runs:
        return False

    # 合并所有文本run的文本
    combined = ''.join(run.text for run in text_runs)
    original = combined

    # 应用文本替换
    for old, new in TEXT_REPLACEMENTS:
        combined = combined.replace(old, new)

    # 替换双引号
    combined = replace_quotation_marks(combined)

    if combined == original:
        return False

    # 检查替换后是否包含[n]（需要上标）
    if not re.search(r'\[\d+\]', combined):
        # 不包含[n]，直接写回第一个run，清空其余
        text_runs[0].text = combined
        for run in text_runs[1:]:
            run.text = ''
        return True

    # 包含[n]，需要将[n]部分转为上标
    # 简化处理：将[n]保留在文本中，但设置上标
    # 这里采用分割方式：将文本分割，[n]部分创建新的上标run
    parts = re.split(r'(\[\d+\])', combined)

    # 清空所有文本run
    for run in text_runs:
        run.text = ''

    # 获取字体信息
    font_info = get_paragraph_font_info(paragraph)

    # 在第一个文本run的位置插入新内容
    # 找到第一个文本run的父元素和位置
    first_run_elem = text_runs[0]._element
    parent = first_run_elem.getparent()
    insert_idx = list(parent).index(first_run_elem)

    # 删除所有文本run
    for run in text_runs:
        parent.remove(run._element)

    # 插入新的run
    for part in parts:
        if not part:
            continue
        new_run = OxmlElement('w:r')
        # 创建rPr
        rPr = OxmlElement('w:rPr')

        # 设置字体
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_CN)
        rFonts.set(qn('w:ascii'), FONT_EN)
        rFonts.set(qn('w:hAnsi'), FONT_EN)
        rPr.append(rFonts)

        # 设置字号
        if font_info['size']:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_info['size'].pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_info['size'].pt * 2)))
            rPr.append(szCs)

        # 如果是[n]，设置上标
        if re.match(r'^\[\d+\]$', part):
            vertAlign = OxmlElement('w:vertAlign')
            vertAlign.set(qn('w:val'), 'superscript')
            rPr.append(vertAlign)

        new_run.append(rPr)

        # 添加文本
        t = OxmlElement('w:t')
        t.text = part
        t.set(qn('xml:space'), 'preserve')
        new_run.append(t)

        parent.insert(insert_idx, new_run)
        insert_idx += 1

    return True


def process_document():
    """主处理函数"""
    doc = Document(INPUT_PATH)

    # 找到参考文献开始位置
    ref_start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '参考文献':
            ref_start_idx = i
            break

    if ref_start_idx is None:
        print("[警告] 未找到'参考文献'章节，将处理所有段落")
        ref_start_idx = len(doc.paragraphs)

    print(f"[信息] 参考文献起始段落: {ref_start_idx}")
    print(f"[信息] 总段落数: {len(doc.paragraphs)}")

    # 处理统计
    stats = {
        'total': 0,
        'replaced': 0,
        'skipped_ref': 0,
        'skipped_omml': 0,
        'skipped_nochange': 0,
    }

    for i, p in enumerate(doc.paragraphs):
        stats['total'] += 1

        # 跳过参考文献部分
        if i >= ref_start_idx:
            stats['skipped_ref'] += 1
            continue

        # 跳过空段落
        if not p.text.strip():
            stats['skipped_nochange'] += 1
            continue

        # 检查是否需要替换
        if not needs_replacement(p.text):
            stats['skipped_nochange'] += 1
            continue

        # 检查是否包含OMML公式
        if has_omml(p):
            # 对含公式段落，使用run级别替换（保留公式）
            print(f"  [处理-含公式] 段落[{i}]: {p.text[:60]}...")
            if replace_in_runs(p):
                stats['replaced'] += 1
            else:
                stats['skipped_nochange'] += 1
            continue

        # 获取字体信息并重建段落
        font_info = get_paragraph_font_info(p)
        print(f"  [处理] 段落[{i}]: {p.text[:60]}...")
        rebuild_paragraph(p, font_info)
        stats['replaced'] += 1

    # 保存
    doc.save(OUTPUT_PATH)
    print(f"\n[OK] 文档已保存: {OUTPUT_PATH}")
    print(f"     大小: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
    print(f"\n[统计]")
    print(f"  总段落: {stats['total']}")
    print(f"  已替换: {stats['replaced']}")
    print(f"  跳过(参考文献): {stats['skipped_ref']}")
    print(f"  跳过(含公式): {stats['skipped_omml']}")
    print(f"  跳过(无需替换): {stats['skipped_nochange']}")


if __name__ == '__main__':
    process_document()
