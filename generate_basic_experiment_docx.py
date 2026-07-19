"""
生成《中国管理科学》规范的基础实验 docx 文档
=============================================
将 基础实验_牛鞭效应对比分析.md 转换为 Word 文档

学术规范:
  - 三线表（顶线1.5pt / 表头线0.75pt / 底线1.5pt）
  - OMML 原生公式（LaTeX → MathML → OMML）
  - 嵌入 6 张图表（PNG 格式）
  - 中文字体：宋体（正文）/ 黑体（标题），英文：Times New Roman
  - 字号：章标题三号黑体，节标题四号黑体，正文五号宋体，
          表标题五号黑体加粗，表内文字五号宋体，表注小五号宋体
  - A4 纵向，上下边距 2.54cm，左右边距 3.17cm
"""

import os
import json
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
matplotlib.rcParams['axes.unicode_minus'] = False
matplotlib.rcParams['font.size'] = 11

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# OMML 公式转换 (LaTeX → MathML → OMML)
# ============================================================
try:
    from latex2mathml.converter import convert as _latex_to_mathml
    from lxml import etree
    _XSLT_PATHS = [
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
    ]
    _xslt_transform = None
    for _p in _XSLT_PATHS:
        if os.path.exists(_p):
            _xslt_tree = etree.parse(_p)
            _xslt_transform = etree.XSLT(_xslt_tree)
            break
    _OMML_AVAILABLE = _xslt_transform is not None
    if not _OMML_AVAILABLE:
        print("[警告] 未找到 MML2OMML.XSL，公式将回退为文本格式")
except Exception as e:
    _OMML_AVAILABLE = False
    print(f"[警告] OMML 公式转换不可用: {e}")


def latex_to_omml(latex_str):
    """将 LaTeX 公式转换为 OMML XML 元素"""
    mathml_str = _latex_to_mathml(latex_str)
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = _xslt_transform(mathml_tree)
    return omml_tree.getroot()


# ============================================================
# 字体与样式配置
# ============================================================
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_HEADING = '黑体'

NODE_NAMES = ['零售商', '批发商', '分销商', '制造商']

# 图表配色
COLOR_BASELINE = '#E74C3C'
COLOR_EXPIDMR = '#3498DB'
COLOR_TS = ['#E74C3C', '#E67E22', '#F39C12', '#C0392B']

# 路径
DATA_JSON = os.path.join('p0_results', '基础实验完整数据_20k.json')
FIG_DIR = 'svg_figures_basic'
PNG_DIR = os.path.join(FIG_DIR, 'png_temp')
os.makedirs(PNG_DIR, exist_ok=True)

# 实验参数
TOTAL_PERIODS = 20000
D = 10; RHO = 0.5; SIGMA_EPS = 5.0; L = 2; P = 5; Z = 2; C_L_RHO = 2.0
SEED = 42; INITIAL_INVENTORY = 10.0


# ============================================================
# 1. 图表生成 (PNG)
# ============================================================

def load_data():
    """加载实验数据"""
    with open(DATA_JSON, 'r', encoding='utf-8') as f:
        return json.load(f)


def generate_comparison_pngs(data):
    """生成3张对比柱状图 PNG (从JSON汇总数据)"""
    baseline = data['baseline']
    exp1 = data['exp1']

    # 图1: 方差比对比
    fig, ax = plt.subplots(figsize=(8, 5))
    x = np.arange(4); width = 0.35
    bwe_base = [baseline['bwe'][str(k)] for k in range(1, 5)]
    bwe_exp = [exp1['bwe'][str(k)] for k in range(1, 5)]
    bars1 = ax.bar(x - width/2, bwe_base, width, label='理性决策',
                   color=COLOR_BASELINE, edgecolor='black', linewidth=0.8)
    bars2 = ax.bar(x + width/2, bwe_exp, width, label='智慧决策',
                   color=COLOR_EXPIDMR, edgecolor='black', linewidth=0.8)
    for bar in bars1:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + max(bwe_base)*0.02,
                f'{h:.2f}', ha='center', va='bottom', fontsize=9)
    for bar in bars2:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + max(bwe_base)*0.02,
                f'{h:.2f}', ha='center', va='bottom', fontsize=9)
    ax.set_xlabel('供应链节点', fontsize=12)
    ax.set_ylabel('方差比 BWE', fontsize=12)
    ax.set_title('两种决策下的方差比对比', fontsize=13, fontweight='bold')
    ax.set_xticks(x); ax.set_xticklabels(NODE_NAMES, fontsize=11)
    ax.legend(fontsize=10, loc='upper left')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_ylim(0, max(max(bwe_base), max(bwe_exp)) * 1.2)
    plt.tight_layout()
    fig.savefig(os.path.join(PNG_DIR, 'fig_bwe_comparison.png'), dpi=200)
    plt.close(fig)

    # 图2: 平均成本对比
    fig, ax = plt.subplots(figsize=(8, 5))
    cost_base = [baseline['avg_cost'][str(k)] for k in range(1, 5)]
    cost_exp = [exp1['avg_cost'][str(k)] for k in range(1, 5)]
    bars1 = ax.bar(x - width/2, cost_base, width, label='理性决策',
                   color=COLOR_BASELINE, edgecolor='black', linewidth=0.8)
    bars2 = ax.bar(x + width/2, cost_exp, width, label='智慧决策',
                   color=COLOR_EXPIDMR, edgecolor='black', linewidth=0.8)
    for bar in bars1:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + max(cost_base)*0.02,
                f'{h:.1f}', ha='center', va='bottom', fontsize=9)
    for bar in bars2:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + max(cost_base)*0.02,
                f'{h:.1f}', ha='center', va='bottom', fontsize=9)
    ax.set_xlabel('供应链节点', fontsize=12)
    ax.set_ylabel('平均成本', fontsize=12)
    ax.set_title('两种决策下的平均成本对比', fontsize=13, fontweight='bold')
    ax.set_xticks(x); ax.set_xticklabels(NODE_NAMES, fontsize=11)
    ax.legend(fontsize=10, loc='upper left')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_ylim(0, max(max(cost_base), max(cost_exp)) * 1.2)
    plt.tight_layout()
    fig.savefig(os.path.join(PNG_DIR, 'fig_cost_comparison.png'), dpi=200)
    plt.close(fig)

    # 图3: 服务水平对比
    fig, ax = plt.subplots(figsize=(8, 5))
    sl_base = [baseline['sl'][str(k)] * 100 for k in range(1, 5)]
    sl_exp = [exp1['sl'][str(k)] * 100 for k in range(1, 5)]
    bars1 = ax.bar(x - width/2, sl_base, width, label='理性决策',
                   color=COLOR_BASELINE, edgecolor='black', linewidth=0.8)
    bars2 = ax.bar(x + width/2, sl_exp, width, label='智慧决策',
                   color=COLOR_EXPIDMR, edgecolor='black', linewidth=0.8)
    for bar in bars1:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + 0.5,
                f'{h:.2f}%', ha='center', va='bottom', fontsize=9)
    for bar in bars2:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + 0.5,
                f'{h:.2f}%', ha='center', va='bottom', fontsize=9)
    ax.axhline(y=97.7, color='gray', linestyle='--', linewidth=1, alpha=0.7)
    ax.text(3.4, 97.9, '理论目标97.7%', fontsize=9, color='gray')
    ax.set_xlabel('供应链节点', fontsize=12)
    ax.set_ylabel('服务水平 SL (%)', fontsize=12)
    ax.set_title('两种决策下的服务水平对比', fontsize=13, fontweight='bold')
    ax.set_xticks(x); ax.set_xticklabels(NODE_NAMES, fontsize=11)
    ax.legend(fontsize=10, loc='lower right')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_ylim(80, 105)
    plt.tight_layout()
    fig.savefig(os.path.join(PNG_DIR, 'fig_sl_comparison.png'), dpi=200)
    plt.close(fig)
    print("  [图] 3张对比柱状图 PNG 已生成")


def generate_timeseries_pngs():
    """重新运行baseline获取时序数据, 生成3张时序图 PNG"""
    from supply_chain_env import SupplyChainEnv, RationalAgent

    print("  [时序图] 重新运行Baseline获取时序数据...")
    env = SupplyChainEnv(
        d=D, rho=RHO, sigma_eps=SIGMA_EPS, L=L, p=P, z=Z,
        C_L_rho=C_L_RHO, initial_inventory=INITIAL_INVENTORY, K=4,
        total_periods=TOTAL_PERIODS, seed=SEED,
    )
    agent = RationalAgent(L=L, p=P, z=Z, C_L_rho=C_L_RHO, sigma_eps=SIGMA_EPS)
    for k in range(1, 5):
        agent.init_node(k)

    env.reset()
    costs = {k: [] for k in range(1, 5)}
    sls = {k: [] for k in range(1, 5)}
    order_history = {k: [] for k in range(1, 5)}
    demand_history = []

    for t in range(TOTAL_PERIODS):
        D_t = env._generate_demand()
        env.customer_demand_history.append(D_t)
        demand_history.append(D_t)
        downstream_demand = {1: D_t}
        for k in range(1, env.K + 1):
            node = env.nodes[k]
            demand_k = downstream_demand.get(k, 0)
            arrived = node.pipeline[0] if len(node.pipeline) > 0 else 0.0
            node.net_stock += arrived
            if len(node.pipeline) > 0:
                node.pipeline.popleft()
            q_t = agent.decide(k, node.net_stock, sum(node.pipeline), demand_k)
            q_t = max(0, q_t)
            node.order_placed = q_t
            node.order_history.append(q_t)
            order_history[k].append(q_t)
            downstream_demand[k + 1] = q_t
            fulfilled = min(max(node.net_stock, 0), demand_k)
            node.net_stock -= fulfilled
            stockout = max(0, demand_k - fulfilled)
            costs[k].append(max(0, node.net_stock) * 1.0 + stockout * 2.0)
            sls[k].append(fulfilled / demand_k if demand_k > 0 else 1.0)
            node.demand_history.append(demand_k)
            node.pipeline.append(q_t)

    # 采样时序数据 (每10步)
    bwe_ts = {k: [] for k in range(1, 5)}
    cost_ts = {k: [] for k in range(1, 5)}
    sl_ts = {k: [] for k in range(1, 5)}
    x_ts = []
    for t in range(0, TOTAL_PERIODS, 10):
        x_ts.append(t)
        for k in range(1, 5):
            cost_ts[k].append(float(np.mean(costs[k][:t+1])))
            sl_ts[k].append(float(np.mean(sls[k][:t+1])))
            if t >= 200:
                wo = order_history[k][t-200:t]
                wd = demand_history[t-200:t]
                vq = float(np.var(wo)); vd = float(np.var(wd))
                bwe_ts[k].append(vq / vd if vd > 0 else 0)
            else:
                bwe_ts[k].append(0)

    # 图4: 方差比时序图
    fig, ax = plt.subplots(figsize=(10, 5))
    for k in range(1, 5):
        ax.plot(x_ts, bwe_ts[k], label=NODE_NAMES[k-1], color=COLOR_TS[k-1],
                linewidth=1.2, alpha=0.85)
    ax.set_xlabel('订货周期', fontsize=12)
    ax.set_ylabel('方差比 BWE', fontsize=12)
    ax.set_title('理性决策下的方差比', fontsize=13, fontweight='bold')
    ax.legend(fontsize=10, loc='upper right')
    ax.grid(alpha=0.3, linestyle='--')
    ax.set_xlim(0, TOTAL_PERIODS)
    plt.tight_layout()
    fig.savefig(os.path.join(PNG_DIR, 'fig_bwe_timeseries.png'), dpi=200)
    plt.close(fig)

    # 图5: 平均成本时序图
    fig, ax = plt.subplots(figsize=(10, 5))
    for k in range(1, 5):
        ax.plot(x_ts, cost_ts[k], label=NODE_NAMES[k-1], color=COLOR_TS[k-1],
                linewidth=1.2, alpha=0.85)
    ax.set_xlabel('订货周期', fontsize=12)
    ax.set_ylabel('累计平均成本', fontsize=12)
    ax.set_title('理性决策下的平均成本', fontsize=13, fontweight='bold')
    ax.legend(fontsize=10, loc='upper left')
    ax.grid(alpha=0.3, linestyle='--')
    ax.set_xlim(0, TOTAL_PERIODS)
    plt.tight_layout()
    fig.savefig(os.path.join(PNG_DIR, 'fig_cost_timeseries.png'), dpi=200)
    plt.close(fig)

    # 图6: 服务水平时序图
    fig, ax = plt.subplots(figsize=(10, 5))
    for k in range(1, 5):
        ax.plot(x_ts, [v * 100 for v in sl_ts[k]], label=NODE_NAMES[k-1],
                color=COLOR_TS[k-1], linewidth=1.2, alpha=0.85)
    ax.axhline(y=97.7, color='gray', linestyle='--', linewidth=1, alpha=0.7)
    ax.text(TOTAL_PERIODS * 0.7, 97.9, '理论目标97.7%', fontsize=9, color='gray')
    ax.set_xlabel('订货周期', fontsize=12)
    ax.set_ylabel('服务水平 SL (%)', fontsize=12)
    ax.set_title('理性决策下的服务水平', fontsize=13, fontweight='bold')
    ax.legend(fontsize=10, loc='lower right')
    ax.grid(alpha=0.3, linestyle='--')
    ax.set_xlim(0, TOTAL_PERIODS)
    ax.set_ylim(70, 102)
    plt.tight_layout()
    fig.savefig(os.path.join(PNG_DIR, 'fig_sl_timeseries.png'), dpi=200)
    plt.close(fig)
    print("  [图] 3张时序图 PNG 已生成")


# ============================================================
# 2. Word 文档辅助函数
# ============================================================

def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=10.5,
                 bold=False, italic=False, color=None):
    """设置run字体（中英文分别设置）"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表用）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge_name, edge_props in [('top', top), ('bottom', bottom),
                                   ('left', left), ('right', right)]:
        edge = tcBorders.find(qn(f'w:{edge_name}'))
        if edge is None:
            edge = OxmlElement(f'w:{edge_name}')
            tcBorders.append(edge)
        for attr in ['val', 'sz', 'color', 'space']:
            edge.attrib.pop(qn(f'w:{attr}'), None)
        if edge_props is None:
            edge.set(qn('w:val'), 'nil')
        else:
            edge.set(qn('w:val'), edge_props.get('val', 'single'))
            edge.set(qn('w:sz'), str(edge_props.get('sz', '6')))
            edge.set(qn('w:color'), edge_props.get('color', '000000'))
            edge.set(qn('w:space'), '0')


def set_cell_text(cell, text, bold=False, font_size=10.5,
                  cn_font=FONT_CN, en_font=FONT_EN,
                  align='center', vertical='center'):
    """设置单元格文本（中英文混排自动切换字体）"""
    cell.text = ''
    para = cell.paragraphs[0]
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_format = para.paragraph_format
    para_format.line_spacing = 1.0
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(0)
    if vertical == 'center':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    elif vertical == 'top':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    run = para.add_run(text)
    set_run_font(run, font_en=en_font, font_cn=cn_font,
                 size=font_size, bold=bold)


def add_heading(doc, text, level=1):
    """添加标题"""
    sizes = {1: 16, 2: 14, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, font_en=FONT_EN, font_cn=FONT_HEADING,
                 size=size, bold=True)
    return p


def add_paragraph(doc, text, size=10.5, bold=False, indent=True, align='left'):
    """添加正文段落"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_formula(doc, latex_str, number=''):
    """添加公式段落（OMML格式，带编号）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    if _OMML_AVAILABLE:
        try:
            omml = latex_to_omml(latex_str)
            p._element.append(omml)
        except Exception as e:
            run = p.add_run(latex_str)
            set_run_font(run, font_en=FONT_EN, size=11, bold=True)
    else:
        run = p.add_run(latex_str)
        set_run_font(run, font_en=FONT_EN, size=11, bold=True)

    if number:
        run = p.add_run(f'    ({number})')
        set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=10.5)
    return p


def add_inline_formula(paragraph, latex_str):
    """在段落中插入行内OMML公式"""
    if _OMML_AVAILABLE:
        try:
            omml = latex_to_omml(latex_str)
            paragraph._element.append(omml)
            return
        except Exception:
            pass
    run = paragraph.add_run(latex_str)
    set_run_font(run, font_en=FONT_EN, size=10.5, bold=True)


def add_table_title(doc, text):
    """添加表标题（表上方，五号黑体加粗，居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, font_en=FONT_EN, font_cn=FONT_HEADING,
                 size=10.5, bold=True)
    return p


def add_table_note(doc, text):
    """添加表注（表下方，小五号宋体，左对齐）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=9)
    return p


def add_figure(doc, png_path, caption, width_inches=5.5):
    """插入图片 + 图题（图下方居中，五号宋体）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(png_path, width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=10.5, bold=True)
    return cap


def make_three_line_table(doc, headers, rows, col_widths=None):
    """创建三线表"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 表头行
    line_15 = {'sz': '12', 'val': 'single', 'color': '000000'}  # 1.5pt
    line_075 = {'sz': '6', 'val': 'single', 'color': '000000'}   # 0.75pt
    no_line = None

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_text(cell, h, bold=True, font_size=10.5,
                      cn_font=FONT_HEADING, align='center')
        set_cell_border(cell, top=line_15, bottom=line_075,
                        left=no_line, right=no_line)

    # 数据行
    for i, row in enumerate(rows):
        is_last = (i == len(rows) - 1)
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            set_cell_text(cell, str(val), bold=False, font_size=10.5,
                          align='center')
            bottom = line_15 if is_last else no_line
            set_cell_border(cell, top=no_line, bottom=bottom,
                            left=no_line, right=no_line)

    return table


# ============================================================
# 3. 文档生成主函数
# ============================================================

def generate_docx():
    """生成完整的docx文档"""
    data = load_data()
    baseline = data['baseline']
    exp1 = data['exp1']

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ---- 文档标题 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run('经典供应链订货决策与牛鞭效应基础实验分析')
    set_run_font(run, font_en=FONT_EN, font_cn=FONT_HEADING,
                 size=18, bold=True)

    # ============================================================
    # 第1节 经典供应链订货决策
    # ============================================================
    add_heading(doc, '1 经典供应链订货决策', level=1)

    add_paragraph(doc,
        '为了让智慧决策机器人能够学到人类理论最优的决策，更好地适应决策环境，'
        '并最终做出最优决策，本文采用Chen刻画的牛鞭效应模型，供应链条按照'
        'Sterman设计的经典啤酒游戏模式运行。')

    add_heading(doc, '1.1 经典供应链模型', level=2)

    add_paragraph(doc,
        '经典多级供应链由零售商、批发商、分销商、制造商四级节点串联组成，'
        '终端顾客需求从零售商逐级向上游传递。定义每个供应链节点为节点 '
        'k（k=1,2,3,4），在每个周期 t 发生如下事件：')

    add_paragraph(doc,
        '第一步：节点 k 收到下游节点 k-1 的订单，观察自身库存水平，'
        '并向上游节点 k+1 按照理性决策原则补货；', indent=True)
    add_paragraph(doc,
        '第二步：节点 k 收到上游节点 k+1 发出的 L 个运输周期前的货物；', indent=True)
    add_paragraph(doc,
        '第三步：节点 k 满足下游节点 k-1 的订单，并更新自身库存和缺货。', indent=True)

    add_heading(doc, '1.2 理性决策', level=2)

    add_paragraph(doc,
        '供应链上每个节点所作决策均为理性决策，按照零售商的订货步骤依次向上游订货。'
        '零售商首先根据最近 p 个周期的需求对当期顾客需求进行简单移动平均'
        '（Simple Moving Average, SMA）预测，得到 L 提前期内的预测需求以及预测误差，'
        '然后根据订至点库存策略（Order-Up-To, OUT）计算期望库存，最后做出补货决策。'
        '该理性决策以最大化服务水平（Service Level, SL）为目标。')

    add_paragraph(doc, '零售商所观察到的周期 t 的需求遵循一阶自回归过程 AR(1)：')

    add_formula(doc, r'D_t = d + \rho D_{t-1} + \varepsilon_t', '1')

    add_paragraph(doc,
        '其中，d 是非负常量，ρ（|ρ|<1）是相关系数，ε_t 是误差项，'
        '服从独立同正态分布。期望库存 S_t 的表达式如下：')

    add_formula(doc, r'S_t = \hat{D}_t^L + z \hat{e}_t^L', '2')

    add_paragraph(doc, '其中，z 表示安全库存系数。SMA 预测过程如下：')

    add_formula(doc,
        r'\hat{D}_t^L = L \cdot \frac{1}{p}\sum_{i=1}^{p} D_{t-i}, \quad '
        r'\hat{e}_t^L = C_{L,\rho} \cdot \frac{\sigma_\varepsilon}{p} '
        r'\sqrt{\sum_{i=1}^{p} e_{t-i}^2}', '3')

    add_paragraph(doc,
        '其中，e_t 表示一个周期的预测误差，即 e_t = D_t - D̂_t¹，'
        'C_{L,ρ} 表示 L、ρ 和 p 的常数。订货策略如下：')

    add_formula(doc, r'q_t = S_t - (NS_t + WIP_t)', '4')

    add_paragraph(doc,
        '其中 NS_t 表示净库存，WIP_t 表示在途库存。牛鞭效应量化方式如下：')

    add_formula(doc, r'BWE = \frac{\mathrm{var}(q_t)}{\mathrm{var}(D_t)}', '5')

    add_paragraph(doc,
        '其中，var(q_t) 表示节点周期 t 时的订货量方差，var(D_t) 表示周期 t 时的'
        '顾客需求量方差。')

    add_heading(doc, '1.3 牛鞭效应', level=2)

    add_paragraph(doc,
        '基于Chen的理性决策和Sterman设计的啤酒游戏模式，本文复盘了动态的牛鞭效应现象。'
        '全部实验经过20000个周期的连续动态订货与发货，供应链上的方差比、平均成本以及'
        '服务水平三个核心指标呈现如下特征。')

    # ---- 1.3.1 方差比 ----
    add_heading(doc, '1）方差比', level=3)

    add_table_title(doc, '表1 理性决策下的牛鞭效应（保留小数点后2位）')

    headers1 = ['节点', '需求（均值）', '订单（均值）', '方差比']
    rows1 = []
    for k in range(1, 5):
        rows1.append([
            NODE_NAMES[k-1],
            f"{baseline['demand_mean'][str(k)]:.2f}",
            f"{baseline['order_mean'][str(k)]:.2f}",
            f"{baseline['bwe'][str(k)]:.2f}",
        ])
    make_three_line_table(doc, headers1, rows1,
                          col_widths=[2.5, 3.5, 3.5, 3.0])

    add_table_note(doc,
        '注：ρ=0.5, L=2, d=10, ε~N(0,5), z=2, C_{L,ρ}=2, p=5，'
        '订货运行周期为20000，各个节点初始库存均为10。')

    add_figure(doc, os.path.join(PNG_DIR, 'fig_bwe_timeseries.png'),
               '图1 理性决策下的方差比')

    add_paragraph(doc,
        '如表1所示，在理性决策下，经过20000个周期的连续动态订货与发货，'
        '供应链上各个节点都表现出订单均值与需求均值近似相等的特征，'
        '该现象与既有文献结论相一致。但在动态订货、发货的场景下，'
        '从零售商到制造商，各个节点订货量的方差与顾客需求量的方差之比'
        '在订货初期呈现出逐级放大并且激增的现象，后期逐渐平稳，'
        '但逐级放大的现象依旧存在。这说明在经典多级供应链模型下，'
        '随着订货周期的增加，供应链上的牛鞭效应有所缓和但仍然较为严重。'
        '制造商20000个周期的方差比达到301.75，约为零售商（4.05）的74倍，'
        '表明需求信息在逐级传递过程中被严重放大。')

    # ---- 1.3.2 平均成本 ----
    add_heading(doc, '2）平均成本', level=3)

    add_figure(doc, os.path.join(PNG_DIR, 'fig_cost_timeseries.png'),
               '图2 理性决策下的平均成本')

    add_paragraph(doc,
        '上图表示在理性决策下，随着订货周期的增加，供应链各个节点从订货开始'
        '到当前订货周期的累计平均成本。从零售商到制造商，各个节点的平均成本'
        '在订货初期呈现出激增的现象并且逐级增加，到订货后期，平均成本逐渐缓和，'
        '最后趋于平稳。其中，制造商20000个周期的平均成本（1212.92）约为'
        '零售商（32.22）的38倍，表明牛鞭效应导致上游节点承担了显著更高的'
        '库存成本与缺货成本。')

    # ---- 1.3.3 服务水平 ----
    add_heading(doc, '3）服务水平', level=3)

    add_paragraph(doc,
        '基于理性决策，本文采用有货率来衡量服务水平，即所有订货次数中有货的比例：')

    add_formula(doc, r'SL = 1 - \frac{\text{缺货次数}}{\text{订货次数}}', '6')

    add_paragraph(doc, '并且服务水平与安全库存系数存在如下关系：')

    add_formula(doc, r'SL = \Phi(z)', '7')

    add_paragraph(doc,
        '其中 Φ(·) 为标准正态分布的累积分布函数。当 z=2 时，'
        '理论目标服务水平为97.7%。')

    add_figure(doc, os.path.join(PNG_DIR, 'fig_sl_timeseries.png'),
               '图3 理性决策下的服务水平')

    add_paragraph(doc,
        '在理性决策下，供应链上各个节点以满足97.7%（z=2）的服务水平为目标。'
        '从实验结果来看，虽然在订货后期服务水平逐渐稳定在一个较高的值，'
        '但除制造商外，其他节点均未稳定达到97.7%的理论目标服务水平。'
        '零售商服务水平最低，约为98.92%，批发商和分销商分别为99.70%和99.76%。'
        '制造商服务水平达99.81%，高于理论目标值，这是由于制造商上游为无限供应源，'
        '不受上游缺货制约。')

    # ============================================================
    # 第2节 智慧决策实验
    # ============================================================
    add_heading(doc, '2 智慧决策实验', level=1)

    add_paragraph(doc,
        '虽然外部市场需求相对稳定，但在经典多级供应链模型下，各级节点根据直接下游的'
        '需求向上游订货，在理性决策下会产生巨大的牛鞭效应。由于DQN本身能够适应动态'
        '变化的环境，并在供应链决策上体现出较高的优越性，本文基于DQN设计了人机协同的'
        '智慧决策机器人（Intelligent Decision-Making Robot, IDMR），并将智慧决策机器人'
        '整合进供应链，用以改善分销商决策。')

    add_paragraph(doc,
        '智慧决策机器人的作用机理是基于Agent与环境进行交互，从环境中获得奖赏并反馈给'
        '动作，包含4个基本要素（策略、奖惩反馈、值函数、环境）以及一个主要要素——'
        '人类订货的最优决策经验即理性决策"老师"。决策"老师"将自己的决策经验教授给'
        '智慧决策机器人，并在发现智慧决策机器人决策失误时及时给以干扰，'
        '即人机协同机制：传授经验、限制决策、惩罚机制。')

    add_heading(doc, '2.1 实验基础设置', level=2)

    add_paragraph(doc,
        '智慧实验将智慧决策机器人接入供应链改善分销商决策，供应链中的其他节点依次'
        '按照理性决策向上游下订单，从方差比、平均成本和服务水平三个方面进行比较，'
        '分析有智慧决策机器人参与下的牛鞭效应新现象。')

    add_paragraph(doc, '为便于比较，智慧实验对顾客需求相关参数的假定与复盘时相同：')

    add_formula(doc,
        r'\rho=0.5,\ L=2,\ d=10,\ \varepsilon\sim N(0,5),\ '
        r'z=2,\ C_{L,\rho}=2,\ p=5', '8')

    add_table_note(doc, '注：每次实验进行20000个周期。IDMR训练步数为10000步，随机种子seed=42。')

    add_paragraph(doc, 'IDMR的奖励函数采用满足率作为核心激励：')

    add_formula(doc,
        r'r_t = \frac{\text{fulfilled}_t}{D_t} - \frac{\text{stockout}_t}{D_t}', '9')

    add_paragraph(doc,
        '该奖励函数直接激励IDMR最大化服务水平，与李勇论文公式11的设计思想一致。'
        '惩罚机制方面，当IDMR的积压库存达到相同需求下经典多级供应链对应节点（分销商）'
        '的平均库存时，决策"老师"禁止IDMR向上游订货。本实验中，经典供应链分销商的'
        '平均库存经2000步预热仿真测得为352.51。')

    add_heading(doc, '2.2 实验结果与分析', level=2)

    # ---- 2.2.1 方差比比较 ----
    add_heading(doc, '2.2.1 方差比比较', level=3)

    add_paragraph(doc,
        '在理性决策下，供应链上各个节点的订货量方差与需求量方差之比会出现1.3节中的'
        '情况，随着供应链层级的增加，这种由需求信息不对称带来的方差比越来越大。'
        '当智慧决策机器人接入供应链用以改善分销商的决策时，方差比有所减小，'
        '特别是分销商与制造商，但批发商的方差比也有一定程度的减小。')

    add_table_title(doc, '表2 两种决策下牛鞭效应对比（保留小数点后2位）')

    headers2 = ['节点', '理性决策\n需求均值', '理性决策\n订单均值', '理性决策\n方差比',
                '智慧决策\n需求均值', '智慧决策\n订单均值', '智慧决策\n方差比']
    rows2 = []
    for k in range(1, 5):
        rows2.append([
            NODE_NAMES[k-1],
            f"{baseline['demand_mean'][str(k)]:.2f}",
            f"{baseline['order_mean'][str(k)]:.2f}",
            f"{baseline['bwe'][str(k)]:.2f}",
            f"{exp1['demand_mean'][str(k)]:.2f}",
            f"{exp1['order_mean'][str(k)]:.2f}",
            f"{exp1['bwe'][str(k)]:.2f}",
        ])
    make_three_line_table(doc, headers2, rows2,
                          col_widths=[1.8, 1.8, 1.8, 1.5, 1.8, 1.8, 1.5])

    add_figure(doc, os.path.join(PNG_DIR, 'fig_bwe_comparison.png'),
               '图4 两种决策下的方差比对比')

    add_paragraph(doc,
        '如表2和上图所示，在需求均值与订单均值近似一致的情况下，智慧决策机器人接入'
        '供应链后，分销商方差比从67.33降至10.65，降幅达84.19%；制造商方差比从301.75'
        '降至22.70，降幅达92.48%。同时，随着顾客需求信息的逐级传递，智慧决策机器人'
        '通过观察当前状态做出最优决策，整个供应链的方差比大小顺序发生了变化：'
        '理性决策下方差比最大的节点为制造商（301.75），智慧决策下方差比最大的节点'
        '变为批发商（15.78），分销商和制造商的方差比均显著降低。')

    # ---- 2.2.2 平均成本比较 ----
    add_heading(doc, '2.2.2 平均成本比较', level=3)

    add_figure(doc, os.path.join(PNG_DIR, 'fig_cost_comparison.png'),
               '图5 两种决策下的平均成本对比')

    add_paragraph(doc,
        '在成本由库存成本和缺货成本两者构成时，理性决策下随着供应链层级的增加，'
        '零售商到制造商的平均成本逐级增加，并且相邻两级之间平均成本的差距越来越大。'
        '在智慧决策下，当智慧决策机器人接入供应链后，供应链的整体成本得到明显减小。'
        '分销商平均成本从317.34降至9.88，制造商平均成本从1212.92降至9.83，'
        '智慧决策有效减小了供应链上游成本，对制造商成本的减小作用尤为显著。'
        '零售商和批发商的平均成本也分别从32.22和102.32降至9.94和9.89，'
        '表明智慧决策对供应链整体成本的改善作用从分销商节点向上、下游两端辐射。')

    # ---- 2.2.3 服务水平比较 ----
    add_heading(doc, '2.2.3 服务水平比较', level=3)

    add_figure(doc, os.path.join(PNG_DIR, 'fig_sl_comparison.png'),
               '图6 两种决策下的服务水平对比')

    add_paragraph(doc,
        '上图表示当 z=2 时，智慧决策机器人代替分销商做决策时，各级的实际服务水平。'
        '在理性决策下，即使理论目标服务水平是97.7%，实际上由于运输提前期的存在，'
        '供应链上除制造商之外，各级都不能稳定达到97.7%的服务水平。'
        '当智慧决策机器人接入供应链用以改善分销商决策时，上游节点的服务水平有所提高：'
        '分销商SL从99.76%提升至100.00%，提高了0.24个百分点，实际服务水平达到100%，'
        '与李勇论文的结论一致。制造商SL从99.81%微降至99.75%，基本保持稳定。'
        '零售商和批发商的服务水平也略有提升，分别为98.98%和99.76%。'
        '这表明IDMR在学习"按需订货"策略的同时，并未牺牲服务水平，'
        '而是通过精准的库存管理实现了成本与服务的双重优化。')

    # ============================================================
    # 第3节 Baseline与需要改进的指标
    # ============================================================
    add_heading(doc, '3 基础实验形成的Baseline与需要改进的指标', level=1)

    add_heading(doc, '3.1 Baseline形成', level=2)

    add_paragraph(doc, '通过上述20000周期的两组对比实验，形成本研究的Baseline数据：')

    add_table_title(doc, '表3 基础实验Baseline汇总')

    headers3 = ['指标', '节点', '理性决策（Baseline）', '智慧决策（Exp_1）', '改进幅度']
    rows3 = []
    # 方差比
    for k in range(1, 5):
        b = baseline['bwe'][str(k)]; e = exp1['bwe'][str(k)]
        imp = (e - b) / b * 100
        rows3.append(['方差比BWE', NODE_NAMES[k-1], f'{b:.2f}', f'{e:.2f}', f'{imp:+.2f}%'])
    # 平均成本
    for k in range(1, 5):
        b = baseline['avg_cost'][str(k)]; e = exp1['avg_cost'][str(k)]
        imp = (e - b) / b * 100
        rows3.append(['平均成本', NODE_NAMES[k-1], f'{b:.2f}', f'{e:.2f}', f'{imp:+.2f}%'])
    # 服务水平
    for k in range(1, 5):
        b = baseline['sl'][str(k)]; e = exp1['sl'][str(k)]
        imp = (e - b) * 100
        rows3.append(['服务水平SL', NODE_NAMES[k-1], f'{b*100:.2f}%', f'{e*100:.2f}%', f'{imp:+.2f}%'])

    make_three_line_table(doc, headers3, rows3,
                          col_widths=[2.5, 2.0, 3.5, 3.5, 2.5])

    add_heading(doc, '3.2 需要改进的指标', level=2)

    add_paragraph(doc,
        '基础实验验证了IDMR在方差比、平均成本和服务水平三个指标上的显著改进效果，'
        '与李勇论文的结论一致。分销商方差比降低84.19%，制造商方差比降低92.48%，'
        '分销商服务水平提升至100%，系统总成本从1664.80降至39.54，降幅达97.63%。'
        '但也暴露出以下需要改进的问题：')

    add_paragraph(doc,
        '1. 批发商方差比改进有限：智慧决策下批发商方差比（15.78）仅较理性决策'
        '（16.25）降低2.89%，且成为智慧决策下方差比最大的节点。批发商节点未部署IDMR，'
        '仍采用理性决策，其方差比改善仅来自下游分销商订货行为变化的间接影响。')

    add_paragraph(doc,
        '2. 零售商节点几乎无改善：零售商方差比从4.05降至4.01，降幅仅0.99%，'
        '服务水平几乎不变。零售商距IDMR部署节点（分销商）最远，受益最小。')

    add_paragraph(doc,
        '3. 制造商方差比仍较高：智慧决策下制造商方差比（22.70）虽然较理性决策'
        '（301.75）大幅降低，但仍高于分销商（10.65），表明需求信息在分销商到制造商'
        '之间的传递仍存在放大效应。')

    add_paragraph(doc,
        '4. 单点优化的局限性：IDMR仅部署在分销商节点，对上游（制造商）和下游'
        '（批发商、零售商）的改善主要通过间接传导实现，改善幅度递减。'
        '这为多智能体协同决策系统的设计提供了改进方向。')

    add_paragraph(doc,
        '上述问题为后续多智能体情绪感知人智协同决策系统的设计提供了改进方向：'
        '通过多智能体协同信息共享机制提升批发商和零售商的方差比改进效果，'
        '通过情绪演化方程与正向激励函数在降低成本的同时进一步维持和提升各节点服务水平，'
        '通过动态事件注入验证系统的鲁棒性。')

    # ============================================================
    # 附录：实验配置参数
    # ============================================================
    add_heading(doc, '附录：实验配置参数', level=1)

    add_table_title(doc, '附表 实验配置参数汇总')

    headers_app = ['参数', '符号', '取值', '说明']
    rows_app = [
        ['基础需求', 'd', '10', 'AR(1)需求模型常量'],
        ['自相关系数', 'ρ', '0.5', 'AR(1)需求模型相关系数'],
        ['需求噪声标准差', 'σ_ε', '5.0', '误差项标准差'],
        ['运输提前期', 'L', '2', '订货至到货的周期数'],
        ['SMA预测窗口', 'p', '5', '移动平均窗口长度'],
        ['安全库存系数', 'z', '2', '对应97.7%理论服务水平'],
        ['需求校正系数', 'C_{L,ρ}', '2.0', 'Lee等(2000)校正系数'],
        ['供应链级数', 'K', '4', '零售商→批发商→分销商→制造商'],
        ['初始库存', 'NS_0', '10', '各节点初始库存'],
        ['仿真周期', 'T', '20000', '连续动态订货与发货周期'],
        ['IDMR训练步数', '—', '10000', 'DQN训练步数'],
        ['随机种子', 'seed', '42', '确保结果可复现'],
        ['单位库存成本', 'h', '1.0', '单位周期单位库存持有成本'],
        ['单位缺货成本', 'b', '2.0', '单位周期单位缺货惩罚成本'],
        ['惩罚阈值', '—', '352.51', '经典供应链分销商平均库存（2000步预热测得）'],
        ['动作空间', '—', '[11, 40]', 'IDMR订货量离散动作范围'],
        ['状态维度', '—', '5', '[S, WIP, q_down, Trans, q_self]'],
    ]
    make_three_line_table(doc, headers_app, rows_app,
                          col_widths=[3.0, 2.0, 3.0, 5.0])

    # ---- 保存 ----
    output_path = '基础实验_牛鞭效应对比分析.docx'
    doc.save(output_path)
    print(f"\n[OK] docx 文档已生成: {output_path}")
    return output_path


# ============================================================
# 主程序
# ============================================================

def main():
    print("=" * 70)
    print("生成《中国管理科学》规范的基础实验 docx 文档")
    print("=" * 70)

    # 1. 加载数据
    print("\n[1/4] 加载实验数据...")
    data = load_data()
    print(f"  Baseline: BWE={data['baseline']['bwe']}")
    print(f"  Exp_1:    BWE={data['exp1']['bwe']}")

    # 2. 生成对比柱状图 PNG
    print("\n[2/4] 生成对比柱状图 PNG...")
    generate_comparison_pngs(data)

    # 3. 生成时序图 PNG（重新运行baseline）
    print("\n[3/4] 生成时序图 PNG...")
    generate_timeseries_pngs()

    # 4. 生成 docx 文档
    print("\n[4/4] 生成 docx 文档...")
    output = generate_docx()

    print(f"\n{'=' * 70}")
    print(f"完成！文档已保存: {output}")
    print(f"{'=' * 70}")


if __name__ == '__main__':
    main()
