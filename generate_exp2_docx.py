"""
生成《中国管理科学》规范的 Exp_2 综合 docx 文档
================================================
将 Exp_2综合分析_人智协同智慧决策.md 转换为 Word 文档

学术规范:
  - 三线表（顶线1.5pt / 表头线0.75pt / 底线1.5pt）
  - OMML 原生公式（LaTeX → MathML → OMML）
  - 嵌入 7 张图表（PNG 格式）
  - 中文字体：宋体（正文）/ 黑体（标题），英文：Times New Roman
  - 字号：章标题三号黑体，节标题四号黑体，正文五号宋体，
          表标题五号黑体加粗，表内文字五号宋体，表注小五号宋体
  - A4 纵向，上下边距 2.54cm，左右边距 3.17cm

工具函数（make_three_line_table / latex_to_omml / add_heading 等）
直接复用自 generate_basic_experiment_docx.py。

数据来源:
  - p0_results/四组对比_20k.json        四组实验对比数据（含剥离实验Exp_1b）
  - p0_results/归因分析_20k.json         Exp_2归因分析数据
  - p0_results/归因分析_exp1b.json       Exp_1b归因分析数据
  - 灾难性遗忘_结果摘要.json             持续学习数据
  - p0_results/exp2_20k_timeseries.json  Exp_2 逐周期时序数据
  - svg_figures_exp2/*.svg               Exp_2 图表（因 cairosvg 不可用，
                                         改用 matplotlib 从 JSON 重新生成 PNG）
  - 灾难性遗忘_性能曲线.png / 灾难性遗忘_噪声鲁棒性.png  持续学习图表
"""

import os
import json
import math
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import networkx as nx
from matplotlib.patches import FancyArrowPatch

matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
matplotlib.rcParams['axes.unicode_minus'] = False
matplotlib.rcParams['font.size'] = 11

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 复用基础实验脚本中的工具函数与样式配置
from generate_basic_experiment_docx import (
    latex_to_omml, _OMML_AVAILABLE,
    set_run_font, set_cell_border, set_cell_text,
    add_heading, add_paragraph, add_formula, add_inline_formula,
    add_table_title, add_table_note, add_figure,
    make_three_line_table,
    FONT_CN, FONT_EN, FONT_HEADING,
)

# ============================================================
# 全局配置
# ============================================================
NODE_NAMES = ['零售商', '批发商', '分销商', '制造商']

COLOR_BASELINE = '#E74C3C'   # 理性决策 - 红
COLOR_EXP1 = '#3498DB'       # 智慧决策 - 蓝
COLOR_EXP1B = '#9B59B6'      # Exp_1b  - 紫
COLOR_EXP2 = '#27AE60'       # Exp_2   - 绿
COLOR_TS = ['#E74C3C', '#E67E22', '#F39C12', '#C0392B']

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FOUR = os.path.join(BASE_DIR, 'p0_results', '四组对比_20k.json')
DATA_ATTR = os.path.join(BASE_DIR, 'p0_results', '归因分析_20k.json')
DATA_ATTR_EXP1B = os.path.join(BASE_DIR, 'p0_results', '归因分析_exp1b.json')
DATA_FORGET = os.path.join(BASE_DIR, '灾难性遗忘_结果摘要.json')
DATA_TS = os.path.join(BASE_DIR, 'p0_results', 'exp2_20k_timeseries.json')

FIG_DIR = os.path.join(BASE_DIR, 'svg_figures_exp2')
PNG_DIR = os.path.join(FIG_DIR, 'png_temp')
os.makedirs(PNG_DIR, exist_ok=True)

# 持续学习已存在的 PNG（位于根目录）
PNG_PERF_CURVE = os.path.join(BASE_DIR, '灾难性遗忘_性能曲线.png')
PNG_NOISE = os.path.join(BASE_DIR, '灾难性遗忘_噪声鲁棒性.png')

TOTAL_PERIODS = 20000


# ============================================================
# 1. 数据加载
# ============================================================

def load_all_data():
    """加载四组对比 / 归因分析 / 持续学习 / 时序数据"""
    with open(DATA_FOUR, 'r', encoding='utf-8') as f:
        four = json.load(f)
    with open(DATA_ATTR, 'r', encoding='utf-8') as f:
        attr = json.load(f)
    attr_exp1b = None
    if os.path.exists(DATA_ATTR_EXP1B):
        with open(DATA_ATTR_EXP1B, 'r', encoding='utf-8') as f:
            attr_exp1b = json.load(f)
    with open(DATA_FORGET, 'r', encoding='utf-8') as f:
        forget = json.load(f)
    ts = None
    if os.path.exists(DATA_TS):
        with open(DATA_TS, 'r', encoding='utf-8') as f:
            ts = json.load(f)
    return four, attr, attr_exp1b, forget, ts


# ============================================================
# 2. 图表生成 (matplotlib → PNG)
#    cairosvg 不可用，全部从 JSON 重新生成
# ============================================================

def _save_fig(fig, name):
    path = os.path.join(PNG_DIR, name)
    fig.savefig(path, dpi=200, bbox_inches='tight')
    plt.close(fig)
    print(f"  [图] {name} 已生成")
    return path


def generate_fig1_bwe(four):
    """图1 四组实验方差比对比柱状图（对数纵轴）"""
    baseline = four['baseline']['bwe']
    exp1 = four['exp1']['bwe']
    exp1b = four['exp1b']['bwe']
    exp2 = four['exp2']['bwe']
    x = np.arange(4)
    width = 0.2

    b = [baseline[str(k)] for k in range(1, 5)]
    e1 = [exp1[str(k)] for k in range(1, 5)]
    e1b = [exp1b[str(k)] for k in range(1, 5)]
    e2 = [exp2[str(k)] for k in range(1, 5)]

    fig, ax = plt.subplots(figsize=(9.5, 5))
    bars1 = ax.bar(x - 1.5*width, b, width, label='Baseline 理性决策',
                   color=COLOR_BASELINE, edgecolor='black', linewidth=0.8)
    bars2 = ax.bar(x - 0.5*width, e1, width, label='Exp_1 智慧决策',
                   color=COLOR_EXP1, edgecolor='black', linewidth=0.8)
    bars3 = ax.bar(x + 0.5*width, e1b, width, label='Exp_1b +情绪',
                   color=COLOR_EXP1B, edgecolor='black', linewidth=0.8)
    bars4 = ax.bar(x + 1.5*width, e2, width, label='Exp_2 人智协同',
                   color=COLOR_EXP2, edgecolor='black', linewidth=0.8)

    for bars in (bars1, bars2, bars3, bars4):
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2, h * 1.05,
                    f'{h:.2f}', ha='center', va='bottom', fontsize=7)

    ax.set_yscale('log')
    ax.set_xlabel('供应链节点', fontsize=12)
    ax.set_ylabel('方差比 BWE（对数刻度）', fontsize=12)
    ax.set_title('四组实验方差比对比', fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(NODE_NAMES, fontsize=11)
    ax.legend(fontsize=9, loc='upper left')
    ax.grid(axis='y', alpha=0.3, linestyle='--', which='both')
    ax.set_ylim(1, 500)
    plt.tight_layout()
    return _save_fig(fig, 'fig1_bwe_comparison_4groups.png')


def generate_fig2_sl(four):
    """图2 四组实验服务水平对比柱状图"""
    baseline = four['baseline']['sl']
    exp1 = four['exp1']['sl']
    exp1b = four['exp1b']['sl']
    exp2 = four['exp2']['sl']
    x = np.arange(4)
    width = 0.2

    b = [baseline[str(k)] * 100 for k in range(1, 5)]
    e1 = [exp1[str(k)] * 100 for k in range(1, 5)]
    e1b = [exp1b[str(k)] * 100 for k in range(1, 5)]
    e2 = [exp2[str(k)] * 100 for k in range(1, 5)]

    fig, ax = plt.subplots(figsize=(9.5, 5))
    bars1 = ax.bar(x - 1.5*width, b, width, label='Baseline 理性决策',
                   color=COLOR_BASELINE, edgecolor='black', linewidth=0.8)
    bars2 = ax.bar(x - 0.5*width, e1, width, label='Exp_1 智慧决策',
                   color=COLOR_EXP1, edgecolor='black', linewidth=0.8)
    bars3 = ax.bar(x + 0.5*width, e1b, width, label='Exp_1b +情绪',
                   color=COLOR_EXP1B, edgecolor='black', linewidth=0.8)
    bars4 = ax.bar(x + 1.5*width, e2, width, label='Exp_2 人智协同',
                   color=COLOR_EXP2, edgecolor='black', linewidth=0.8)

    for bars in (bars1, bars2, bars3, bars4):
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.3,
                    f'{h:.2f}%', ha='center', va='bottom', fontsize=7)

    ax.axhline(y=97.7, color='gray', linestyle='--', linewidth=1, alpha=0.7)
    ax.text(3.45, 97.9, '理论目标97.7%', fontsize=9, color='gray')
    ax.set_xlabel('供应链节点', fontsize=12)
    ax.set_ylabel('服务水平 SL (%)', fontsize=12)
    ax.set_title('四组实验服务水平对比', fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(NODE_NAMES, fontsize=11)
    ax.legend(fontsize=9, loc='lower left')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_ylim(75, 106)
    plt.tight_layout()
    return _save_fig(fig, 'fig2_sl_comparison_4groups.png')


def generate_fig3_emotion_corr(ts, attr):
    """图3 情绪-决策相关性散点图（4节点子图，抽样）"""
    corr = attr['emotion_decision_correlation']
    fig, axes = plt.subplots(2, 2, figsize=(10, 8))
    axes = axes.flatten()
    step = 20  # 抽样

    for idx, k in enumerate(range(1, 5)):
        ax = axes[idx]
        emo = ts['emotion_history'][str(k)][::step]
        order = ts['order_history'][str(k)][::step]
        ax.scatter(emo, order, s=6, alpha=0.35,
                   color=COLOR_TS[idx], edgecolors='none')
        r = corr[str(k)]['pearson_emotion_order']
        ax.set_title(f'{NODE_NAMES[idx]}  r(E,q)={r:.3f}',
                     fontsize=11, fontweight='bold')
        ax.set_xlabel('情绪状态 E', fontsize=10)
        ax.set_ylabel('订货量 q', fontsize=10)
        ax.axhline(y=0, color='gray', linewidth=0.5, alpha=0.5)
        ax.axvline(x=0, color='gray', linewidth=0.5, alpha=0.5)
        ax.grid(alpha=0.3, linestyle='--')

    fig.suptitle('各节点情绪状态与订货决策相关性', fontsize=13, fontweight='bold')
    plt.tight_layout()
    return _save_fig(fig, 'fig3_emotion_decision_corr.png')


def generate_fig4_contagion_network(attr):
    """图4 情绪传染网络图"""
    ca = attr['contagion_analysis']
    path_counts = ca['path_counts']
    total = ca['total_events']

    G = nx.DiGraph()
    for name in NODE_NAMES:
        G.add_node(name)

    edges = []
    for path, cnt in path_counts.items():
        src, dst = path.split('→')
        G.add_edge(src, dst, weight=cnt)
        edges.append((src, dst, cnt))

    fig, ax = plt.subplots(figsize=(8.5, 5))
    # 水平布局：零售商在左，制造商在右
    pos = {NODE_NAMES[i]: (i * 2.0, 0) for i in range(4)}

    # 节点颜色：下游恐慌更重，颜色更深红
    node_colors = ['#C0392B', '#E67E22', '#F39C12', '#F1C40F']
    nx.draw_networkx_nodes(G, pos, ax=ax, node_size=2200,
                           node_color=node_colors,
                           edgecolors='black', linewidths=1.5)
    nx.draw_networkx_labels(G, pos, ax=ax, font_size=11,
                            font_color='white', font_weight='bold')

    max_cnt = max(e[2] for e in edges)
    for src, dst, cnt in edges:
        width = 1.0 + 6.0 * cnt / max_cnt
        rad = 0.15
        # 正向箭头
        arrow = FancyArrowPatch(
            pos[src], pos[dst],
            connectionstyle=f'arc3,rad={rad}',
            arrowstyle='-|>', mutation_scale=22,
            linewidth=width, color='#2C3E50', alpha=0.75,
            shrinkA=30, shrinkB=30)
        ax.add_patch(arrow)
        # 边权标注
        mid_x = (pos[src][0] + pos[dst][0]) / 2
        mid_y = (pos[src][1] + pos[dst][1]) / 2 + 0.35
        ax.text(mid_x, mid_y, f'{cnt}\n({cnt/total*100:.1f}%)',
                ha='center', va='center', fontsize=9,
                bbox=dict(boxstyle='round,pad=0.3',
                          facecolor='white', edgecolor='#2C3E50', alpha=0.9))

    ax.set_title(f'情绪传染网络（共 {total} 次事件）',
                 fontsize=13, fontweight='bold')
    ax.set_xlim(-1, 7)
    ax.set_ylim(-1.5, 1.5)
    ax.axis('off')
    ax.text(0, -1.1, '恐慌蔓延方向 →', fontsize=10,
            color='#2C3E50', ha='center', style='italic')
    plt.tight_layout()
    return _save_fig(fig, 'fig4_contagion_network.png')


def generate_fig5_blocking(attr):
    """图5 正向激励阻断效应图（恐慌 vs 中性过度订货比例）"""
    over = attr['overorder_analysis']
    panic = [over['panic'][str(k)] * 100 for k in range(1, 5)]
    neutral = [over['neutral'][str(k)] * 100 for k in range(1, 5)]
    x = np.arange(4)
    width = 0.35

    fig, ax = plt.subplots(figsize=(8.5, 5))
    bars1 = ax.bar(x - width / 2, panic, width, label='恐慌状态',
                   color=COLOR_EXP1, edgecolor='black', linewidth=0.8)
    bars2 = ax.bar(x + width / 2, neutral, width, label='中性状态',
                   color=COLOR_EXP2, edgecolor='black', linewidth=0.8)

    for bars in (bars1, bars2):
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2, h + 1,
                    f'{h:.2f}%', ha='center', va='bottom', fontsize=9)

    # 标注阻断效应
    for k in range(4):
        diff = neutral[k] - panic[k]
        sign = '+' if diff >= 0 else ''
        y_pos = max(panic[k], neutral[k]) + 8
        ax.text(x[k], y_pos, f'阻断效应\n{sign}{diff:.2f}pp',
                ha='center', va='bottom', fontsize=8,
                color='#C0392B' if diff >= 0 else '#7F8C8D',
                fontweight='bold')

    ax.axhline(y=0, color='black', linewidth=0.6)
    ax.set_xlabel('供应链节点', fontsize=12)
    ax.set_ylabel('过度订货比例 (%)', fontsize=12)
    ax.set_title('恐慌 vs 中性状态下的过度订货比例（阻断效应）',
                 fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(NODE_NAMES, fontsize=11)
    ax.legend(fontsize=10, loc='upper right')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_ylim(0, 110)
    plt.tight_layout()
    return _save_fig(fig, 'fig5_blocking_effect.png')


def generate_fig6_ablation(four):
    """图6 剥离实验效应分解图（分销商与制造商BWE对比）"""
    exp1 = four['exp1']['bwe']
    exp1b = four['exp1b']['bwe']
    exp2 = four['exp2']['bwe']

    k3_vals = [exp1['3'], exp1b['3'], exp2['3']]
    k4_vals = [exp1['4'], exp1b['4'], exp2['4']]
    labels = ['Exp_1\n(无情绪无协同)', 'Exp_1b\n(+情绪)', 'Exp_2\n(+协同+动态)']
    x = np.arange(3)
    width = 0.35

    fig, ax = plt.subplots(figsize=(8.5, 5))
    bars1 = ax.bar(x - width/2, k3_vals, width, label='分销商 (k=3)',
                   color=COLOR_EXP1, edgecolor='black', linewidth=0.8)
    bars2 = ax.bar(x + width/2, k4_vals, width, label='制造商 (k=4)',
                   color=COLOR_EXP2, edgecolor='black', linewidth=0.8)

    for bars in (bars1, bars2):
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.3,
                    f'{h:.2f}', ha='center', va='bottom', fontsize=9)

    # 标注效应分解箭头
    ax.annotate('', xy=(1, k3_vals[1]+0.5), xytext=(0, k3_vals[0]+0.5),
                arrowprops=dict(arrowstyle='->', color=COLOR_EXP1B, lw=2))
    ax.text(0.5, max(k3_vals[0], k3_vals[1])+1.5, '情绪效应\n-1.3%',
            ha='center', fontsize=8, color=COLOR_EXP1B, fontweight='bold')

    ax.annotate('', xy=(2, k4_vals[2]+0.5), xytext=(1, k4_vals[1]+0.5),
                arrowprops=dict(arrowstyle='->', color=COLOR_EXP2, lw=2))
    ax.text(1.5, max(k4_vals[1], k4_vals[2])+1.5, '协同效应\n-55.8%',
            ha='center', fontsize=8, color=COLOR_EXP2, fontweight='bold')

    ax.set_xlabel('实验组', fontsize=12)
    ax.set_ylabel('方差比 BWE', fontsize=12)
    ax.set_title('剥离实验效应分解', fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=10)
    ax.legend(fontsize=10, loc='upper right')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_ylim(0, max(k4_vals) * 1.3)
    plt.tight_layout()
    return _save_fig(fig, 'fig6_ablation_summary.png')


def generate_all_figures(four, attr, forget, ts):
    """生成全部 PNG 图表"""
    print("  [图] 从 JSON 重新生成 PNG 图表（cairosvg 不可用）...")
    paths = {}
    paths['fig1'] = generate_fig1_bwe(four)
    paths['fig2'] = generate_fig2_sl(four)
    if ts is not None:
        paths['fig3'] = generate_fig3_emotion_corr(ts, attr)
    else:
        paths['fig3'] = None
    paths['fig4'] = generate_fig4_contagion_network(attr)
    paths['fig5'] = generate_fig5_blocking(attr)
    paths['fig6'] = generate_fig6_ablation(four)
    # 图7 / 图8 使用持续学习已生成的 PNG
    paths['fig7'] = PNG_PERF_CURVE if os.path.exists(PNG_PERF_CURVE) else None
    paths['fig8'] = PNG_NOISE if os.path.exists(PNG_NOISE) else None
    return paths


# ============================================================
# 3. 文档生成主函数
# ============================================================

def _pct(v, decimals=2):
    return f'{v * 100:.{decimals}f}%'


def generate_docx(four, attr, attr_exp1b, forget, fig_paths):
    """生成完整的 Exp_2 docx 文档"""
    baseline = four['baseline']
    exp1 = four['exp1']
    exp1b = four['exp1b']
    exp2 = four['exp2']
    corr = attr['emotion_decision_correlation']
    ca = attr['contagion_analysis']
    over = attr['overorder_analysis']
    summ = forget['summary']
    perc = forget['perception_stats']['C_有EWC有噪声']

    doc = Document()

    # ---- 页面设置：A4 纵向 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ---- 文档标题 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run('Exp_2综合分析：人智协同智慧决策系统对牛鞭效应的缓解机制')
    set_run_font(run, font_en=FONT_EN, font_cn=FONT_HEADING,
                 size=18, bold=True)

    # 摘要式引言
    add_paragraph(doc,
        '实验周期：20000 cycles × 4 nodes = 80000 steps；随机种子 42；'
        '仿真框架 PettingZoo AECEnv。在基础实验（Baseline 与 Exp_1）'
        '完成20000周期对标验证的基础上，本研究设计四组对比实验（含剥离实验'
        ' Exp_1b），通过控制变量法精确分离情绪机制与协同通信的独立效应。'
        'Exp_2 引入人智协同三要素——情绪演化模块、正向激励机制与多智能体'
        '协同通信，并叠加76次动态突发事件（53次需求突变 + 23次供应中断），'
        '验证损失厌恶放大、正向激励阻断与协同鲁棒性三个研究假设。')

    # ============================================================
    # 第1节 实验概述
    # ============================================================
    add_heading(doc, '1 实验概述', level=1)

    add_heading(doc, '1.1 实验目标', level=2)

    add_paragraph(doc,
        '在基础实验（Baseline + Exp_1）已完成20000周期对标验证的基础上，'
        '本研究设计四组对比实验（含剥离实验 Exp_1b），通过控制变量法精确'
        '分离情绪机制与协同通信的独立效应。Exp_2 引入人智协同三要素——'
        '情绪演化模块、正向激励机制与多智能体协同通信，并叠加76次动态'
        '突发事件（53次需求突变 + 23次供应中断），验证以下三个研究假设：')

    add_paragraph(doc,
        'H1（损失厌恶放大假设）：动态环境中的缺货事件通过情绪演化方程将'
        '决策者情绪状态推向恐慌饱和（E_t → -1），使有效缺货惩罚权重放大约70%，'
        '导致订货决策偏离理性最优，加剧牛鞭效应。', indent=False)
    add_paragraph(doc,
        'H2（正向激励阻断假设）：库存精准匹配正向激励函数通过将 DQN 优化目标'
        '从"最小化缺货"转变为"最大化精准匹配"，在系统层面阻断'
        '"缺货→恐慌→过度订货→上游缺货→更大恐慌"的恶性循环，'
        '使分销商 BWE 降低80%以上。', indent=False)
    add_paragraph(doc,
        'H3（协同鲁棒性假设）：多智能体协同信息共享机制使上游节点能够提前感知'
        '下游需求突变，相较于单智能体方案，制造商 BWE 进一步降低45%以上，'
        '系统平均服务水平提升至99%以上，且在情绪感知噪声干扰下通过 EWC 与 PER '
        '机制仍保持稳定。', indent=False)

    add_heading(doc, '1.2 实验配置', level=2)

    add_paragraph(doc,
        'Exp_2 在 PettingZoo AECEnv 框架下运行，四级供应链'
        '（零售商→批发商→分销商→制造商）每个节点部署独立的情绪感知机器人，'
        '关键参数如下表所示。')

    add_table_title(doc, '实验配置参数')
    headers_cfg = ['参数类别', '参数名称', '取值', '说明']
    rows_cfg = [
        ['需求模型', 'AR(1)参数', 'd=10, ρ=0.5, σ_ε=5', '与 Baseline 完全一致'],
        ['情绪演化', '惯性系数 α', '0.7', '情绪粘性，越大越持久'],
        ['情绪演化', '敏感度 γ', '2.0', '对反馈信号的响应强度'],
        ['情绪传染', '传染概率', '0.30', '30% 概率向上游传染'],
        ['情绪传染', '传染强度 η', '0.40', '恐慌蔓延幅度'],
        ['动态事件', '需求突变概率', '0.00265', '约 53 次/20000 周期'],
        ['动态事件', '供应中断概率', '0.00115', '约 23 次/20000 周期'],
        ['情绪调节', '恐慌放大系数', '1.0 + 0.5|E_t|', 'E_t<0 时订货量放大'],
        ['情绪调节', '乐观缩减系数', '1.0 - 0.3 E_t', 'E_t>0 时过度部分缩减'],
    ]
    make_three_line_table(doc, headers_cfg, rows_cfg,
                          col_widths=[2.6, 2.6, 3.4, 4.0])

    add_paragraph(doc,
        '其中，顾客需求遵循一阶自回归过程 AR(1)：')
    add_formula(doc, r'D_t = d + \rho D_{t-1} + \varepsilon_t', '1')

    add_paragraph(doc,
        '情绪演化方程采用双曲正切函数刻画情绪状态的饱和特征：')
    add_formula(doc, r'E_t = \tanh\left(\alpha E_{t-1} + \gamma \Phi_t\right)', '2')

    add_paragraph(doc,
        '其中，α 为情绪惯性系数，γ 为敏感度系数，Φ_t 为反馈信号。'
        '当情绪状态 E_t<0 时（恐慌），订货量按恐慌放大系数放大：')
    add_formula(doc, r'q_t^{panic} = q_t \cdot \left(1.0 + 0.5|E_t|\right)', '3')

    add_paragraph(doc,
        '当 E_t>0 时（乐观），订货量中超出需求的部分按乐观缩减系数缩减：')
    add_formula(doc, r'q_t^{optimistic} = D_t + (q_t - D_t)\left(1.0 - 0.3 E_t\right)', '4')

    add_paragraph(doc,
        '为验证系统的感知鲁棒性，持续学习实验在情绪状态输入端叠加情绪感知噪声：')
    add_formula(doc,
        r'E_{perceived} = \mathrm{clip}\left(E_{true} + \mathcal{N}(0, \sigma_{noise}),\ -1,\ 1\right)', '5')

    add_paragraph(doc, '牛鞭效应与服务水平定义为：')
    add_formula(doc, r'BWE_k = \frac{\mathrm{var}(q_k)}{\mathrm{var}(D)}, \quad SL = \frac{\mathrm{fulfilled}}{D}', '6')

    # ============================================================
    # 第2节 四组对比实验结果
    # ============================================================
    add_heading(doc, '2 四组对比实验结果', level=1)

    # ---- 2.1 方差比对比 ----
    add_heading(doc, '2.1 方差比（BWE）对比', level=2)

    add_table_title(doc, '表1 四组实验方差比对比（20000周期）')
    headers2 = ['节点', 'Baseline', 'Exp_1', 'Exp_1b', 'Exp_2',
                '情绪效应(1→1b)', '协同效应(1b→2)']
    rows2 = []
    for k in range(1, 5):
        b = baseline['bwe'][str(k)]
        e1 = exp1['bwe'][str(k)]
        e1b = exp1b['bwe'][str(k)]
        e2 = exp2['bwe'][str(k)]
        emo_eff = (e1b - e1) / e1 * 100 if e1 != 0 else 0
        coo_eff = (e2 - e1b) / e1b * 100 if e1b != 0 else 0
        rows2.append([
            NODE_NAMES[k - 1],
            f'{b:.2f}', f'{e1:.2f}', f'{e1b:.2f}', f'{e2:.2f}',
            f'{emo_eff:+.1f}%', f'{coo_eff:+.1f}%',
        ])
    make_three_line_table(doc, headers2, rows2,
                          col_widths=[1.5, 1.7, 1.7, 1.7, 1.7, 2.2, 2.2])

    if fig_paths.get('fig1'):
        add_figure(doc, fig_paths['fig1'], '图1 四组实验方差比对比')

    add_paragraph(doc,
        '剥离实验揭示，情绪机制独立效应（Exp_1→Exp_1b）在各节点上均极小'
        '（|变化|<1.3%），说明在平稳 AR(1) 需求下 DQN 已学到近似最优策略，'
        '情绪模块对 BWE 的边际影响可忽略。协同通信独立效应（Exp_1b→Exp_2）'
        '在批发商（-47.9%）和制造商（-55.8%）节点上表现显著，证明上游 BWE '
        '降幅主要由协同机制贡献。制造商 BWE 从 Baseline 的301.75降至 Exp_2 '
        '的10.07（-96.66%），满足 H3 中"制造商 BWE 进一步降低45%以上"的条件。')

    # ---- 2.2 服务水平对比 ----
    add_heading(doc, '2.2 服务水平（SL）对比', level=2)

    add_table_title(doc, '表2 四组实验服务水平对比（20000周期）')
    headers3 = ['节点', 'Baseline', 'Exp_1', 'Exp_1b', 'Exp_2',
                '情绪效应(1→1b)', '协同效应(1b→2)']
    rows3 = []
    for k in range(1, 5):
        b = baseline['sl'][str(k)]
        e1 = exp1['sl'][str(k)]
        e1b = exp1b['sl'][str(k)]
        e2 = exp2['sl'][str(k)]
        emo_eff = (e1b - e1) * 100
        coo_eff = (e2 - e1b) * 100
        rows3.append([
            NODE_NAMES[k - 1],
            _pct(b), _pct(e1), _pct(e1b), _pct(e2),
            f'{emo_eff:+.2f}pp', f'{coo_eff:+.2f}pp',
        ])
    sys_b = sum(baseline['sl'][str(k)] for k in range(1, 5)) / 4
    sys_e1 = sum(exp1['sl'][str(k)] for k in range(1, 5)) / 4
    sys_e1b = sum(exp1b['sl'][str(k)] for k in range(1, 5)) / 4
    sys_e2 = sum(exp2['sl'][str(k)] for k in range(1, 5)) / 4
    rows3.append([
        '系统均值', _pct(sys_b), _pct(sys_e1), _pct(sys_e1b), _pct(sys_e2),
        f'{(sys_e1b - sys_e1) * 100:+.2f}pp',
        f'{(sys_e2 - sys_e1b) * 100:+.2f}pp',
    ])
    make_three_line_table(doc, headers3, rows3,
                          col_widths=[1.5, 1.7, 1.7, 1.7, 1.7, 2.2, 2.2])

    if fig_paths.get('fig2'):
        add_figure(doc, fig_paths['fig2'], '图2 四组实验服务水平对比')

    add_paragraph(doc,
        'Exp_1b 系统平均 SL（99.61%）与 Exp_1（99.62%）几乎一致，证明情绪'
        '机制在无动态事件时对服务水平无显著影响。Exp_2 系统平均 SL 降至'
        '94.73%，主要因零售商 SL 骤降16.20pp（98.96%→82.76%），这是76次'
        '动态突发事件与情绪放大效应叠加的结果。剥离实验精确归因表明，'
        'SL 下降主要由动态事件触发而非协同机制本身缺陷。')

    # ---- 2.3 平均成本对比 ----
    add_heading(doc, '2.3 平均成本对比', level=2)

    add_table_title(doc, '表3 四组实验平均成本对比（20000周期）')
    headers4 = ['节点', 'Baseline', 'Exp_1*', 'Exp_1b', 'Exp_2', '协同效应(1b→2)']
    rows4 = []
    for k in range(1, 5):
        b = baseline['avg_cost'][str(k)]
        e1 = exp1['avg_cost'][str(k)]
        e1b = exp1b['avg_cost'][str(k)]
        e2 = exp2['avg_cost'][str(k)]
        coo_eff = (e2 - e1b) / e1b * 100 if e1b != 0 else 0
        rows4.append([
            NODE_NAMES[k - 1],
            f'{b:.2f}', f'{e1:.2f}', f'{e1b:.2f}', f'{e2:.2f}',
            f'{coo_eff:+.1f}%',
        ])
    coo_total = (exp2["total_cost"] - exp1b["total_cost"]) / exp1b["total_cost"] * 100
    rows4.append([
        '系统总成本',
        f'{baseline["total_cost"]:.2f}',
        f'{exp1["total_cost"]:.2f}',
        f'{exp1b["total_cost"]:.2f}',
        f'{exp2["total_cost"]:.2f}',
        f'{coo_total:+.1f}%',
    ])
    make_three_line_table(doc, headers4, rows4,
                          col_widths=[1.8, 1.8, 1.8, 1.8, 1.8, 2.5])

    add_paragraph(doc,
        '注：Exp_1 成本采用简化公式（与李勇等[5]一致），Baseline、Exp_1b、'
        'Exp_2 均采用库存持有+缺货惩罚公式，故 Exp_1 成本不与其他三组直接'
        '比较。Exp_1b 系统总成本（634.96）相较 Baseline（1664.80）降低'
        '61.9%，但高于 Exp_2（458.34）。协同机制使 Exp_2 在引入动态事件的'
        '情况下仍实现系统成本净降低27.8%（相对 Exp_1b）。', indent=False)

    # ---- 2.4 剥离实验效应分解 ----
    add_heading(doc, '2.4 剥离实验效应分解', level=2)

    add_table_title(doc, '表4 剥离实验效应分解（20000周期）')
    headers_abl = ['分析对象', '指标', '情绪效应(1→1b)', '协同效应(1b→2)', '联合效应(1→2)']
    rows_abl = []
    # 分销商 BWE
    e1_k3 = exp1['bwe']['3']; e1b_k3 = exp1b['bwe']['3']; e2_k3 = exp2['bwe']['3']
    rows_abl.append(['分销商', 'BWE变化',
        f'{e1_k3:.2f}→{e1b_k3:.2f}（{(e1b_k3-e1_k3)/e1_k3*100:+.1f}%）',
        f'{e1b_k3:.2f}→{e2_k3:.2f}（{(e2_k3-e1b_k3)/e1b_k3*100:+.1f}%）',
        f'{e1_k3:.2f}→{e2_k3:.2f}（{(e2_k3-e1_k3)/e1_k3*100:+.1f}%）'])
    # 分销商 SL
    s1_k3 = exp1['sl']['3']; s1b_k3 = exp1b['sl']['3']; s2_k3 = exp2['sl']['3']
    rows_abl.append(['分销商', 'SL变化',
        f'{s1_k3*100:.2f}→{s1b_k3*100:.2f}（{(s1b_k3-s1_k3)*100:+.2f}pp）',
        f'{s1b_k3*100:.2f}→{s2_k3*100:.2f}（{(s2_k3-s1b_k3)*100:+.2f}pp）',
        f'{s1_k3*100:.2f}→{s2_k3*100:.2f}（{(s2_k3-s1_k3)*100:+.2f}pp）'])
    # 制造商 BWE
    e1_k4 = exp1['bwe']['4']; e1b_k4 = exp1b['bwe']['4']; e2_k4 = exp2['bwe']['4']
    rows_abl.append(['制造商', 'BWE变化',
        f'{e1_k4:.2f}→{e1b_k4:.2f}（{(e1b_k4-e1_k4)/e1_k4*100:+.1f}%）',
        f'{e1b_k4:.2f}→{e2_k4:.2f}（{(e2_k4-e1b_k4)/e1b_k4*100:+.1f}%）',
        f'{e1_k4:.2f}→{e2_k4:.2f}（{(e2_k4-e1_k4)/e1_k4*100:+.1f}%）'])
    # 制造商 SL
    s1_k4 = exp1['sl']['4']; s1b_k4 = exp1b['sl']['4']; s2_k4 = exp2['sl']['4']
    rows_abl.append(['制造商', 'SL变化',
        f'{s1_k4*100:.2f}→{s1b_k4*100:.2f}（{(s1b_k4-s1_k4)*100:+.2f}pp）',
        f'{s1b_k4*100:.2f}→{s2_k4*100:.2f}（{(s2_k4-s1b_k4)*100:+.2f}pp）',
        f'{s1_k4*100:.2f}→{s2_k4*100:.2f}（{(s2_k4-s1_k4)*100:+.2f}pp）'])
    make_three_line_table(doc, headers_abl, rows_abl,
                          col_widths=[1.5, 1.5, 3.2, 3.2, 3.2])

    if fig_paths.get('fig6'):
        add_figure(doc, fig_paths['fig6'], '图3 剥离实验效应分解')

    add_paragraph(doc,
        '（1）情绪机制独立效应极小但机制存在性成立。Exp_1b 相较 Exp_1，'
        '分销商 BWE 仅变化-1.3%，制造商 BWE 微升+0.4%，说明 DQN 在平稳'
        '环境下已学到近似最优策略，情绪模块对 BWE 的边际影响可忽略。但 '
        'Exp_1b 中分销商情绪均值 E=-0.778（95.9%恐慌占比），证明情绪演化'
        '方程确实将决策者推向恐慌饱和状态，H1 的机制存在性得到验证。',
        indent=False)

    add_paragraph(doc,
        '（2）协同通信贡献主要 BWE 降幅。Exp_2 相较 Exp_1b，制造商 BWE '
        '从22.78骤降至10.07（-55.8%），远超情绪机制的独立贡献（+0.4%），'
        '证明 H3 中"制造商 BWE 进一步降低45%以上"的假设主要由协同通信'
        '机制实现。', indent=False)

    add_paragraph(doc,
        '（3）SL 下降归因于动态事件。Exp_1b 系统 SL（99.61%）与 Exp_1'
        '（99.62%）几乎一致，而 Exp_2 系统 SL 降至94.73%，证明 SL 下降'
        '并非协同机制本身导致，而是76次动态突发事件与情绪放大效应叠加'
        '的结果。', indent=False)

    # ---- 2.5 情绪波动指数 ----
    add_heading(doc, '2.5 情绪波动指数', level=2)

    add_table_title(doc, '表5 Exp_1b 与 Exp_2 各节点情绪波动指数对比')
    headers5 = ['节点', '情绪方差 σ_E²', '情绪标准差 σ_E', '恐慌占比', '乐观占比']
    rows5 = []
    for k in range(1, 5):
        var = exp2['emotion_variance'][str(k)]
        std = math.sqrt(var)
        c = corr[str(k)]
        total_count = c['panic_count'] + c['neutral_count'] + c['optimistic_count']
        panic_ratio = c['panic_count'] / TOTAL_PERIODS * 100
        opt_ratio = c['optimistic_count'] / TOTAL_PERIODS * 100
        rows5.append([
            NODE_NAMES[k - 1],
            f'{var:.4f}', f'{std:.4f}',
            f'{panic_ratio:.2f}%', f'{opt_ratio:.2f}%',
        ])
    make_three_line_table(doc, headers5, rows5,
                          col_widths=[2.0, 3.0, 3.0, 2.5, 2.5])

    add_paragraph(doc,
        'Exp_1b 分销商情绪方差（0.0305）仅为 Exp_2（0.0634）的48.1%，'
        '说明76次动态突发事件显著放大了情绪波动。Exp_2 情绪波动指数呈显著的'
        '下游放大梯度——零售商 σ_E（0.5922）是制造商（0.2205）的2.69倍。'
        '这一梯度模式验证了 H1 的核心机制：需求突变首先冲击零售商，恐慌情绪'
        '通过传染机制逐级向上游蔓延，但传染强度在每一级衰减，形成"恐慌衰减链"。'
        '批发商至制造商的恐慌占比均超过87%，说明情绪传染机制确实使系统进入了'
        '持续轻度恐慌的稳态。')

    # ============================================================
    # 第3节 归因分析
    # ============================================================
    add_heading(doc, '3 归因分析', level=1)

    # ---- 3.1 Exp_1b 剥离实验归因分析 ----
    add_heading(doc, '3.1 Exp_1b 剥离实验归因分析', level=2)

    if attr_exp1b is not None:
        add_table_title(doc, '表6 Exp_1b 分销商情绪状态与订货决策的 Pearson 相关系数')
        r_val = attr_exp1b['pearson_emotion_order']['r']
        p_val = attr_exp1b['pearson_emotion_order']['p_value']
        p_str = f'{p_val:.2e}' if p_val > 0 else '≈ 0'
        make_three_line_table(doc,
            ['分析对象', 'r(E, q)', 'p 值', '显著性'],
            [['分销商 (k=3)', f'{r_val:.3f}', p_str, '***']],
            col_widths=[3.0, 2.5, 3.5, 2.0])

        add_table_title(doc, '表7 Exp_1b 分销商不同情绪状态下的订货统计')
        gs = attr_exp1b['group_stats']
        rows_e1b = []
        emo_labels = [('恐慌', '恐慌(E<-0.3)'), ('焦虑', '焦虑(-0.3≤E<-0.1)'),
                      ('中性', '中性(|E|≤0.1)'), ('自信', '自信(0.1<E≤0.3)'),
                      ('乐观', '乐观(E>0.3)')]
        for key, label in emo_labels:
            g = gs[key]
            rows_e1b.append([label, str(g['count']), f'{g["ratio"]*100:.2f}%',
                f'{g["mean_order"]:.2f}', f'{g["std_order"]:.2f}',
                f'{g["overorder_ratio"]*100:.2f}%'])
        make_three_line_table(doc,
            ['情绪状态', '周期数', '占比', '平均订货量', '订货标准差', '过度订货比例'],
            rows_e1b, col_widths=[2.8, 1.5, 1.5, 2.0, 2.0, 2.2])

        add_paragraph(doc,
            'Exp_1b 分销商情绪与订货量的 Pearson 相关系数 r=-0.116'
            f'（p={p_str}），达到极显著水平，证明即使在无动态事件的平稳环境下，'
            '情绪状态仍统计显著地影响订货决策。恐慌状态下（95.9%周期）平均'
            '订货量为20.23，接近需求均值（D₀=20），而自信/乐观状态下（0.6%周期）'
            '平均订货量为36.31，远高于需求均值。这一反直觉现象表明，DQN 在'
            '长期训练中已学到"恐慌反馈下维持理性订货"的策略，情绪对决策的'
            '扰动被 DQN 的学习能力部分抵消，解释了 Exp_1b 中情绪机制对 BWE '
            '独立效应极小（-1.3%）的原因。')
    else:
        add_paragraph(doc, 'Exp_1b 归因分析数据未找到，跳过此节。')

    # ---- 3.2 Exp_2 情绪-决策相关性 ----
    add_heading(doc, '3.2 Exp_2 情绪-决策相关性分析', level=2)

    add_table_title(doc, '表8 Exp_2 各节点情绪状态与订货决策的 Pearson 相关系数')
    headers6 = ['节点', 'r(E, q)', 'p 值', 'r(E, q_excess)', 'p 值', '显著性']
    rows6 = []
    for k in range(1, 5):
        c = corr[str(k)]
        p_order = c['p_value_order']
        p_excess = c['p_value_excess']
        # p 值极小时显示为科学计数或近似 0
        p_order_str = (f'{p_order:.2e}' if p_order > 0 else '≈ 0')
        p_excess_str = (f'{p_excess:.2e}' if p_excess > 0 else '≈ 0')
        rows6.append([
            NODE_NAMES[k - 1],
            f'{c["pearson_emotion_order"]:.3f}', p_order_str,
            f'{c["pearson_emotion_excess"]:.3f}', p_excess_str,
            '***',
        ])
    make_three_line_table(doc, headers6, rows6,
                          col_widths=[1.8, 2.0, 2.6, 2.4, 2.6, 1.6])

    add_table_title(doc, '表9 Exp_2 不同情绪状态下的平均订货量')
    headers7 = ['节点', '恐慌(E<-0.3)', '中性(|E|≤0.1)', '乐观(E>0.3)', '恐慌-中性差']
    rows7 = []
    for k in range(1, 5):
        c = corr[str(k)]
        panic_str = f'{c["panic_mean_order"]:.2f} (n={c["panic_count"]})'
        neutral_str = f'{c["neutral_mean_order"]:.2f} (n={c["neutral_count"]})'
        opt_str = f'{c["optimistic_mean_order"]:.2f} (n={c["optimistic_count"]})'
        diff = c['panic_mean_order'] - c['neutral_mean_order']
        rows7.append([
            NODE_NAMES[k - 1],
            panic_str, neutral_str, opt_str,
            f'{diff:+.2f}',
        ])
    make_three_line_table(doc, headers7, rows7,
                          col_widths=[1.6, 3.2, 3.2, 3.2, 2.2])

    if fig_paths.get('fig3'):
        add_figure(doc, fig_paths['fig3'], '图3 情绪-决策相关性散点图')

    add_paragraph(doc,
        '（1）零售商、批发商、分销商的情绪-订货量呈正相关（r>0），即乐观'
        '状态下订货量更高，恐慌状态下订货量更低。这一现象的微观机制为：'
        '恐慌情绪主要由需求低谷期的库存积压触发（积压反馈权重 w_e=0.3），'
        '而非单纯由缺货触发。当需求处于低谷时，基础订货量 q_t 本就偏低，'
        '即使恐慌放大系数（1.0+0.5|E_t|）生效，最终订货量仍低于需求高峰期'
        '的中性状态。', indent=False)

    add_paragraph(doc,
        '（2）制造商呈负相关（r=-0.213），即恐慌时订货量更高。这是因为'
        '制造商距离终端需求最远，其情绪主要由上游传染驱动而非直接需求冲击，'
        '恐慌时传染冲击叠加放大系数导致订货量上升。', indent=False)

    add_paragraph(doc,
        '（3）所有节点的相关性均达到 p<0.001 的极显著水平，证明情绪状态'
        '确实显著影响订货决策，H1 的核心假设得到统计验证。', indent=False)

    # ---- 3.3 情绪传染路径 ----
    add_heading(doc, '3.3 情绪传染路径分析', level=2)

    add_table_title(doc, '表10 情绪传染事件统计')
    headers8 = ['传染路径', '事件次数', '占比', '平均延迟']
    rows8 = []
    path_order = ['零售商→批发商', '批发商→分销商', '分销商→制造商']
    total_ev = ca['total_events']
    for path in path_order:
        cnt = ca['path_counts'].get(path, 0)
        rows8.append([path, str(cnt), f'{cnt / total_ev * 100:.1f}%', '1-5周期'])
    rows8.append(['总计', str(total_ev), '100%', '—'])
    make_three_line_table(doc, headers8, rows8,
                          col_widths=[4.0, 2.5, 2.5, 3.0])

    if fig_paths.get('fig4'):
        add_figure(doc, fig_paths['fig4'], '图5 情绪传染网络图')

    add_paragraph(doc,
        '情绪传染事件共检测到2830次，呈现显著的逐级衰减特征。'
        '零售商→批发商的传染占比最高（65.0%），符合"恐慌从终端向上游蔓延"'
        '的理论预期。传染强度沿供应链递减（1839→594→397），验证了情绪传染'
        '机制中衰减系数的设计合理性。')

    # ---- 3.4 正向激励阻断效应 ----
    add_heading(doc, '3.4 正向激励阻断效应分析', level=2)

    add_table_title(doc, '表11 恐慌 vs 中性状态下的过度订货比例')
    headers9 = ['节点', '恐慌时过度订货比例', '中性时过度订货比例', '阻断效应']
    rows9 = []
    for k in range(1, 5):
        pr = over['panic'][str(k)] * 100
        nr = over['neutral'][str(k)] * 100
        block = nr - pr
        sign = '+' if block >= 0 else ''
        rows9.append([
            NODE_NAMES[k - 1],
            f'{pr:.2f}%', f'{nr:.2f}%',
            f'{sign}{block:.2f}pp',
        ])
    make_three_line_table(doc, headers9, rows9,
                          col_widths=[2.0, 3.5, 3.5, 2.5])

    if fig_paths.get('fig5'):
        add_figure(doc, fig_paths['fig5'], '图6 正向激励阻断效应图')

    add_paragraph(doc,
        '在下游三个节点（零售商、批发商、分销商）中，恐慌状态下的过度订货'
        '比例均低于中性状态，证明正向激励机制在恐慌时成功"阻断"了过度订货'
        '行为。阻断效应在批发商节点最为显著（86.55%→44.20%，降低42.35个'
        '百分点）。制造商出现反向效应（恐慌时过度订货更多），与其负相关系数'
        '一致，反映了上游传染驱动的恐慌模式不同于下游的需求驱动恐慌。')

    # ============================================================
    # 第4节 持续学习鲁棒性分析
    # ============================================================
    add_heading(doc, '4 持续学习鲁棒性分析', level=1)

    add_heading(doc, '4.1 实验设计', level=2)

    add_paragraph(doc,
        '基于弹性权重巩固（EWC）与优先经验回放（PER）机制，设计任务切换场景'
        '验证灾难性遗忘的抑制效果：')

    add_paragraph(doc,
        'Task1（平稳需求）：ρ=0.5, σ_ε=5，训练15000步；', indent=False)
    add_paragraph(doc,
        'Task2（反向剧烈波动）：ρ=-0.5, σ_ε=20，训练3000步；', indent=False)
    add_paragraph(doc,
        '三组对比：A（无 EWC 无噪声）、B（有 EWC 无噪声）、'
        'C（有 EWC + 情绪感知噪声 σ_noise=0.15）。', indent=False)

    add_heading(doc, '4.2 实验结果', level=2)

    add_table_title(doc, '表12 持续学习任务切换前后 Task1 性能对比')
    headers10 = ['指标', '组别', 'Task1训练后', 'Task2训练后', '变化量', '遗忘率']
    rows10 = []
    # 分销商 BWE
    metric_groups = [
        ('分销商BWE', 'bwe_distributor', False),
        ('服务水平SL', 'service_level', True),
        ('平均奖励', 'avg_reward', False),
    ]
    group_labels = [
        ('A_无EWC无噪声', 'A'),
        ('B_有EWC无噪声', 'B'),
        ('C_有EWC有噪声', 'C'),
    ]
    for metric_name, metric_key, is_sl in metric_groups:
        for gkey, glabel in group_labels:
            d = summ[metric_key][gkey]
            before = d['before']
            after = d['after']
            change = d['change']
            if is_sl:
                before_str = f'{before * 100:.2f}%'
                after_str = f'{after * 100:.2f}%'
                change_str = f'{change * 100:+.2f}pp'
                forget_str = '—'
            else:
                before_str = f'{before:.3f}'
                after_str = f'{after:.3f}'
                change_str = f'{change:+.3f}'
                if metric_name == '分销商BWE':
                    fr = forget['forgetting'][gkey]['bwe_forgetting_rate']
                    forget_str = f'{fr * 100:+.2f}%'
                else:
                    forget_str = '—'
            rows10.append([metric_name, glabel, before_str, after_str,
                           change_str, forget_str])
    make_three_line_table(doc, headers10, rows10,
                          col_widths=[2.2, 1.2, 2.4, 2.4, 2.0, 2.2])

    if fig_paths.get('fig7'):
        add_figure(doc, fig_paths['fig7'], '图7 持续学习性能曲线')

    add_heading(doc, '4.3 情绪感知误差统计', level=2)

    add_table_title(doc, '表13 C 组情绪感知误差统计（σ_noise=0.15）')
    headers11 = ['统计量', '数值']
    rows11 = [
        ['误差均值', f'{perc["error_mean"]:.4f}'],
        ['误差标准差', f'{perc["error_std"]:.4f}'],
        ['平均绝对误差 (MAE)', f'{perc["error_mae"]:.4f}'],
        ['样本数', str(perc['n_samples'])],
    ]
    make_three_line_table(doc, headers11, rows11,
                          col_widths=[5.0, 4.0])

    if fig_paths.get('fig8'):
        add_figure(doc, fig_paths['fig8'], '图8 噪声鲁棒性分析')

    add_heading(doc, '4.4 关键发现', level=2)

    add_paragraph(doc,
        '（1）EWC 在无噪声条件下的保护效果：B 组（有 EWC）相较 A 组（无 EWC），'
        'Task1 服务水平保持更好（+1.42pp vs +0.10pp），平均奖励提升更显著'
        '（+0.057 vs +0.004），证明 EWC 正则化有效约束了 Q 网络参数偏移，'
        '保护了旧任务知识。', indent=False)

    add_paragraph(doc,
        '（2）情绪感知噪声的负面影响：C 组（有 EWC + 噪声）的服务水平下降1.26pp，'
        '平均奖励下降0.047，表明情绪感知噪声（MAE=0.1126）会削弱 EWC 的保护效果。'
        '这一发现揭示了感知精度对持续学习鲁棒性的关键影响，为实际部署中的'
        '传感器校准提供了指导。', indent=False)

    add_paragraph(doc,
        '（3）BWE 的非典型变化：三组实验中 BWE 在 Task2 训练后均下降（约-53%），'
        '未出现预期的"灾难性遗忘"现象。这表明 Task2 的反向需求模式（ρ=-0.5）'
        '与 Task1 的正向需求模式（ρ=0.5）形成的对比训练效应，可能增强了 Q 网络'
        '对需求分布漂移的泛化能力，而非导致遗忘。', indent=False)

    # ============================================================
    # 第5节 假设验证结论
    # ============================================================
    add_heading(doc, '5 假设验证结论', level=1)

    add_heading(doc, '5.1 H1（损失厌恶放大假设）—— 部分验证', level=2)

    add_paragraph(doc, '验证依据：', bold=True, indent=False)
    add_paragraph(doc,
        '剥离实验 Exp_1b 证实情绪演化方程确实将决策者情绪推向恐慌饱和'
        '（分销商 E=-0.778，95.9%恐慌占比），且情绪-订货量 Pearson '
        '相关系数达极显著水平（r=-0.116，p<10⁻⁶⁰）；Exp_2 情绪波动指数'
        '呈下游放大梯度（零售商 σ_E=0.59 > 制造商0.22），证明缺货事件'
        '确实通过情绪演化方程将零售商情绪推向恐慌饱和；Exp_2 各节点'
        '情绪-订货量 Pearson 相关系数全部达到 p<10⁻⁴⁸ 的极显著水平，'
        '统计上确认情绪显著影响决策；Exp_2 零售商 SL 下降16.16%'
        '（98.92%→82.76%），证明恐慌放大效应在动态扰动下确实导致决策'
        '偏离理性最优。')

    add_paragraph(doc, '未完全验证的部分：', bold=True, indent=False)
    add_paragraph(doc,
        '剥离实验 Exp_1b 揭示，情绪机制对 BWE 的独立效应极小（分销商'
        '-1.3%），说明 DQN 的学习能力可在平稳环境下抵消情绪扰动。情绪'
        '放大效应需在动态突发事件触发下才显著显现，而非在所有条件下持续'
        '放大。此外，恐慌状态下的订货量反而低于中性状态（零售商16.08 '
        'vs 23.30），这与"恐慌导致过度订货"的直接预期不符，深层原因是'
        '恐慌情绪主要由需求低谷期的积压反馈触发。')

    add_heading(doc, '5.2 H2（正向激励阻断假设）—— 验证成立', level=2)

    add_paragraph(doc, '验证依据：', bold=True, indent=False)
    add_paragraph(doc,
        '分销商 BWE 从 Baseline 的67.33降至 Exp_2 的9.81（-85.43%），'
        '超过 H2 假设的"80%以上"阈值；剥离实验进一步揭示，Exp_1b（启用'
        '情绪+正向激励）分销商 BWE 为10.51，与 Exp_1（10.65）接近，说明'
        '正向激励机制在无动态事件时已有效引导 DQN 趋向精准匹配；恐慌时'
        '过度订货比例在下游三节点均低于中性状态（零售商41.60%<58.14%、'
        '批发商44.20%<86.55%、分销商48.25%<67.44%），证明正向激励成功'
        '阻断了"恐慌→过度订货"的恶性循环；制造商 BWE 从 Exp_1 的22.70'
        '降至 Exp_2 的10.07（-55.66%），证明阻断效应沿供应链传导。')

    add_heading(doc, '5.3 H3（协同鲁棒性假设）—— 部分验证', level=2)

    add_paragraph(doc, '验证依据：', bold=True, indent=False)
    add_paragraph(doc,
        '剥离实验精确归因显示，协同通信机制贡献了制造商 BWE 降幅的绝大'
        '部分（Exp_1b→Exp_2: -55.8%），远超情绪机制的独立贡献（+0.4%），'
        '超过 H3 假设的"45%以上"阈值；EWC 在无噪声条件下有效保护旧任务'
        '知识（SL +1.42pp，奖励 +0.057）；情绪感知噪声下系统仍保持基本'
        '功能（BWE 遗忘率-53.77%，与无噪声组接近）。')

    add_paragraph(doc, '未完全验证的部分：', bold=True, indent=False)
    add_paragraph(doc,
        '系统平均 SL 为94.73%，低于 H3 假设的"99%以上"阈值。剥离实验'
        '精确归因表明，SL 下降并非协同机制本身导致——Exp_1b 系统 SL'
        '（99.61%）与 Exp_1（99.62%）几乎一致，而 Exp_2 的 SL 下降主要'
        '由76次动态突发事件与情绪放大效应叠加导致（零售商 SL 从98.96%降至'
        '82.76%）。这一发现揭示了协同机制在 BWE 控制上的有效性，同时指出了'
        '动态事件冲击下终端节点 SL 保护的工程优化方向。')

    # ============================================================
    # 第6节 工程优化建议
    # ============================================================
    add_heading(doc, '6 工程优化建议', level=1)

    add_paragraph(doc, '基于实验结果的分析，提出以下优化方向：')

    add_heading(doc, '6.1 差异化情绪调节参数', level=2)

    add_paragraph(doc,
        '当前所有节点使用统一的恐慌放大系数（1.0+0.5|E_t|）。建议针对零售商'
        '降低该系数至 1.0+0.2|E_t|，以缓解终端节点的过度订货压力，'
        '预期可将零售商 SL 从82.76%提升至95%以上。')

    add_heading(doc, '6.2 情绪反馈信号解耦', level=2)

    add_paragraph(doc,
        '当前情绪演化方程中，积压反馈（w_e=0.3）与缺货反馈（w_s=1.0）共同'
        '驱动情绪。建议将积压反馈的符号反转——积压应触发"乐观"（库存充足）'
        '而非"焦虑"，使情绪演化更符合实际管理心理。')

    add_heading(doc, '6.3 感知噪声自适应校准', level=2)

    add_paragraph(doc,
        '持续学习实验显示，情绪感知噪声（MAE=0.1126）会削弱 EWC 保护效果。'
        '建议在 Q 网络输入端引入感知噪声自适应层，通过元学习动态校准感知误差，'
        '提升实际部署鲁棒性。')

    # ============================================================
    # 参考文献
    # ============================================================
    add_heading(doc, '参考文献', level=1)

    refs = [
        '[1] Sterman J D. Modeling managerial behavior: Misperceptions of '
        'feedback in a dynamic decision making experiment[J]. Management '
        'Science, 1989, 35(3): 321-339.',
        '[2] Chen F, Drezner Z, Ryan J K, et al. Quantifying the bullwhip '
        'effect in a simple supply chain[J]. Management Science, 2000, '
        '46(3): 436-443.',
        '[3] Kahneman D, Tversky A. Prospect theory: An analysis of '
        'decision under risk[J]. Econometrica, 1979, 47(2): 263-291.',
        '[4] Lee H L, Padmanabhan V, Whang S. Information distortion in a '
        'supply chain: The bullwhip effect[J]. Management Science, 1997, '
        '43(4): 546-558.',
        '[5] 李勇, 陈元, 于辉. 缓解牛鞭效应的新途径：人机协同的智慧决策'
        '机器人[J]. 中国管理科学, 2022.',
        '[6] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control '
        'through deep reinforcement learning[J]. Nature, 2015, 518(7540): '
        '529-533.',
        '[7] Kirkpatrick J, Pascanu R, Rabinowitz N, et al. Overcoming '
        'catastrophic forgetting in neural networks[J]. Proceedings of the '
        'National Academy of Sciences, 2017, 114(13): 3521-3526.',
        '[8] Schaul T, Quan J, Antonoglou I, et al. Prioritized experience '
        'replay[C]//ICLR, 2016.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.first_line_indent = Pt(-21)
        run = p.add_run(ref)
        set_run_font(run, size=10.5)

    # ---- 保存 ----
    output_path = os.path.join(BASE_DIR,
                               'Exp_2综合分析_人智协同智慧决策.docx')
    doc.save(output_path)
    print(f"\n[OK] docx 文档已生成: {output_path}")
    return output_path


# ============================================================
# 主程序
# ============================================================

def main():
    print('=' * 70)
    print('生成《中国管理科学》规范的 Exp_2 综合分析 docx 文档')
    print('=' * 70)

    # 1. 加载数据
    print('\n[1/3] 加载实验数据...')
    four, attr, attr_exp1b, forget, ts = load_all_data()
    print(f"  四组对比: baseline BWE={four['baseline']['bwe']}")
    print(f"           exp1     BWE={four['exp1']['bwe']}")
    print(f"           exp1b    BWE={four['exp1b']['bwe']}")
    print(f"           exp2     BWE={four['exp2']['bwe']}")
    print(f"  归因分析: 传染事件={attr['contagion_analysis']['total_events']}")
    if attr_exp1b:
        print(f"  Exp_1b归因: r={attr_exp1b['pearson_emotion_order']['r']:.3f}")
    print(f"  持续学习: 组别={list(forget['summary']['bwe_distributor'].keys())}")
    print(f"  时序数据: {'已加载' if ts else '未找到'}")

    # 2. 生成图表
    print('\n[2/3] 生成 PNG 图表...')
    fig_paths = generate_all_figures(four, attr, forget, ts)

    # 3. 生成 docx
    print('\n[3/3] 生成 docx 文档...')
    output = generate_docx(four, attr, attr_exp1b, forget, fig_paths)

    print(f"\n{'=' * 70}")
    print(f'完成！文档已保存: {output}')
    print(f"{'=' * 70}")


if __name__ == '__main__':
    main()
