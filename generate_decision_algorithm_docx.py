"""
理性决策与智慧决策核心算法及公式 docx 生成脚本
================================================
功能：
1. 复现智慧决策（IDMR）与理性决策（SMA+OUT）的核心算法
2. 标准OMML公式格式（LaTeX→MathML→OMML）
3. 三线表格式
4. 用于论文的理性决策与智慧决策研究段落
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# OMML 公式转换
# ============================================================
try:
    from latex2mathml.converter import convert as _latex_to_mathml
    from lxml import etree
    _XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
    _xslt_tree = etree.parse(_XSLT_PATH)
    _xslt_transform = etree.XSLT(_xslt_tree)
    _OMML_AVAILABLE = True
except Exception as e:
    _OMML_AVAILABLE = False
    print(f"[警告] OMML 不可用: {e}")


def latex_to_omml(latex_str, display='inline'):
    """LaTeX -> MathML -> OMML"""
    mathml_str = _latex_to_mathml(latex_str, display=display)
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = _xslt_transform(mathml_tree)
    return omml_tree.getroot()


# ============================================================
# 字体与样式
# ============================================================
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_HEADING_CN = '黑体'
COLOR_HEADING = RGBColor(0x1F, 0x3A, 0x5F)
COLOR_FORMULA = RGBColor(0x00, 0x00, 0x80)
COLOR_FIGURE = RGBColor(0x70, 0x70, 0x70)


def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=10, bold=False, color=None):
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def set_three_line_table(table):
    """三线表格式：顶线1.5pt、表头下线0.75pt、底线1.5pt，无竖线"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')

    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '0'); top.set(qn('w:color'), '000000')
    tblBorders.append(top)

    left = OxmlElement('w:left'); left.set(qn('w:val'), 'nil'); tblBorders.append(left)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0'); bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)

    right = OxmlElement('w:right'); right.set(qn('w:val'), 'nil'); tblBorders.append(right)
    insideH = OxmlElement('w:insideH'); insideH.set(qn('w:val'), 'nil'); tblBorders.append(insideH)
    insideV = OxmlElement('w:insideV'); insideV.set(qn('w:val'), 'nil'); tblBorders.append(insideV)
    tblPr.append(tblBorders)

    # 表头单元格底边框
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '6')
        bottom_border.set(qn('w:space'), '0')
        bottom_border.set(qn('w:color'), '000000')
        existing_bottom = tcBorders.find(qn('w:bottom'))
        if existing_bottom is not None:
            tcBorders.remove(existing_bottom)
        tcBorders.append(bottom_border)


def set_cell_content(cell, text, bold=False, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格内容（支持行内公式 $...$）"""
    import re
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pattern = r'(\$[^$]+\$)'
    segments = re.split(pattern, text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith('$') and seg.endswith('$'):
            formula = seg[1:-1]
            if _OMML_AVAILABLE:
                try:
                    omml = latex_to_omml(formula)
                    p._element.append(omml)
                    continue
                except Exception as e:
                    print(f"[提示] 公式转换失败: {formula} ({e})")
            run = p.add_run(formula)
            set_run_font(run, size=size, bold=True, color=COLOR_FORMULA)
            run.font.italic = True
        else:
            run = p.add_run(seg)
            set_run_font(run, size=size, bold=bold)


def set_cell_vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def add_heading(doc, text, level=2, size=14):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = FONT_HEADING_CN
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_HEADING_CN)
    return h


def add_paragraph(doc, text, indent=True, size=11):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    # 解析行内公式 $...$
    import re
    pattern = r'(\$[^$]+\$)'
    segments = re.split(pattern, text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith('$') and seg.endswith('$'):
            formula = seg[1:-1]
            if _OMML_AVAILABLE:
                try:
                    omml = latex_to_omml(formula)
                    p._element.append(omml)
                    continue
                except Exception as e:
                    print(f"[提示] 行内公式转换失败: {formula} ({e})")
            run = p.add_run(formula)
            set_run_font(run, size=size, bold=True, color=COLOR_FORMULA)
            run.font.italic = True
        else:
            run = p.add_run(seg)
            set_run_font(run, size=size)
    return p


def add_table_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(caption_text)
    set_run_font(run, size=10, bold=True)


def add_block_formula(doc, latex_str, formula_num=None):
    """添加块级公式（居中，可带编号）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if _OMML_AVAILABLE:
        try:
            omml = latex_to_omml(latex_str, display='block')
            p._element.append(omml)
        except Exception as e:
            print(f"[提示] 块级公式转换失败: {latex_str} ({e})")
            run = p.add_run(latex_str)
            set_run_font(run, size=12, bold=True, color=COLOR_FORMULA)
            run.font.italic = True
    else:
        run = p.add_run(latex_str)
        set_run_font(run, size=12, bold=True, color=COLOR_FORMULA)
        run.font.italic = True
    # 公式编号
    if formula_num:
        run = p.add_run(f'    ({formula_num})')
        set_run_font(run, size=11, bold=False)


def create_table(doc, data, col_widths, first_col_left=False):
    n_rows = len(data)
    n_cols = len(data[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            set_cell_vertical_center(cell)
            if i == 0:
                set_cell_content(cell, cell_text, bold=True, size=9)
            else:
                if first_col_left and j == 0:
                    set_cell_content(cell, cell_text, bold=False, size=9,
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
                else:
                    set_cell_content(cell, cell_text, bold=False, size=9)
    set_three_line_table(table)
    return table


# ============================================================
# 创建文档
# ============================================================
def create_decision_algorithm_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)

    # ============================================================
    # 标题
    # ============================================================
    heading = doc.add_heading(level=1)
    run = heading.add_run('理性决策与智慧决策核心算法及公式')
    run.font.name = FONT_HEADING_CN
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_HEADING_CN)

    add_paragraph(doc,
        '本部分系统复现理性决策（SMA+OUT）与智慧决策（IDMR-DQN）的核心算法及公式，'
        '为论文的理性决策与智慧决策研究段落提供标准公式依据。'
        '所有公式采用LaTeX→MathML→OMML转换链生成Word原生公式格式，确保可编辑性与排版规范性。')

    # ============================================================
    # 1. 供应链需求模型
    # ============================================================
    add_heading(doc, '1 供应链需求模型', level=2, size=14)

    add_paragraph(doc,
        '顾客需求遵循一阶自回归过程AR(1)，该模型既反映了真实需求的自相关性，'
        '又为动态突发事件提供了可控的扰动基底：')

    add_block_formula(doc, r'D_t = d + \rho \cdot D_{t-1} + \varepsilon_t', formula_num='1')

    add_paragraph(doc,
        '其中 $d$ 为需求基础水平（$d=10$），$\\rho$ 为自相关系数（$\\rho=0.5$，$|\\rho|<1$ 保证平稳性），'
        '$\\varepsilon_t \\sim N(0, \\sigma_\\varepsilon^2)$ 为独立同正态分布的误差项（$\\sigma_\\varepsilon=5$），'
        '初始需求 $D_0 = d/(1-\\rho) = 20$。')

    # ============================================================
    # 2. 理性决策基线（SMA+OUT）
    # ============================================================
    add_heading(doc, '2 理性决策基线（SMA+OUT）', level=2, size=14)

    add_paragraph(doc,
        '零售商、批发商与制造商（k=1,2,4）采用Chen等(2000)的理性决策基线，'
        '包含SMA移动平均预测与OUT订至点策略两个核心环节。')

    add_heading(doc, '2.1 SMA移动平均预测', level=3, size=12)

    add_paragraph(doc,
        '根据最近 $p$ 个周期的需求预测当期需求：')

    add_block_formula(doc, r'\hat{D}_t^1 = \frac{1}{p} \sum_{i=1}^{p} D_{t-i}', formula_num='2')

    add_paragraph(doc,
        'L提前期内的预测需求为单步预测的线性扩展：')

    add_block_formula(doc, r'\hat{D}_t^L = L \cdot \hat{D}_t^1', formula_num='3')

    add_heading(doc, '2.2 预测误差估计', level=3, size=12)

    add_paragraph(doc,
        '基于历史预测误差估计L期预测的标准差：')

    add_block_formula(doc, r'\hat{\sigma}_{e_t}^L = C_{L,\rho} \cdot \frac{\sigma_\varepsilon}{p} \cdot \sqrt{\sum_{i=1}^{p} e_{t-i}^2}', formula_num='4')

    add_paragraph(doc,
        '其中 $e_t = D_t - \\hat{D}_t^1$ 为单步预测误差，$C_{L,\\rho}$ 为预测误差缩放常数（$C_{L,\\rho}=2.0$）。')

    add_heading(doc, '2.3 OUT订至点策略', level=3, size=12)

    add_paragraph(doc,
        '期望库存（订至点）由预测需求与安全库存两部分构成：')

    add_block_formula(doc, r'S_t = \hat{D}_t^L + z \cdot \hat{\sigma}_{e_t}^L', formula_num='5')

    add_paragraph(doc,
        '其中 $z$ 为安全库存系数（$z=2$ 对应97.7%服务水平）。'
        '订货决策为订至点与当前可用库存的差值：')

    add_block_formula(doc, r'q_t = \max\left(0, \, S_t - (NS_t + WIP_t)\right)', formula_num='6')

    add_paragraph(doc,
        '其中 $NS_t$ 为净库存，$WIP_t$ 为在途库存，$\\max(\\cdot, 0)$ 保证不允许负订货。')

    # 理性决策参数表
    add_table_caption(doc, '表1  理性决策（SMA+OUT）参数配置')
    table1_data = [
        ['参数', '符号', '取值', '说明'],
        ['需求基础水平', '$d$', '10', 'AR(1)常数项'],
        ['自相关系数', '$\\rho$', '0.5', '保证平稳性'],
        ['误差标准差', '$\\sigma_\\varepsilon$', '5', '正态分布'],
        ['运输延迟', '$L$', '2', '订货后到货周期数'],
        ['SMA预测窗口', '$p$', '5', '移动平均周期'],
        ['安全库存系数', '$z$', '2', '对应97.7% SL'],
        ['预测误差常数', '$C_{L,\\rho}$', '2.0', '误差缩放'],
        ['初始库存', '$NS_0$', '10.0', '各节点统一'],
    ]
    create_table(doc, table1_data,
                col_widths=[Cm(3.0), Cm(2.5), Cm(2.0), Cm(5.5)],
                first_col_left=True)

    # ============================================================
    # 3. IDMR智慧决策机器人（DQN）
    # ============================================================
    add_heading(doc, '3 IDMR智慧决策机器人（DQN）', level=2, size=14)

    add_paragraph(doc,
        'IDMR（Intelligent Decision-making Robot）智慧决策机器人部署于分销商节点（k=3），'
        '基于深度Q网络（DQN）进行决策。其核心架构包含状态空间、动作空间、奖励函数、'
        'Q网络与训练算法五个模块。')

    add_heading(doc, '3.1 状态空间', level=3, size=12)

    add_paragraph(doc,
        'IDMR观察5维状态向量，涵盖库存水平、在途库存、下游订单与运输信息：')

    add_block_formula(doc, r's_t = \left(S_{t-1}^3, \, WIP_{t-1}^3, \, q_{t-1}^2, \, Trans_{t-2}^3, \, q_{t-1}^3\right)', formula_num='7')

    add_table_caption(doc, '表2  IDMR状态空间分量定义')
    table2_data = [
        ['分量', '含义', '归一化'],
        ['$S_{t-1}^3$', '分销商上期净库存', '/100'],
        ['$WIP_{t-1}^3$', '分销商上期在途库存', '/100'],
        ['$q_{t-1}^2$', '批发商上期订单', '/50'],
        ['$Trans_{t-2}^3$', '分销商前期运输货物', '/50'],
        ['$q_{t-1}^3$', '分销商上期订货量', '/50'],
    ]
    create_table(doc, table2_data,
                col_widths=[Cm(3.0), Cm(6.0), Cm(3.0)],
                first_col_left=True)

    add_heading(doc, '3.2 动作空间', level=3, size=12)

    add_paragraph(doc,
        'IDMR的动作空间为30个离散订货量：')

    add_block_formula(doc, r'a_t \in \{11, 12, 13, \ldots, 40\}', formula_num='8')

    add_paragraph(doc,
        '该范围依据AR(1)需求均值 $d/(1-\\rho)=20$ 对称设计，覆盖合理订货区间。')

    add_heading(doc, '3.3 奖励函数', level=3, size=12)

    add_paragraph(doc,
        '采用论文公式(11)的稠密化版本，以fill_rate为主导奖励信号：')

    add_block_formula(doc, r'r_t = \text{fill\_rate}_t - \text{holding\_penalty}_t', formula_num='9')

    add_block_formula(doc, r'\text{fill\_rate}_t = \frac{F_t}{D_t}, \quad \text{holding\_penalty}_t = 0.0001 \cdot \max(0, NS_t)', formula_num='10')

    add_paragraph(doc,
        '其中 $F_t$ 为实际满足量，$D_t$ 为需求量。'
        'fill_rate主导奖励信号直接对应服务水平目标，'
        '极轻微库存惩罚（系数0.0001）防止无界囤积但不压制订货积极性。')

    add_heading(doc, '3.4 Q网络结构', level=3, size=12)

    add_paragraph(doc,
        'Q网络采用三层全连接神经网络，隐藏层使用ReLU激活函数：')

    add_block_formula(doc, r'Q(s, a; \theta) \approx f_{\theta}\left(\text{ReLU}\left(W_2 \cdot \text{ReLU}(W_1 s + b_1) + b_2\right) W_3 + b_3\right)', formula_num='11')

    add_paragraph(doc,
        '网络结构为 $5 \\to 64 \\to 64 \\to 30$，其中输入层5维对应状态空间，'
        '输出层30维对应动作空间Q值。前向传播为：')

    add_block_formula(doc, r'z_1 = W_1 s + b_1, \quad a_1 = \text{ReLU}(z_1)', formula_num='12')
    add_block_formula(doc, r'z_2 = W_2 a_1 + b_2, \quad a_2 = \text{ReLU}(z_2)', formula_num='13')
    add_block_formula(doc, r'Q(s, a; \theta) = W_3 a_2 + b_3', formula_num='14')

    add_heading(doc, '3.5 DQN训练算法', level=3, size=12)

    add_paragraph(doc,
        'DQN的目标函数（损失函数）为均方TD误差：')

    add_block_formula(doc, r"L(\theta) = \mathbb{E}_{(s,a,r,s') \sim U(D)} \left[\left(y - Q(s, a; \theta)\right)^2\right]", formula_num='15')

    add_paragraph(doc,
        '其中目标Q值 $y$ 由目标网络计算：')

    add_block_formula(doc, r"y = r + \gamma \cdot \max_{a'} Q(s', a'; \theta^-)", formula_num='16')

    add_paragraph(doc,
        '其中 $\\gamma$ 为折扣因子（$\\gamma=0.9$），$\\theta^-$ 为目标网络参数，每 $N$ 步同步一次。'
        '反向传播采用Adam优化器，参数更新规则为：')

    add_block_formula(doc, r'm_t = \beta_1 m_{t-1} + (1-\beta_1) g_t, \quad v_t = \beta_2 v_{t-1} + (1-\beta_2) g_t^2', formula_num='17')
    add_block_formula(doc, r'\theta_{t+1} = \theta_t - \eta \cdot \frac{\hat{m}_t}{\sqrt{\hat{v}_t} + \epsilon}', formula_num='18')

    add_paragraph(doc,
        '其中 $\\hat{m}_t = m_t / (1-\\beta_1^t)$，$\\hat{v}_t = v_t / (1-\\beta_2^t)$ 为偏差校正，'
        '$\\eta=10^{-3}$ 为学习率，$\\beta_1=0.9$，$\\beta_2=0.999$，$\\epsilon=10^{-8}$。')

    add_heading(doc, '3.6 ε-greedy探索策略', level=3, size=12)

    add_paragraph(doc,
        'IDMR采用ε-greedy策略平衡探索与利用：')

    add_block_formula(doc, r'a_t = \begin{cases} \arg\max_a Q(s_t, a; \theta), & \text{概率 } 1-\varepsilon \\ \text{random} \in [11, 40], & \text{概率 } \varepsilon \end{cases}', formula_num='19')

    add_paragraph(doc,
        'ε线性衰减策略：')

    add_block_formula(doc, r'\varepsilon_t = \max\left(0.01, \, 1.0 - (1.0 - 0.01) \cdot \frac{t}{T}\right)', formula_num='20')

    add_paragraph(doc,
        '其中 $T$ 为总训练步数（$T=40000$），ε从1.0线性衰减至0.01。')

    # DQN超参数表
    add_table_caption(doc, '表3  IDMR智慧决策（DQN）超参数配置')
    table3_data = [
        ['超参数', '符号', '取值', '说明'],
        ['学习率', '$\\eta$', '$1\\times 10^{-3}$', 'Adam优化器'],
        ['折扣因子', '$\\gamma$', '0.9', '未来回报折扣'],
        ['批大小', 'batch', '32', '经验回放采样'],
        ['经验池容量', '|D|', '20000', 'Replay Buffer'],
        ['回放起始', '$N_{min}$', '100', '最小经验数'],
        ['ε起始', '$\\varepsilon_0$', '1.0', '探索率初始'],
        ['ε终止', '$\\varepsilon_{min}$', '0.01', '探索率最小'],
        ['目标网络更新频率', '$N_{update}$', '10', '步'],
        ['网络结构', '—', '5-64-64-30', '输入-隐藏-输出'],
        ['激活函数', '—', 'ReLU', '隐藏层'],
        ['权重初始化', '—', 'He', '$\\sqrt{2/n}$'],
        ['训练步数', '$T$', '40000', 'DQN训练步数'],
    ]
    create_table(doc, table3_data,
                col_widths=[Cm(3.5), Cm(2.0), Cm(2.5), Cm(5.0)],
                first_col_left=True)

    # ============================================================
    # 4. 人机协同三大机制
    # ============================================================
    add_heading(doc, '4 人机协同三大机制', level=2, size=14)

    add_paragraph(doc,
        'IDMR的人机协同机制通过"传授经验-限制决策-惩罚机制"三个环节，'
        '将人类理性决策的经验知识嵌入DQN训练过程。')

    add_heading(doc, '4.1 传授经验', level=3, size=12)

    add_paragraph(doc,
        '决策"老师"（理性决策）将理论最优决策经验教授给IDMR。'
        'IDMR通过观察环境状态，在理性决策指导下做出回报最大的决策。'
        '该机制将SMA+OUT策略的预测结果作为IDMR的状态输入分量，'
        '使智能体能够学习理性决策的"按需订货"特征。')

    add_heading(doc, '4.2 限制决策（ε-greedy探索）', level=3, size=12)

    add_paragraph(doc,
        'IDMR以 $1-\\varepsilon$ 概率选择Q值最大的动作，以 $\\varepsilon$ 概率在 $[11, 40]$ 范围内随机探索。'
        '该机制通过公式(19)与公式(20)实现，确保智能体在利用已学知识与探索新策略之间取得平衡。')

    add_heading(doc, '4.3 惩罚机制', level=3, size=12)

    add_paragraph(doc,
        '当IDMR库存超过经典供应链对应节点平均库存的5倍时，'
        '决策"老师"禁止IDMR向上游订货（强制 $a_t = 0$）：')

    add_block_formula(doc, r'\text{if } NS_t > 5 \cdot NS_{\text{avg}} \text{ then } a_t = 0', formula_num='21')

    add_paragraph(doc,
        '其中 $NS_{\\text{avg}}$ 为经典供应链对应节点的平均库存水平。'
        '该机制防止智能体学习到"无界囤货"的次优策略，'
        '将人类决策者的库存管理经验以硬约束形式注入训练过程。')

    # ============================================================
    # 5. 评估指标
    # ============================================================
    add_heading(doc, '5 评估指标', level=2, size=14)

    add_paragraph(doc,
        '本研究从牛鞭效应（BWE）、平均成本、服务水平（SL）三个维度评估决策效果。')

    add_heading(doc, '5.1 牛鞭效应（BWE）', level=3, size=12)

    add_paragraph(doc,
        '牛鞭效应方差比衡量节点订单方差相对顾客需求方差的放大倍数：')

    add_block_formula(doc, r'\text{BWE}_k = \frac{\text{var}(q_k)}{\text{var}(D)}', formula_num='22')

    add_paragraph(doc,
        '其中 $q_k$ 为节点 $k$ 的订货量序列，$D$ 为顾客需求序列。BWE值越大，表明订单波动在逐级传递中被放大得越严重。')

    add_heading(doc, '5.2 平均成本', level=3, size=12)

    add_paragraph(doc,
        '平均成本由库存持有成本与缺货惩罚成本两部分构成：')

    add_block_formula(doc, r'\text{Cost}_k = \frac{1}{T}\sum_{t=1}^{T}\left[h \cdot \max(0, NS_{k,t}) + b \cdot \max(0, D_{k,t} - F_{k,t})\right]', formula_num='23')

    add_paragraph(doc,
        '其中 $h=1.0$ 为单位库存成本，$b=2.0$ 为单位缺货成本，'
        '$NS_{k,t}$ 为净库存，$F_{k,t}$ 为实际满足量，$T$ 为评估周期。')

    add_heading(doc, '5.3 服务水平（SL）', level=3, size=12)

    add_paragraph(doc,
        '服务水平衡量节点对下游需求的满足能力：')

    add_block_formula(doc, r'\text{SL}_k = \frac{1}{T}\sum_{t=1}^{T}\min\left(1, \frac{F_{k,t}}{D_{k,t}}\right)', formula_num='24')

    add_paragraph(doc,
        '其中 $F_{k,t}$ 为实际满足量，$D_{k,t}$ 为需求量。SL值越接近1，表明需求满足能力越强。')

    # 评估指标表
    add_table_caption(doc, '表4  评估指标与公式对照')
    table4_data = [
        ['指标', '公式', '含义'],
        ['BWE', '$\\text{BWE}_k = \\text{var}(q_k)/\\text{var}(D)$', '订单方差放大倍数'],
        ['Cost', '$\\text{Cost}_k = \\frac{1}{T}\\sum[h\\cdot NS + b\\cdot (D-F)]$', '库存+缺货成本'],
        ['SL', '$\\text{SL}_k = \\frac{1}{T}\\sum\\min(1, F/D)$', '需求满足率'],
    ]
    create_table(doc, table4_data,
                col_widths=[Cm(2.0), Cm(7.0), Cm(4.0)],
                first_col_left=True)

    # ============================================================
    # 6. 实验结果对比
    # ============================================================
    add_heading(doc, '6 理性决策与智慧决策实验结果对比', level=2, size=14)

    add_paragraph(doc,
        '基于上述核心算法，本研究在四级供应链仿真环境下进行三组对比实验。'
        '表5呈现了理性决策与智慧决策（IDMR）在方差比、平均成本与服务水平三大指标上的对比结果。')

    add_table_caption(doc, '表5  理性决策与智慧决策（IDMR）核心指标对比')
    table5_data = [
        ['节点', '决策方式', 'BWE(理性)', 'BWE(IDMR)', 'SL(理性)', 'SL(IDMR)'],
        ['零售商 (k=1)', '理性决策', '4.01', '4.01', '99.2%', '99.2%'],
        ['批发商 (k=2)', '理性决策', '15.45', '16.31', '99.8%', '99.7%'],
        ['分销商 (k=3)', '智慧决策', '62.10', '10.43', '99.8%', '90.1%'],
        ['制造商 (k=4)', '理性决策', '279.52', '15.57', '100%', '100%'],
    ]
    create_table(doc, table5_data,
                col_widths=[Cm(2.5), Cm(2.0), Cm(2.5), Cm(2.5), Cm(2.0), Cm(2.0)],
                first_col_left=True)

    add_paragraph(doc,
        '实验结果表明，IDMR智慧决策使分销商BWE从62.10降至10.43（降幅83.2%），'
        '制造商BWE从279.52降至15.57（降幅94.4%），牛鞭效应得到极大缓解。'
        '上游（制造商）受益最显著，符合Lee等(1997)提出的牛鞭效应上游放大特征。'
        'IDMR在分销商节点有效阻断了方差放大传播，制造商端BWE降幅最为显著。')

    add_paragraph(doc,
        '理性决策的"按需订货"特征表现为 $\\bar{q} \\approx \\bar{D}$（偏差率<1.2%），'
        '符合SMA+OUT策略"围绕预测需求补货"的设计；'
        '而IDMR智慧决策的"主动压缩"特征表现为 $\\bar{q} < \\bar{D}$（偏差率-22.6%），'
        '通过DQN端到端学习突破了传统"按需补货"框架，主动降低向上游的订货量，'
        '从源头切断牛鞭效应的传播路径。这是IDMR使分销商BWE降低83%、制造商BWE降低94%的根本原因。')

    # ============================================================
    # 7. 算法总结
    # ============================================================
    add_heading(doc, '7 算法总结', level=2, size=14)

    add_table_caption(doc, '表6  理性决策与智慧决策算法对照')
    table6_data = [
        ['维度', '理性决策（SMA+OUT）', '智慧决策（IDMR-DQN）'],
        ['理论基础', '理性预期理论', '深度强化学习'],
        ['决策依据', '需求预测+安全库存', 'Q值函数最大化'],
        ['核心公式', '$q_t = S_t - (NS_t + WIP_t)$', '$a_t = \\arg\\max_a Q(s_t, a; \\theta)$'],
        ['优化目标', '最小化预测误差', '最大化累计奖励'],
        ['订货特征', '$\\bar{q} \\approx \\bar{D}$（按需订货）', '$\\bar{q} < \\bar{D}$（主动压缩）'],
        ['信息需求', '完整需求历史', '局部状态观测'],
        ['人机协同', '无', '传授经验+限制决策+惩罚'],
        ['BWE控制', '被动追随需求波动', '主动阻断方差传播'],
    ]
    create_table(doc, table6_data,
                col_widths=[Cm(2.5), Cm(5.0), Cm(5.0)],
                first_col_left=True)

    add_paragraph(doc,
        '理性决策与智慧决策的本质差异在于：理性决策依赖"预测+安全库存"范式，'
        '订货量被动追随需求波动（$\\bar{q} \\approx \\bar{D}$）；'
        '而智慧决策IDMR通过值函数近似直接优化长期累计奖励，'
        '订货量主动压缩订单波动（$\\bar{q} < \\bar{D}$），'
        '从源头切断牛鞭效应的传播路径。人机协同三大机制将人类决策者的经验知识'
        '以"传授-限制-惩罚"的形式嵌入DQN训练过程，'
        '使智能体既具备深度强化学习的优化能力，又保留人类决策的合理性约束。')

    # ============================================================
    # 保存
    # ============================================================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        '理性决策与智慧决策核心算法及公式.docx'
    )

    # 处理文件占用
    def _try_save(doc, path):
        try:
            if os.path.exists(path):
                os.remove(path)
            doc.save(path)
            return True
        except PermissionError:
            return False

    import datetime
    if not _try_save(doc, output_path):
        v2_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            '理性决策与智慧决策核心算法及公式_v2.docx'
        )
        if not _try_save(doc, v2_path):
            ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = os.path.join(
                os.path.dirname(os.path.abspath(__file__)),
                f'理性决策与智慧决策核心算法及公式_{ts}.docx'
            )
            doc.save(output_path)
            print(f"[提示] 原文件与v2均被占用，保存为时间戳版本")
        else:
            output_path = v2_path
            print(f"[提示] 原文件被占用，保存为: {output_path}")

    print(f"[OK] 文档已生成: {output_path}")
    print(f"     大小: {os.path.getsize(output_path) / 1024:.1f} KB")

    # 统计OMML公式数量
    import zipfile
    with zipfile.ZipFile(output_path, 'r') as z:
        with z.open('word/document.xml') as f:
            content = f.read().decode('utf-8')
            omath_count = content.count('<m:oMath>')
            omathpara_count = content.count('<m:oMathPara>')
            print(f"     OMML 公式: {omath_count} 个 oMath + {omathpara_count} 个 oMathPara")


if __name__ == '__main__':
    create_decision_algorithm_docx()
