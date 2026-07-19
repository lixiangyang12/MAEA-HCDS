"""
Markdown 转 DOCX 转换脚本
=========================
将 实验设计方案.md 转换为 Word 文档，保留：
  - 标题层级 (Heading 1/2/3)
  - 段落 (含加粗/斜体)
  - 表格 (Markdown 表格)
  - 数学公式 (OMML 原生公式格式，LaTeX→MathML→OMML)
  - 代码块
  - 列表 (无序/有序)
  - 水平线
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# LaTeX → MathML → OMML 转换
try:
    from latex2mathml.converter import convert as _latex_to_mathml
    from lxml import etree
    _XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
    _xslt_tree = etree.parse(_XSLT_PATH)
    _xslt_transform = etree.XSLT(_xslt_tree)
    _OMML_AVAILABLE = True
except Exception as e:
    _OMML_AVAILABLE = False
    print(f"[警告] OMML 公式转换不可用，将回退为文本格式: {e}")


def latex_to_omml(latex_str, display='inline'):
    """将 LaTeX 公式转换为 OMML XML 元素 (Office 原生公式格式)

    Args:
        latex_str: LaTeX 公式字符串 (如 r'\\rho = 0.5')
        display: 'inline' (行内) 或 'block' (块级，生成 oMathPara 容器)

    Returns:
        lxml.etree._Element: OMML XML 元素，可直接追加到段落
    """
    # 1. LaTeX → MathML
    mathml_str = _latex_to_mathml(latex_str, display=display)
    # 2. MathML → OMML (通过微软 MML2OMML.XSL)
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = _xslt_transform(mathml_tree)
    return omml_tree.getroot()


# ============================================================
# 样式配置
# ============================================================
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_HEADING_CN = '黑体'
FONT_CODE = 'Consolas'

COLOR_HEADING = RGBColor(0x1F, 0x3A, 0x5F)   # 深蓝
COLOR_FORMULA = RGBColor(0x00, 0x00, 0x80)   # 深蓝
COLOR_CODE = RGBColor(0x33, 0x33, 0x33)     # 深灰
COLOR_TABLE_HEADER = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_HEADER_BG = '4472C4'             # 表头背景蓝


def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=11, bold=False, color=None):
    """设置run字体（中英文分别设置）"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def add_paragraph_with_runs(doc, text, style=None, alignment=None):
    """解析段落文本中的加粗/斜体/行内公式，添加到doc

    行内公式 $...$ 使用 OMML 原生公式格式（若可用），否则回退为文本。
    """
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment

    # 解析 **bold**, *italic*, `code`, $formula$
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=11, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            set_run_font(run, size=11, bold=False)
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_run_font(run, font_en=FONT_CODE, size=10, color=COLOR_CODE)
        elif part.startswith('$') and part.endswith('$'):
            formula = part[1:-1]
            _add_inline_formula(p, formula)
        else:
            run = p.add_run(part)
            set_run_font(run, size=11)
    return p


def _add_inline_formula(paragraph, latex_str):
    """将行内公式添加到段落（优先使用OMML，失败则回退为文本）"""
    if _OMML_AVAILABLE:
        try:
            omml = latex_to_omml(latex_str)
            paragraph._element.append(omml)
            return
        except Exception as e:
            print(f"[提示] 公式转换失败，回退为文本: {latex_str} ({e})")
    # Fallback: 文本格式
    run = paragraph.add_run(latex_str)
    set_run_font(run, font_en=FONT_EN, size=11, bold=True, color=COLOR_FORMULA)
    run.font.italic = True


def add_table_from_md(doc, rows):
    """从markdown表格行创建Word表格"""
    # 解析表格
    parsed = []
    for row in rows:
        # 去除首尾 |，按 | 分割
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    if len(parsed) < 2:
        return

    # 检测分隔行（如 |---|:---:|）
    if re.match(r'^[\s|:-]+$', parsed[1][0]) and len(parsed[1]) == 1:
        # 旧解析方式：整行是分隔符
        header = parsed[0]
        body = parsed[2:]
    else:
        header = parsed[0]
        # 跳过分隔行
        if all(re.match(r'^[\s:-]+$', c) for c in parsed[1]):
            body = parsed[2:]
        else:
            body = parsed[1:]

    n_cols = len(header)
    n_rows = len(body) + 1

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 解析加粗等
        run = p.add_run(re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text))
        set_run_font(run, size=10, bold=True, color=COLOR_TABLE_HEADER)
        # 设置表头背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), COLOR_TABLE_HEADER_BG)
        tcPr.append(shd)

    # 数据行
    for i, row in enumerate(body):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 解析加粗并处理行内公式
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
            # 按 $...$ 拆分，公式部分使用OMML
            formula_pattern = r'(\$[^$]+\$)'
            segments = re.split(formula_pattern, text)
            for seg in segments:
                if not seg:
                    continue
                if seg.startswith('$') and seg.endswith('$'):
                    # 行内公式 → OMML
                    formula = seg[1:-1]
                    if _OMML_AVAILABLE:
                        try:
                            omml = latex_to_omml(formula)
                            p._element.append(omml)
                            continue
                        except Exception:
                            pass
                    run = p.add_run(formula)
                    set_run_font(run, font_en=FONT_EN, size=10, bold=True,
                                 color=COLOR_FORMULA)
                    run.font.italic = True
                else:
                    run = p.add_run(seg)
                    set_run_font(run, size=10, bold=False)
    return table


def add_code_block(doc, code_lines):
    """添加代码块"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 添加边框
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)
    # 背景色
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

    run = p.add_run('\n'.join(code_lines))
    set_run_font(run, font_en=FONT_CODE, size=9, color=COLOR_CODE)


def add_formula_block(doc, formula_text):
    """添加块级公式（居中显示，优先使用OMML原生公式格式）"""
    # 去除 $$
    formula = formula_text.strip().strip('$').strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    if _OMML_AVAILABLE:
        try:
            omml = latex_to_omml(formula, display='block')
            p._element.append(omml)
            return
        except Exception as e:
            print(f"[提示] 块级公式转换失败，回退为文本: {formula} ({e})")

    # Fallback: 文本格式
    run = p.add_run(formula)
    set_run_font(run, font_en=FONT_EN, size=12, bold=True, color=COLOR_FORMULA)
    run.font.italic = True


def add_horizontal_line(doc):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# 主转换函数
# ============================================================
def md_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为 DOCX"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 设置标题样式
    for level, size in [(1, 18), (2, 15), (3, 13), (4, 12)]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = FONT_HEADING_CN
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.font.color.rgb = COLOR_HEADING
        rPr = heading_style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_HEADING_CN)
        rFonts.set(qn('w:ascii'), FONT_HEADING_CN)
        rFonts.set(qn('w:hAnsi'), FONT_HEADING_CN)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 代码块开始
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # 块级公式 $$...$$
        if stripped.startswith('$$'):
            formula_lines = [stripped]
            if not stripped.endswith('$$') or stripped == '$$':
                i += 1
                while i < n and not lines[i].strip().endswith('$$'):
                    formula_lines.append(lines[i].strip())
                    i += 1
                if i < n:
                    formula_lines.append(lines[i].strip())
            else:
                # 单行 $$...$$
                pass
            add_formula_block(doc, '\n'.join(formula_lines))
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # 去除markdown格式符号
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            heading = doc.add_heading(level=level)
            run = heading.add_run(text)
            run.font.name = FONT_HEADING_CN
            run.font.size = Pt([18, 15, 13, 12][level - 1])
            run.font.bold = True
            run.font.color.rgb = COLOR_HEADING
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), FONT_HEADING_CN)
            i += 1
            continue

        # 水平线
        if stripped in ('---', '***', '___'):
            add_horizontal_line(doc)
            i += 1
            continue

        # 表格检测
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = [stripped]
            i += 1
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # 无序列表
        list_match = re.match(r'^[-*]\s+(.+)$', stripped)
        if list_match:
            text = list_match.group(1)
            add_paragraph_with_runs(doc, text, style='List Bullet')
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^\d+\.\s+(.+)$', stripped)
        if ol_match:
            text = ol_match.group(1)
            add_paragraph_with_runs(doc, text, style='List Number')
            i += 1
            continue

        # 普通段落
        add_paragraph_with_runs(doc, stripped)
        i += 1

    doc.save(docx_path)
    print(f"[OK] 已生成: {docx_path}")
    print(f"     大小: {os.path.getsize(docx_path) / 1024:.1f} KB")


if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base_dir, '实验设计方案.md')
    # 优先使用原文件名，若被占用则使用带 _v2 后缀的文件名
    docx_path = os.path.join(base_dir, '实验设计方案.docx')
    try:
        # 尝试删除旧文件（若被占用会抛出异常）
        if os.path.exists(docx_path):
            os.remove(docx_path)
    except PermissionError:
        docx_path = os.path.join(base_dir, '实验设计方案_v2.docx')
        print(f"[提示] 原文件被占用，保存为: {docx_path}")
    md_to_docx(md_path, docx_path)
