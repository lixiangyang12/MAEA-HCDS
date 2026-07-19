from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

with open('c:/个人资料/申博材料/东北财经大学/李宇教授学术研究阅读报告.md', 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal'].font.size = Pt(11)

title = doc.add_heading('李宇教授学术研究阅读报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].font.bold = True
title.runs[0].font.color.rgb = RGBColor(30, 50, 80)

subtitle = doc.add_paragraph('东北财经大学发展规划与学科建设处处长、教授、博士生导师')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

lines = content.split('\n')
current_table = None
table_data = []

for line in lines:
    if line.startswith('## '):
        if table_data:
            if len(table_data) > 0 and len(table_data[0]) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell
            table_data = []
        heading_text = line[3:]
        heading = doc.add_heading(heading_text, level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        doc.add_paragraph()
        
    elif line.startswith('### '):
        if table_data:
            if len(table_data) > 0 and len(table_data[0]) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell
            table_data = []
        heading_text = line[4:]
        heading = doc.add_heading(heading_text, level=2)
        heading.runs[0].font.size = Pt(12)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 77, 153)
        
    elif line.startswith('#### '):
        if table_data:
            if len(table_data) > 0 and len(table_data[0]) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell
            table_data = []
        heading_text = line[5:]
        para = doc.add_paragraph(heading_text)
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(51, 102, 153)
        
    elif line.startswith('|'):
        parts = [p.strip() for p in line.split('|') if p.strip()]
        if len(parts) >= 2 and '---' not in ''.join(parts):
            table_data.append(parts)
        continue
        
    elif line.startswith('- '):
        if table_data:
            if len(table_data) > 0 and len(table_data[0]) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell
            table_data = []
        text = line[2:]
        para = doc.add_paragraph(text, style='List Bullet')
        para.runs[0].font.size = Pt(11)
        
    elif re.match(r'^\d+\.\s+', line):
        if table_data:
            if len(table_data) > 0 and len(table_data[0]) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell
            table_data = []
        match = re.match(r'(\d+)\.\s+(.*)', line)
        if match:
            text = match.group(2)
            para = doc.add_paragraph(text, style='List Number')
            para.runs[0].font.size = Pt(11)
        
    elif line.startswith('> '):
        if table_data:
            if len(table_data) > 0 and len(table_data[0]) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell
            table_data = []
        text = line[2:].strip()
        if text:
            para = doc.add_paragraph(text)
            if para.runs:
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.italic = True
                para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            para.paragraph_format.left_indent = Inches(0.5)
            doc.add_paragraph()
        
    elif line.strip() and not line.startswith('---') and not line.startswith('**'):
        if table_data:
            if len(table_data) > 0 and len(table_data[0]) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell
            table_data = []
        para = doc.add_paragraph(line)
        para.runs[0].font.size = Pt(11)
        para.paragraph_format.line_spacing = 1.5

if table_data:
    if len(table_data) > 0 and len(table_data[0]) > 0:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell

doc.save('c:/个人资料/申博材料/东北财经大学/李宇教授学术研究阅读报告.docx')
print('文档已生成')
