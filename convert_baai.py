# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

INPUT = r'c:\个人资料\申博材料\东北财经大学\BAAI2026参会指南.md'
OUTPUT = r'c:\个人资料\申博材料\东北财经大学\BAAI2026参会指南.docx'

with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()

# ---------- 全局样式 ----------
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---------- 辅助函数 ----------
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 50, 80)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    for part in re.split(r'(\*\*.*?\*\*)', text):
        run = p.add_run(part.replace('**', ''))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if part.startswith('**'):
            run.font.bold = True
    return p

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 下划线装饰
    p_border = p.paragraph_format
    return p

def add_sub_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 77, 153)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_small_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 102, 153)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bold_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    # Clear default and re-add
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # Header bg
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear',
        })
        shading.append(shd)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            # Alternating row bg
            if r % 2 == 0:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F7FB',
                    qn('w:val'): 'clear',
                })
                shading.append(shd)
    doc.add_paragraph()  # spacer
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)

# ---------- 解析 ----------
i = 0
table_data = None  # (headers, rows)

def flush_table():
    global table_data
    if table_data and table_data[0]:
        add_table(table_data[0], table_data[1])
    table_data = None

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Title (h1)
    if line.startswith('# ') and not line.startswith('## '):
        flush_table()
        add_title(line[2:])
        i += 1
        continue

    # Section heading (##)
    if line.startswith('## '):
        flush_table()
        add_section_heading(line[3:])
        i += 1
        continue

    # Sub heading (###)
    if line.startswith('### '):
        flush_table()
        add_sub_heading(line[4:])
        i += 1
        continue

    # Divider
    if line.strip() == '---':
        flush_table()
        add_divider()
        i += 1
        continue

    # Bold-only line (like **时间**：...)
    if line.startswith('**') and not line.startswith('|'):
        flush_table()
        add_bold_body(line)
        i += 1
        continue

    # Table
    if line.startswith('|'):
        # Collect consecutive table lines
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        # Parse
        headers = []
        rows = []
        for tl in table_lines:
            parts = [p.strip() for p in tl.split('|') if p.strip()]
            if '---' in ''.join(parts):
                continue
            if not headers:
                headers = parts
            else:
                # Ensure same length
                while len(parts) < len(headers):
                    parts.append('')
                rows.append(parts[:len(headers)])
        if headers:
            flush_table()
            add_table(headers, rows)
        continue

    # Bullet list
    if line.startswith('- '):
        flush_table()
        add_bullet(line[2:])
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.*)', line)
    if m:
        flush_table()
        add_numbered(m.group(2))
        i += 1
        continue

    # Body text
    if line.strip():
        flush_table()
        add_body(line)
    else:
        flush_table()
        # skip empty lines
    i += 1

flush_table()
doc.save(OUTPUT)
print('BAAI2026参会指南.docx 已生成')