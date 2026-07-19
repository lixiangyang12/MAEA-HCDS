"""读取改进牛鞭效应新途径_629.docx的完整内容并保存为txt"""
from docx import Document
import os

src = r"c:\个人资料\申博材料\企业运营与科研管理数据库\改进牛鞭效应新途径：人智协同决策系统\改进牛鞭效应新途径：人智协同决策系统_629.docx"
dst = r"c:\个人资料\申博材料\企业运营与科研管理数据库\改进牛鞭效应新途径：人智协同决策系统\改进牛鞭效应新途径_629_content.txt"

doc = Document(src)

lines = []
para_count = 0
table_count = 0

# 读取段落
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        lines.append(f"[{i}] {text}")
        para_count += 1

# 读取表格
lines.append("\n\n===== 表格内容 =====\n")
for ti, table in enumerate(doc.tables):
    table_count += 1
    lines.append(f"\n--- 表 {ti+1} ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
        lines.append(f"  行{ri+1}: " + " || ".join(cells))

# 写入文件
with open(dst, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print(f"段落数: {para_count}")
print(f"表格数: {table_count}")
print(f"输出文件: {dst}")
print(f"文件大小: {os.path.getsize(dst)} bytes")

# 打印前30段预览
print("\n===== 前30段预览 =====")
for i in range(min(30, len(lines))):
    print(lines[i])
