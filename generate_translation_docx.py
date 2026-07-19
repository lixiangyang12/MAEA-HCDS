"""
生成翻译后的docx文档
论文：Adaptive Inventory Strategies using Deep Reinforcement Learning for Dynamic Agri-Food Supply Chains
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# 创建文档
doc = Document()

# 设置页面格式
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 0:
            run.font.size = Pt(16)
        elif level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)
    return heading

def add_paragraph(text, bold=False, italic=False, indent=True):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_figure_placeholder(fig_num, caption_cn):
    """添加图片占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Figure {fig_num} - 原图保留]")
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Figure {fig_num} {caption_cn}")
    run2.font.name = '宋体'
    run2.font.size = Pt(10)
    run2.bold = True

def add_table_placeholder(table_num, caption_cn, headers, rows):
    """添加表格占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Table {table_num} {caption_cn}")
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True

    # 创建表格
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for i, row in enumerate(rows):
        row_cells = table.rows[i+1].cells
        for j, cell in enumerate(row):
            row_cells[j].text = str(cell) if cell else ""

# ==================== 开始翻译内容 ====================

# 标题
add_heading("基于深度强化学习的动态农产品供应链自适应库存策略", level=0)

# 作者信息
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Amandeep Kaur* 和 Gyan Prakash")
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ABV-印度信息技术与管理学院瓜廖尔分校，中央邦，印度")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Email: amandeepkaur@iiitm.ac.in; gyan@iiitm.ac.in")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 摘要
add_heading("摘要", level=2)
add_paragraph("农产品（agricultural products）常常受到生产和需求季节性波动的影响。预测和管理响应这些变化的库存水平具有挑战性，容易导致库存过剩或缺货（stockouts）。此外，现有文献未考虑食品供应链各层级利益相关者之间的协调（coordination）。为弥补这些研究空白，本研究聚焦于需求和提前期（lead time）不确定性下的农产品库存管理。通过实施有效的库存补货策略（inventory replenishment policy），实现整个供应链总利润最大化。然而，由于这些不确定性和产品保质期（shelf-life）的存在，问题复杂性增加，使得传统方法难以生成最优解集。因此，本研究提出一种新型深度强化学习（Deep Reinforcement Learning, DRL）算法，结合基于价值（value-based）和基于策略（policy-based）的DRL方法的优势，用于不确定性下的库存优化。所提算法通过共享利润最大化优化目标，协调利益相关者的利益和目标，激励其协作，同时考虑易腐性（perishability）和不确定性。通过在连续动作空间（continuous action space）中选择最优订货量，所提算法有效解决了库存优化挑战。为严格评估该算法，采用了生鲜农产品供应链库存的实证数据。实验结果证实，所提库存补货策略在随机需求模式和提前期场景下表现出改进的性能。研究结果为政策制定者提供了管理启示，以在不确定性下更有效地管理农产品库存。")

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run("关键词：")
run.bold = True
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run2 = p.add_run("农产品供应链；需求不确定性；深度强化学习；分布式近端策略优化；连续动作空间")
run2.font.name = '宋体'
run2.font.size = Pt(12)
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# 1. 引言
add_heading("1. 引言", level=1)

add_paragraph("农产品供应链构成了复杂的网络，涉及各参与方协同合作，从农场到客户交付产品和服务[1]。直接参与供应链的主要利益相关者包括农民、分销商、零售商和消费者，这些利益相关者之间的产品流、信息流和资金流构成了复杂的多级（multi-echelon）结构。2022年，全球农业供应链市场规模为8.6亿美元，预计到2030年将达到18.22亿美元，复合年增长率（CAGR）为10.02%[2]。农产品供应链对实现可持续发展目标（Sustainable Development Goals, SDGs）至关重要，特别是SDG 2，旨在确保每个人都能获得充足、安全和营养的食物。《2030年可持续发展议程》，特别是SDG 12，设定了将零售和消费层面的人均全球食物浪费减半，并减少整个生产和供应链中食物损失的目标[3]。")

add_paragraph("尽管农产品供应链与其他一般供应链有相似之处，但具有季节性、强易腐性、高时效性和温控存储等特定特征，使其采购、加工和分销管理颇具挑战性。食品产品一旦生产，需要经过加工、存储、配送和零售，以满足客户的多样化需求[4]。近期报告表明，约13%的食物在收获和零售之间损失。此外，17%的世界粮食产量在家庭、餐饮服务和零售层面被浪费[5]。水果和蔬菜约占食物损失的32%，其次是非素食产品占12.4%[6]。因此，实施有效的库存管理策略对于在这些产品失去市场价值之前销售至关重要。")

add_paragraph("农产品易腐性是影响农产品供应链整体决策的另一个重要问题。由于许多产品的易腐特性，季节性和需求波动对这些产品的库存管理提出了独特挑战。持有过高库存将增加采购成本、持有成本和浪费成本，而持有过低库存则导致客户流失。这些链中的产品一旦生产，其价值和质量都会退化[7]，同时还面临日益增多的法规和环境压力，如包装、可持续性问题、废物减少和回收。因此，在规划库存管理策略时，必须考虑农产品的易腐性。这种方法有助于降低库存成本，提升多级供应链网络的整体效率和效益。")

add_paragraph("在多级农产品供应链中，随机补货提前期和需求变异性显著影响库存管理、服务水平、成本和协调工作[8]。这些不确定性的综合效应为供应链管理者创造了更加复杂和具有挑战性的环境。提前期的不可预测性需要维持更高的安全库存（safety stock）以避免缺货，这反过来影响再订货点（reorder points）并增加持有成本[9]。同样，波动的需求模式需要额外的缓冲库存来满足客户订单的意外激增。这种双重不确定性通常导致比仅有一个因素变化时更高的库存水平。随机提前期和需求之间的相互作用可能导致更频繁的缺货和延期交货（backorders），潜在地削弱客户满意度并导致销售损失[10]。")

add_paragraph("此外，由于需要更频繁的订单、加急运输或临时从替代供应商采购，这种综合不确定性可能增加订货和运输成本。供应和需求双方的变异性也可能需要更复杂的预测和库存优化技术，潜在地增加运营复杂性和相关成本[11]。这些不确定性的连锁反应可以传播到整个多级结构，在上游移动时放大——表现为所谓的牛鞭效应（bullwhip effect）[12]。这可能导致更高层级的订货量波动加剧，进一步复杂化整个供应链的库存管理和协调工作。")

add_paragraph("在农产品供应链复杂和动态的环境中，多种因素导致决策和利润优化的复杂性。这些因素包括不确定的需求和供应、波动的客户偏好、食品产品的易腐特性、季节性变化以及严格的质量和安全标准。鉴于这种动态环境，持续监控和优化整个供应链的盈利能力是一项艰巨的挑战[13]。有效的库存优化通过自动确定适当的订货量，在确保产品及时可用方面发挥着关键作用。历史上，研究者和从业者采用了各种经典方法来应对农产品供应链管理中的挑战。农产品供应链管理中的许多经典方法基于博弈论（game theory）[14-17]、启发式方法（heuristics）[18]、控制理论（control theory）[19]、混合整数非线性规划（Mixed Integer Non-Linear Programming, MINLP）[20]、层次分析法（Analytical Hierarchy Process, AHP）和遗传算法（Genetic algorithm）[21]。虽然这些传统解决方案在具有特定约束的静态环境中有效，但它们经常无法跟上农产品供应链的复杂性和不断变化的条件。当面临该领域固有的多方面挑战时，这些方法的局限性变得明显。快速变化的农产品供应链格局需要更灵活、自适应和稳健的方法。这些方法应能够处理固有的不确定性，整合实时数据，并适应不断变化的市场条件。因此，越来越需要创新方法，能够在保持运营效率和盈利能力的同时，有效应对现代农产品供应链的复杂性。")

add_paragraph("近年来，机器学习（Machine Learning, ML）方法在应对农产品供应链管理复杂性方面获得了显著关注。这些先进技术为在高动态环境中提高效率、减少浪费和优化产品采购流程提供了有前景的解决方案。强化学习（Reinforcement Learning, RL）已成为一种强大的ML方法，能够在不需要预定义网络模型的情况下学习最优策略。这使其特别适合农产品供应链典型的动态环境。然而，RL算法在处理高维状态空间时性能可能受限[22-24]。为克服这些限制，深度学习与RL相结合，在各种供应链优化场景中创建了深度强化学习（DRL）。这种混合方法为复杂随机环境中的序贯决策提供了稳健的框架[25-26]。Chong等[27]提出了用于服装供应链优化的DRL模型，考虑了售罄率、服务水平和库存销售比等因素。然而，他们的模型未考虑不确定需求场景，限制了其现实适用性。Wu等[28]引入了一种无导数RL方法用于多级供应链优化，专门设计用于处理复杂随机系统。Demizu等[29]开发了基于模型的DRL用于新智能手机的库存管理。他们的方法结合了贝叶斯神经网络进行概率预测和模型无关元学习进行需求预测。最近，Yavuz和Kaya[30]提出使用DRL算法解决易腐产品的动态定价和订货问题。他们在研究中将深度Q学习（Deep Q Learning, DQN）用于离散动作空间，将软演员-评论家（Soft Actor-Critic, SAC）用于连续动作空间。该研究为基于价格和年龄依赖随机需求的易腐产品动态定价问题提供了显著解决方案。尽管取得了这些进展，挑战仍然存在。当处理由随机需求、提前期和产品易腐性导致的大状态空间时，这些模型的性能可能受限。随着状态空间的增长，SAC等方法中使用的神经网络需要更大以捕获日益增加的复杂性，导致更高的计算需求。")

# 2. 相关工作
add_heading("2. 相关工作", level=1)
add_paragraph("为全面调查现有研究领域，对文献进行了深入回顾。本节探讨了食品供应链背景下的经典方法以捕获更广泛的运营方面，以及供应链优化中基于强化学习（RL）的方法，以提供对当前知识状态的宝贵见解。")

add_heading("2.1 经典方法", level=2)
add_paragraph("库存管理对于优化运营、满足客户需求、降低成本以及在整个供应链中保持农产品质量至关重要。有效的库存管理通过最小化过量库存和减少浪费，实现资源高效利用，从而节约成本并提高盈利能力。近年来，文献中讨论了采用经典方法的农产品供应链综合研究。例如，[31]中的作者提出了具有Weibull生存和死亡特性的变质物品的订货和定价策略。作者考虑了连续随机需求，目标是基于连续审查（r, Q）策略最大化系统的总平均期望利润。为求解复杂的随机微分方程，采用直接方法-泰勒级数展开获得最优策略。在[32]中，作者为经营生鲜农产品的零售商制定了单阶段和两阶段定价与库存决策模型。Liu等[33]研究了易腐农产品的最优采购和库存补货决策，考虑了存储成本、库存短缺和过剩、价格和需求波动以及产品变质。他们关注批发商的有限期库存模型，其中需求受当前和过去价格影响。Sindhuja等[34]开发了一种变质乳制品库存模型，纳入基于质量的需求以减缓变质率并最小化总支出和成本。此外，Banerjee和Agrawal[35]设计了一种变质物品库存模型，考虑了最初受销售价格影响、随后受新鲜度条件影响的需求动态。")

add_paragraph("多位研究者提出了优化农产品供应链的各种求解方法，包括博弈论、模糊层次分析法（AHP）、线性规划（LP）、混合整数线性规划（MILP）、混合整数规划（MIP）、MINLP、遗传算法（GA）、多目标遗传算法-II（MOGA-II）、非支配排序遗传算法（NSGA）和粒子群优化（PSO）。Validi等[36]引入了一种基于逼近理想解排序法（Technique for Order of Preference by Similarity to Ideal Solution, TOPSIS）的稳健解决方案，用于爱尔兰牛奶分销的两层乳制品供应链。他们应用TOPSIS根据总成本和CO₂排放之间的权衡对运输路线进行排序。此外，Mirakhorli等[37]利用模糊多目标线性规划（FMOLP）方法优化面包生产工厂的供应链，旨在实现同时最小化总成本和交付时间的帕累托最优解。Galal等[38]研究了两级农产品供应链，探讨了不同订货量对成本、排放和服务水平的影响。他们开发了离散事件仿真模型，以考虑随机需求和提前期变异性，以及对服务水平和产品生命周期的影响。Miranda等[39]通过GA和TOPSIS进行多目标优化，对绿色供应链进行建模和优化，然后通过橙汁多级供应链的开发和分析进行验证。然而，这些模型未将供应链中的不确定性作为主要问题考虑。为解决这一问题，Zhao和Wang[40]提出了一种不确定性下库存控制的GA优化模型。此外，Mousavi等[41]提出了改进的PSO来解决两级供应链中的库存控制问题，作为MINLP优化问题。虽然这两种元启发式优化技术在多维复杂搜索空间中可能有效，但在某些场景下可能不太适用，如不确定需求下的库存优化。这两种算法可能难以找到在各种需求场景下表现良好的稳健解。")

add_paragraph("当前关于经典库存管理方法的文献往往缺乏将创新技术全面整合到供应链实践中。特别是在农产品供应链中，在探索不确定性下的库存管理及其对供应链动态更广泛影响方面存在明显空白。")

add_heading("2.2 基于RL的方法", level=2)
add_paragraph("由于季节性波动、市场趋势或不可预见事件，库存需求模式可能随时间表现出非平稳性。基于RL的方法可以通过基于近期经验持续更新策略来适应非平稳环境，使其在动态变化的需求场景中保持性能。基本上，RL及其变体将问题表述为马尔可夫决策过程（Markov Decision Process, MDP），不需要显式数学模型即可在高度复杂的场景中开发最优策略。在供应链场景中，RL技术已被探索用于优化各个方面，如订货决策[25]、易腐产品库存管理[24]以及采购和分销功能。Kara等关注多零售商框架内易腐库存物品的订货策略，使用RL时序差分算法——特别是Q-Learning和状态-动作-奖励-（下一）状态-（下一）动作（SARSA）——来表述问题。Dogan和Güner[22]基于随机动态规划和基于代理的仿真，在多零售商竞争环境中提出了类似的学习模型。遗憾的是，此类算法只能处理有限状态-动作空间的问题，这引发了许多关于在高度复杂供应链网络中进行库存管理的DRL解决方案的研究。")

add_paragraph("最近，Boute等[26]描述了DRL算法在库存控制中的关键设计选择。另一项研究[42]提出了不确定性环境下易腐材料批量问题的DRL解决方案，但其性能限于单一供应商设置。DRL架构被引入以解决多维动作空间问题，但仍需要重新调整网络架构和训练算法以适应连续动作空间。保持DRL的类似风格但采用不同类别的DRL——称为策略梯度（Policy Gradient）——被用于处理随机供应链环境下的多级发散系统。特别是，近端策略优化（Proximal Policy Optimization, PPO）已被证明是随机多级供应链优化中大规模连续动作空间的稳健算法[43]。Geevers等[44]的研究调查了PPO算法在多级供应链网络的线性、发散和一般结构中的性能，但其性能仅限于特定场景。Hubbs等[45]开发了基于PPO的化学生产动态调度，但其性能限于单阶段供应链网络。此外，Tian等[46]提出了带有PPO和GRU注意力机制的A2C用于仓库库存补货。")

add_paragraph("尽管DRL在应对包括库存管理、生产计划、质量控制和可持续性改进等各种供应链挑战方面的应用日益增多，但文献中关于其应用于复杂农产品供应链仍存在显著空白。我们对现有文献的全面回顾显示，在农产品供应链背景下深入探索先进RL模型的研究明显稀缺。虽然DRL在其他供应链领域显示出前景，但其在解决农产品系统特定挑战方面的潜力在很大程度上尚未开发。")

# 3. 问题描述
add_heading("3. 问题描述", level=1)
add_heading("3.1 问题表述", level=1)
add_paragraph("在多级农产品供应链网络中，考虑了三个利益相关者——农业农场、分销中心和零售商——它们构成了如图1所示的复杂网络。农民在农场生产新鲜农产品并将其销售给分销中心。分销中心进一步将其销售给零售商。由于所考虑的产品是易腐的且保质期短，假设所有具有不同剩余保质期的产品在每个周期开始时从分销中心合并并一起运输到零售商。分销中心接收零售商的订单，主要负责订单处理、补充从农民采购的新鲜农产品库存，以及沿上游和下游供应链运输和分销。在高需求方差和不可预测的提前期下，需要做出库存决策以优化相关库存成本（如持有成本、缺货成本和运输成本），从而最大化整个供应链网络的收入。考虑到这一场景，开发了库存优化模型。模型表述中使用的主要符号——参数和变量汇总于表1。")

add_paragraph("为构建生鲜农产品三级供应链网络模型，做出以下假设：")
add_paragraph("• 每个层级有有限数量的利益相关者", indent=False)
add_paragraph("• 可用的运输车辆有限", indent=False)
add_paragraph("• 不允许同一层级实体之间的产品流动", indent=False)
add_paragraph("• 每个零售商可以在特定补货间隔从单一分销中心订货", indent=False)

add_figure_placeholder("1", "多级农产品供应链网络")

# 表1占位符
add_table_placeholder("1", "主要符号", ["符号", "含义"], [
    ["集合", ""],
    ["p∈P", "农业农场收获的作物集合"],
    ["j∈F", "农业农场位置集合"],
    ["k∈K", "分销中心集合"],
    ["m∈M", "车辆集合"],
    ["t∈T", "时间段集合"],
    ["c∈C", "零售商集合"],
    ["参数", ""],
    ["c_{p,j,t}", "时间段t从农业农场j采购产品p的采购成本"],
    ["c'_{p,c,k}", "时间段t从分销中心k采购产品p的采购成本"],
    ["c_f^p", "产品p的固定采购成本"],
    ["c_{h,p,k}", "时间段t分销中心k持有产品p的持有成本"],
    ["c_{h,p,c}", "时间段t零售商c持有产品p的持有成本"],
    ["v_{t,p,k}", "时间段t分销中心k可用的产品p现有库存"],
    ["v_{t,p,c}", "时间段t零售商c可用的产品p现有库存"],
    ["d_{t,p,c}", "时间段t零售商c对产品p的需求"],
    ["c_{w,p,k}", "分销中心k产品p的浪费成本"],
    ["c_{w,p,c}", "零售商c产品p的浪费成本"],
    ["c_{sh,p,k}", "分销中心k产品p的缺货成本"],
    ["c_{sh,p,c}", "零售商c产品p的缺货成本"],
    ["变量", ""],
    ["q_{j,p,k,t}", "时间段t分销中心k从农场j采购产品p的数量"],
    ["q_{t,c,p,k}", "时间段t零售商c从分销中心k采购产品p的数量"],
    ["l", "分销中心的补货提前期（天）"],
    ["l'", "零售商的补货提前期（天）"],
])

doc.add_paragraph()
add_paragraph("库存流从分销中心在每个日开始时通过从农业农民采购来提高其库存水平开始。设k = {1, 2, ..., K}为从农民采购产品订单的分销中心。在库存采购过程中，考虑了采购成本、持有成本、交付成本（或运输成本）、浪费成本、缺货和断货成本，目标函数是确定最优订货量以最大化总收入和库存周转率，同时最小化库存相关成本和断货：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("min (C_p^t + C_h^t + C_trans^t + C_w^t + C_sh^t)    (1)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("其中采购成本、持有成本、浪费成本、缺货成本和运输成本分别按公式(2)-(7)计算。产品变质取决于新鲜度，可计算为 q_t^p = q_p · e^(-δ_p·μ_p·F_{p,t})。其中δ_p表示农产品p特有的浪费惩罚系数，μ_p是新鲜农产品p随时间退化的敏感因子（μ_p > 0），F_{p,t}∈[0,1]表示产品p在时间t的新鲜度水平，F_{p,t}=1表示新采购的物品，F_{p,t}=0表示完全变质。")

add_paragraph("对于高度易腐产品（如叶菜），δ_p=0.18和μ_p=6.0；对于低易腐性物品，δ_p=0.05和μ_p=2.0。")

add_paragraph("订单履行过程遵循先进先出（First-In-First-Out, FIFO）的订单管道。每个下达的订单都与一个提前期计数器相关联，在每个时间步递减。在下一个时隙(t+1)，分销中心和零售商的库存水平更新为：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("I_k(t+1) = I_k(t) + q_{p,k,t} - Σq_{p,c,t}, ∀k∈K    (10)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("I_c(t+1) = I_c(t) + q_{c,p,t} - d̃_c(t+1), ∀c∈C    (11)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("目标是在以下约束条件下最小化所有时间段的总库存相关成本：约束C1和C2分别对分销中心和零售商的库存物品设置上限，不超过最大库存持有能力。约束C3和C4假设每个分销中心和零售商在每个时间段t内分别只能从单一农业农场和单一分销中心订货。约束C5确保运输车辆总数不超过最大允许数量。约束C6检查每辆运输车辆的最大容量并防止超载。")

add_heading("3.2 随机环境下的需求和补货提前期建模", level=2)
add_paragraph("农产品供应链在高度随机的环境中运营，具有众多不确定性和变异性来源。这些包括受季节性、天气条件和消费者偏好变化等因素影响的波动客户需求模式；以及受运输延迟和供应中断影响的不可预测的补货提前期。农产品的易腐特性引入了额外的复杂性，因为产品质量和保质期可能差异显著。这种多方面的随机环境需要复杂的库存管理策略，能够适应快速变化的条件，平衡多个竞争目标，并在信息不完整和结果不确定的情况下做出最优决策。")

add_paragraph("客户需求和补货提前期的随机性使用多种概率分布建模：特别是正态分布（Normal N(μ,σ²)）、伽马分布（Gamma(α,β)）、威布尔分布（Weibull(u,λ)）和指数分布（Exponential Exp(λ)）。这种方法使我们能够捕获农产品供应链中常见的各种需求和提前期模式及不确定性。正态分布用于表示围绕均值的对称需求波动，伽马分布用于建模季节性产品中常见的正偏需求场景。威布尔分布以其灵活性著称，用于捕获更复杂的需求模式，包括具有不同偏度和峰度的模式。对于实际提前期场景，考虑指数分布。通过纳入这些不同的分布，我们提出的库存优化模型开发了适用于各种现实场景的稳健和适应性框架，增强了其在农产品领域不同产品类别和市场条件下的适用性。")

# 4. DRL框架
add_heading("4. 基于DRL的智能库存管理框架", level=1)
add_paragraph("农产品供应链网络的成本优化涉及复杂交互、多个决策点、不确定需求和供应动态，可以用MDP映射。它为不确定性下和随时间的决策问题建模和求解提供了稳健的框架。")

add_heading("A. MDP模型", level=2)
add_paragraph("在库存管理的背景下，目标是通过最小化持有成本、订货成本和缺货来优化整个农产品供应链的利润。供应链中的决策通常随时间顺序做出。MDP本质上捕获了序贯决策的概念，这在某一层级的行动影响下游和上游运营时至关重要。本研究考虑由单一农业农场、单一分销中心和多个零售商组成的网络。然而，该网络可以通过在基础供应链中添加更多分销中心和零售商作为子网络来进一步扩展。MDP的关键组成部分包括状态、动作、奖励，定义如下：")

add_paragraph("• 状态：状态捕获时间段t库存状态的基本信息。状态包括(1)分销中心和零售商的库存水平，(2)补货提前期，(3)客户需求的信息", indent=False)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("x_t = {I_{c,t}, I_{k,t}, l̃_k, l̃_c, d̃_{c,t}, t_n}, k∈K, c∈C    (13)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("其中I_{k,t}和I_{c,t}分别表示分销中心和零售商的当前库存位置，l̃_k和l̃_c分别表示分销中心和零售商的预测补货提前期，d̃_{c,t}表示零售商端最终客户的预测需求，t_n∈{1,2,...,7}表示对应一周中各天的7维向量。")

add_paragraph("• 动作：动作空间考虑分销中心和零售商对应的订货量", indent=False)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("y_t = {q_{k,t}, q_{c,t}}, ∀k∈K, c∈C    (14)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("此处，动作空间考虑连续订货量。连续动作空间允许在优化订货量以满足特定目标方面具有更大的灵活性，同时最小化库存相关成本、避免缺货或平衡库存水平。")

add_paragraph("• 奖励：每个代理在特定状态x_t下执行动作y_t后获得的奖励。净利润通过考虑基础供应链中各实体产生的收入以及零售商和分销中心产生的库存成本来计算", indent=False)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("z_t = R_k(t) + ΣR_c(t) - (C_k(t) + ΣC_c(t)), 若满足C1至C6    (15)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_heading("B. DQN算法", level=2)
add_paragraph("DQN算法通过用神经网络逼近价值函数来获得最优库存补货策略。它以权重更新对应每个状态对的形式将信息存储在神经网络中。在所考虑的场景中，环境持续观察状态空间信息，以决定维持高效供应链的最优订货量。DQN算法使用主网络和目标网络来逼近最优动作价值函数Q(x,y)，该函数估计在给定状态x下采取动作y的期望累积奖励z_t。主网络负责选择动作，而目标网络为学习提供稳定的目标值。动作价值函数使用贝尔曼方程更新：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Q(x,y) = z_t + γ·max_{y'} Q(x',y')    (19)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("其中x'=x(t+1)表示下一状态。学习过程涉及将经验(x,y,z,x')存储在经验回放缓冲区中。在训练期间，从缓冲区采样小批量经验来更新主网络。损失函数L(θ)衡量预测值和目标值之间的差异，使用梯度下降最小化：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("L(θ) = E_{(x,y,z,x')}[(z + γ·max_{y'} Q(x',y';θ_) - Q(x,y;θ))²]    (20)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_figure_placeholder("2", "DQN算法的库存优化")

add_heading("C. 软演员-评论家（SAC）算法", level=2)
add_paragraph("SAC算法设计用于处理随机环境中的连续动作空间。基本上，SAC结合了基于价值和基于策略的RL方法的优势。该算法旨在最大化期望回报和策略熵之间的权衡。熵项通过防止策略变得过于确定性来鼓励探索。SAC算法由一个演员（策略）网络、两个评论家（价值）网络和一个价值函数组成，每个都由神经网络参数化。")

add_paragraph("SAC的总体目标是在探索期间纳入熵的同时最大化期望回报：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("J(π) = Σ_{t=0}^{T} E_{(x_t,y_t)~ρ_π}[z(x_t,y_t) + α·H(π(·|x_t))]    (21)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("其中α是平衡奖励最大化和熵的温度参数，H(π(·|x_t)) = -E[log π(y_t|x_t)]表示策略的熵。SAC使用两个Q函数Q_{θ1}(x,y)和Q_{θ2}(x,y)来减轻价值估计中的正偏差。通过将熵最大化整合到RL目标中，SAC确保了探索和利用之间的平衡权衡。双Q函数、价值函数和随机策略网络的使用，以及特定的损失函数更新，为算法提供了高效处理复杂和不确定环境所需的稳健性。")

add_heading("D. 提出的协作A3C-DPPO算法", level=2)
add_paragraph("为处理控制产品订货量和实现最优决策所需的连续动作空间，采用了A3C-DPPO算法。该算法以分布式方式执行训练过程，并遵循与异步优势演员-评论家（Asynchronous Advantage Actor-Critic, A3C）多代理学习类似的训练（Zhang等，2019）。与A3C类似，DPPO有多个放置在不同位置的学习代理，它们独立与环境交互、进行模型训练，并通过收集来自本地网络的每个代理的梯度来定期更新全局PPO网络。")

add_paragraph("在所考虑的农产品供应链中，每个零售商（作为本地代理）连接到单一分销中心（作为全局模型），该中心通过收集本地模型的信息来更新，并将全局参数同步到本地参数，遵循A3C-DPPO算法的训练过程。类似地，每个分销中心连接到单一农业农场。分销中心的库存模型采用了类似的分布式训练框架，其中供应商（农业农场）作为全局代理，分销中心作为本地代理，以实现利益相关者之间的协作。")

add_paragraph("全局模型通过收集每个代理的梯度来更新相应参数，并将更新后的模型发送给每个本地代理进行下一轮训练。在分布式设置中，PPO可以与A3C协调，以高效利用多个代理收集的信息，同时保持策略稳定性。协作方面源于A3C的探索和PPO的优化之间的协同作用。它们共同协作改进策略和价值函数，以实现期望目标。")

add_paragraph("所提A3C-DPPO架构在分层供应链网络中引入了协作学习框架。零售商作为本地代理与分销中心交互。分销中心作为全局代理，聚合来自多个零售商的反馈。每个代理维护一个本地演员-评论家网络，根据其与环境交互独立学习策略和价值函数。这些本地梯度定期在相应的全局模型（分销中心）处聚合，然后更新共享参数并重新分配给本地代理。")

add_paragraph("演员网络学习策略π_{θk}(v_k(t)|u_k(t))，该策略确定给定当前状态u_k(t)下动作的概率分布。评论家网络估计价值函数V_{φk}(u_k(t))，评估给定状态u_k(t)的期望回报。零售商k的策略更新旨在通过调整策略参数θk来最大化期望奖励：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("θ_k(t+1) = θ_k(t) + α_a · ∇_{θk} log π_{θk}(v_k(t)|u_k(t)) · A_k^π    (25)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("其中α_a是演员网络的学习率，∇_{θk} log π_{θk}(v_k(t)|u_k(t))是所采取动作对数概率的梯度。为更新演员网络的权重，需要计算采取某一动作相对于其他动作的优势，使得策略朝着增加采取更好动作机会的方向更新。对应每个零售商的优势函数A_k^π定义为：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A_k^π = r_k(t) + γ·V_{φk}(u_k(t+1)) - V_{φk}(u_k(t))    (26)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("全局代理使用DPPO方法更新其策略和价值函数。分销中心通过协调所有零售商的行动来优化整体库存管理。分销中心的策略更新通过裁剪策略比率确保稳定更新：")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("θ_d(t+1) = θ_d(t) + α_a · min(∇_{θd} log π_{θd}(v(t)|u(t)), clip(η_t^d, 1-ε, 1+ε)·∇_{θd} log π_{θd}(v(t)|u(t))) · A^π    (28)")
run.font.name = 'Cambria Math'
run.font.size = Pt(11)

add_paragraph("其中A^π = r(t) + γ·V_{φd}(u(t+1)) - V_{φd}(u(t))表示分销中心的优势函数，η_t^d(θd)表示当前策略和旧策略之间的概率比率。如果比率η_t^d(θd) > 1，意味着新策略很可能被选择，否则旧策略继续。这通过调整优化目标来防止新策略显著偏离旧策略，同时通过自适应散度惩罚系数避免对优化目标的过高估计。")

add_paragraph("为确保异步学习过程中的稳定性和同步性，纳入了全局参数的指数移动平均作为稳定机制。更新全局参数后，每个本地零售商接收更新后的全局参数：θ_k ← θ_d, φ_k ← φ_d, ∀k∈{1,2,...,K}。这确保了所提算法有效管理异步梯度更新，同时在制药供应链场景中保持稳定性和同步性。总结混合A3C-DPPO算法的伪代码在算法1中概述。")

# 算法1
add_heading("算法1：农产品供应链的协作A3C-DPPO算法", level=3)
add_paragraph("1: 初始化全局演员评论家网络的参数θ_d和φ_d", indent=False)
add_paragraph("2: 为每个零售商初始化本地演员评论家网络的参数θ_k和φ_k", indent=False)
add_paragraph("3: 初始化本地和全局网络的演员和评论家网络超参数", indent=False)
add_paragraph("4: 对于每个epoch：", indent=False)
add_paragraph("5:   重置全局演员(Δθ_d)和全局评论家(Δφ_d)的累积梯度", indent=False)
add_paragraph("6:   对于每个本地代理（零售商）k并行执行：", indent=False)
add_paragraph("7:     模拟本地代理与环境的交互，收集状态-动作-奖励-下一状态元组", indent=False)
add_paragraph("8:     计算本地代理的优势函数A_k^π", indent=False)
add_paragraph("9:     计算本地演员的策略梯度∇_{θk}", indent=False)
add_paragraph("10:    计算时序差分误差δ_{kt}和本地评论家的价值梯度∇_{φk}", indent=False)
add_paragraph("11:    将这些梯度累积到全局累积梯度(Δθ_d)和(Δφ_d)中", indent=False)
add_paragraph("12:  结束循环", indent=False)
add_paragraph("13:  使用累积梯度更新全局演员和评论家参数", indent=False)
add_paragraph("14:  将更新后的全局参数与所有本地代理同步", indent=False)
add_paragraph("15: 结束循环", indent=False)

# 5. 数值实验
add_heading("5. 数值实验", level=1)
add_paragraph("为评估所提算法的有效性，在需求和提前期不确定性下进行了库存管理的实验仿真。在农产品供应链场景中，设置由单一农业农场和单一分销中心组成，服务于不同数量的零售商，形成代表性子网络。实验仿真在配备Core i9处理器、32GB内存和RTX 3060显卡的HP工作站上进行。所提A3C-DPPO算法使用Python 3.6和Pytorch工具实现，并使用真实数据集评估性能。A3C-DPPO算法的超参数如表2所示。")

add_table_placeholder("2", "A3C-DPPO算法超参数", ["参数", "值"], [
    ["隐藏层大小", "2"],
    ["节点大小", "{64, 128}"],
    ["学习率（演员网络）", "0.00005"],
    ["学习率（评论家网络）", "0.0001"],
    ["裁剪参数ε", "[0.1, 0.2, 0.3]"],
    ["激活函数（策略网络）", "tanh"],
    ["激活函数（价值网络）", "ReLU"],
])

add_heading("A. 数据和参数", level=2)
add_paragraph("本研究考虑三种农场新鲜农产品——石榴（产品1）、杨梅（产品2）和苹果（产品3）——从农场供应到分销中心，随后分销给多个零售商。与每种产品在整个供应链中的库存成本、生产成本、单位销售价格和其他库存相关成本汇总于表3至表5。")

add_table_placeholder("3(a)", "农业农场相关成本参数", ["产品", "单位库存成本($)", "单位生产成本($)"], [
    ["1", "0.04", "3"],
    ["2", "0.12", "8.2"],
    ["3", "0.14", "9.5"],
])

add_table_placeholder("4(a)", "分销中心相关成本参数", ["产品", "初始库存", "单位采购价($)", "单位售价($)", "单位持有成本($)", "单位缺货成本($)", "单位浪费成本($)", "固定订货价($)", "最大库存容量"], [
    ["p=1", "80", "4.25", "8.28", "0.17", "0.24", "0.02", "300", "500"],
    ["p=2", "80", "10.05", "17", "0.2", "0.34", "0.14", "300", "500"],
    ["p=3", "80", "11.25", "19.2", "0.36", "0.44", "0.17", "300", "500"],
])

add_table_placeholder("5(a)", "分销中心下零售商1的参数", ["产品", "初始库存", "单位售价($)", "单位持有成本($)", "单位缺货成本($)", "单位浪费成本($)", "固定订货价($)", "最大库存容量", "需求模式"], [
    ["p=1", "60", "25.8", "0.2", "0.32", "0.04", "200", "200", "正态"],
    ["p=2", "20", "35", "0.28", "0.38", "0.16", "200", "200", "威布尔"],
    ["p=3", "30", "50", "0.5", "0.46", "0.19", "200", "200", "伽马"],
])

add_paragraph("对于需求变化，客户遵循正态分布N(μ,σ²)的混合。例如，周一至周四的客户需求遵循具有最低均值的正态分布，周五的需求遵循具有中等均值的正态分布，周末的需求遵循具有最高均值的正态分布，如表6所示。此外，由于假设客户需求是独立同分布（i.i.d）且具有固有波动，利用伽马分布来捕获需求模式中的随机性和变异性，由尺度参数m和形状参数u表征。")

add_table_placeholder("6(a)", "正态分布需求模式", ["星期", "需求模式"], [
    ["周一至周三", "最低 N(3,1.5)"],
    ["周四至周五", "中等 N(6,1)"],
    ["周六至周日", "最高 N(12,2)"],
])

add_heading("B. 基准方案", level=2)
add_paragraph("所提A3C-DPPO算法的性能与以下基准算法进行比较：")
add_paragraph("• (s,S)策略：考虑周期性审查(s,S)策略作为经典库存策略，其中当库存位置降至s单位以下时重新订货，需要在每个时间段t∈T订货以将库存补充至S单位。在此策略下，每隔t时间段审查库存位置。这里，s被视为再订货点，S被称为订货至水平。", indent=False)
add_paragraph("  a) 确定性需求场景：在此情况下，假设需求遵循名义分布，相应参数值采用[47]中提出的值。", indent=False)
add_paragraph("  b) 随机需求场景：为建模需求不确定性，采用椭球不确定性集。定义椭球分布的参数按照[47]中概述的方法进行调优。", indent=False)
add_paragraph("• DQN算法：DQN策略是基于价值的策略，用于优化复杂库存运营。DQN的超参数以与[30]中考虑的类似方式进行调优。", indent=False)
add_paragraph("• SAC算法：为评估连续动作空间随机环境下的性能，考虑SAC算法。相关超参数根据[30]中提出的配置进行调优。", indent=False)
add_paragraph("• 集中式PPO：PPO算法在集中式学习设置中实现，其中全局代理观察所有环境状态和动作，通过聚合经验实现策略更新[40]。", indent=False)
add_paragraph("• 联邦DRL：采用联邦强化学习方法，实现跨多个代理的去中心化训练。本地模型定期与全局模型同步，以保护隐私并减少通信开销，同时仍受益于共享策略改进[48]。", indent=False)

add_heading("C. 收敛性分析", level=2)
add_paragraph("本节通过分析两个关键参数——折扣因子和裁剪因子——来检验所提算法的收敛行为。如图3(a)所示，折扣因子决定了决策过程中即时奖励和未来奖励之间的平衡。研究发现，较高的折扣因子值，特别是当γ接近1时，导致更好的性能。这表明当算法更加重视长期后果时，能够获得最优结果，这对于管理易腐商品的库存和应对农产品供应链典型的季节性变化至关重要。")

add_figure_placeholder("3(a)", "不同折扣因子下的收敛行为")
add_figure_placeholder("3(b)", "不同裁剪因子下的收敛行为")
add_figure_placeholder("4(a)", "与基准方法的收敛比较")
add_figure_placeholder("4(b)", "不同需求场景下的平均奖励")

add_paragraph("图3(b)关注裁剪因子，这是A3C-DPPO算法的一个组件，限制新策略在更新期间可以偏离旧策略的程度。考虑裁剪因子值0.1、0.2和0.3，研究发现0.2产生最佳结果。这个中间值可能在允许必要的策略调整和保持学习稳定性之间取得了最佳平衡。")

add_paragraph("此外，为评估所提方法的学习效率和策略稳定性，比较了A3C-DPPO与基准方法的收敛行为，如图4(a)所示。结果表明，A3C-DPPO以最快的收敛速率实现了最高的平均奖励，约在300个epoch后稳定。这归因于其协作分布式架构，允许本地代理通过与环境的并行交互异步更新共享策略。共享梯度更新导致更快的适应和更高效的样本学习，即使在需求不确定性和可变提前期下也是如此。联邦DRL也表现良好，但由于其周期性模型聚合方案和缺乏代理间协调而收敛较慢。集中式PPO收敛适中，受其依赖集中式采样的阻碍，这在多代理环境中变得不太可扩展。SAC虽然在连续控制方面有效，但由于其对熵正则化和探索参数的敏感性，显示出较慢的学习进度和较低的最终奖励值。如图所示，DQN表现最差，突显了基于价值的方法在高维连续动作空间中的局限性。")

add_paragraph("表7和图4(b)说明了在考虑各种需求分布的情况下，最优补货水平下供应链实现的最大总利润。结果表明，所提算法对库存管理固有的复杂动态表现出卓越的适应性，特别是在农产品供应链背景下。具体而言，该算法对需求不确定性、季节性波动和其他动态因素（如产品易腐性和可变提前期）显示出增强的响应能力。")

add_table_placeholder("7", "所提算法在供应链中获得的总利润", ["产品", "需求分布", "制造商", "分销中心", "零售商", "供应链系统"], [
    ["产品1", "威布尔", "59,656", "1,823", "1,173", "62,652"],
    ["", "伽马", "56,322", "1,652", "998", "58,972"],
    ["产品2", "威布尔", "58,980", "1,225", "1,000", "61,205"],
    ["", "伽马", "60,151", "1,532", "1,144", "62,827"],
    ["产品3", "威布尔", "57,443", "1,342", "1,298", "60,083"],
    ["", "伽马", "59,662", "1,409", "1,287", "62,358"],
])
add_paragraph("注：利润单位为美元。", indent=False)

add_heading("D. 与基准方案的性能比较", level=2)
add_paragraph("在本小节中，比较了所提基于A3C-DPPO的库存策略与基准算法的性能。如图5所示，所提协作A3C-DPPO框架始终优于所有基准，在随机条件下展示了卓越的学习能力和韧性。在低需求方差(0%)下，A3C-DPPO相比DQN提高21.7%，相比SAC提高5.7%，相比集中式PPO提高7.7%，相比联邦DRL提高3.7%。随着需求方差增加到80%，A3C-DPPO优于DQN 220.0%，优于SAC 33.3%，优于集中式PPO 18.5%，优于联邦DRL 6.7%。这些发现突显了A3C-DPPO框架在不确定环境中的稳健性，而最先进的DRL方法则遭受急剧退化。")

add_figure_placeholder("5", "不同需求方差下的性能比较")
add_figure_placeholder("6", "A3C-DPPO算法在不同提前期分布下的性能")

add_paragraph("图6描绘了所提算法在随机提前期场景下的性能。提前期不确定性使用指数分布和伽马分布映射，代表不同的需求分布场景。散点图显示了四个场景中库存成本如何随提前期变化。每个场景的趋势线（虚线）表明，随着提前期增加，库存成本通常上升。这表明较长的提前期与较高的库存成本相关。点的广泛散布也表明任何给定提前期下成本的显著变异性，突显了供应链管理的复杂性和不可预测性。此外，所提算法通过有效处理补货提前期的轻微变化，展示了稳健性，从而很好地适应不同的提前期情况。")

add_table_placeholder("8", "不同提前期下的成本比较", ["λ", "平均提前期", "DQN最优成本", "DQN差距(%)", "SAC最优成本", "SAC差距(%)", "集中式PPO最优成本", "PPO差距(%)", "联邦DRL最优成本", "DRL差距(%)", "协作A3C-DPPO最优成本"], [
    ["0.1", "10天", "6872±95", "40.5%", "5564±82", "13.6%", "4998±70", "2.0%", "4960±65", "1.3%", "4898±52"],
    ["0.5", "2天", "5440±88", "25.8%", "4986±74", "15.3%", "4455±63", "3.0%", "4388±58", "1.5%", "4324±49"],
    ["1.0", "1天", "4876±76", "25.0%", "4265±62", "9.4%", "3980±55", "2.1%", "3945±51", "1.2%", "3900±47"],
    ["1.5", "0.67天", "4198±70", "23.8%", "3782±58", "11.5%", "3560±50", "5.1%", "3458±44", "2.0%", "3392±42"],
    ["2.0", "0.5天", "3870±62", "61.6%", "2900±47", "21.0%", "2600±40", "8.4%", "2498±38", "4.2%", "2398±35"],
])

add_paragraph("此外，图7描绘了基准方案在不同提前期场景下平均库存成本的变化。从图中可以看出，所提协作A3C-DPPO算法在所有提前期范围内始终产生最低成本，在短期和长期场景中均优于基准方法。具体而言，在0-2天范围内，A3C-DPPO在所有方法中实现最低库存成本，相比SAC和联邦DRL改进约6-10%。随着提前期增加到6-8天，A3C-DPPO产生的成本比集中式PPO低约8%，比DQN低约10%。")

add_figure_placeholder("7", "不同提前期场景下的性能比较")
add_figure_placeholder("8", "不同提前期场景下的收敛时间")

add_paragraph("图8说明了以训练迭代次数衡量的收敛时间。观察到，随着提前期增加，收敛时间增加。在所有算法中，所提算法在所有提前期场景下始终实现最快收敛。当提前期延长到最具挑战性的8天以上场景时，A3C-DPPO仍在170次迭代以下收敛，优于SAC、集中式PPO、联邦DRL和DQN。A3C-DPPO的优越收敛行为可归因于其协作梯度共享机制和并行化的演员-评论家更新，这增强了样本效率和学习稳定性。")

add_paragraph("为验证所提协作框架的可扩展性和有效性，设置了包含3个农场、4个分销中心和10个零售商的供应链网络。每个农场生产不同类别的产品，具有不同的易腐性和提前期特征。如表9所示，所提算法实现了最高的平均累积奖励，优于联邦DRL、集中式PPO和SAC。")

add_table_placeholder("9", "大型供应链网络下的性能比较", ["方法", "平均奖励", "平均库存成本", "收敛(迭代次数)"], [
    ["SAC", "52,800", "17,300", "190"],
    ["集中式PPO", "56,100", "16,000", "165"],
    ["联邦DRL", "58,300", "15,200", "150"],
    ["协作A3C-DPPO", "61,200", "14,500", "135"],
])

add_heading("E. 敏感性分析", level=2)
add_paragraph("图9展示了不同易腐率(δ)下平均库存成本分量的敏感性。随着δ从0.05增加到0.20，所有成本分量呈上升趋势，其中浪费惩罚经历最显著的激增。这验证了易腐性建模在农产品供应链中的影响。浪费成本的激增强调了整合新鲜度感知决策的重要性，这正是所提A3C-DPPO框架所实现的。此外，由于需要更频繁的交付和保守的库存持有策略，运输和断货成本也上升。总体而言，分析证实了所提协作A3C-DPPO模型在高易腐性压力下的稳健性和适应性。")

add_figure_placeholder("9", "易腐率(δ)对库存成本分量的影响")

# 6. 理论贡献和管理启示
add_heading("6. 理论贡献和管理启示", level=1)
add_heading("6.1 理论贡献", level=2)
add_paragraph("本研究通过引入新型A3C-DPPO算法，对农产品供应链管理的理论理解做出了重要贡献。这种新方法专门设计用于应对农产品供应链网络面临的独特挑战，特别关注管理源自随机需求和可变补货提前期的不确定性。所提算法在几个关键方面推进了现有研究。首先，该框架有效处理了需求和提前期的固有不确定性，这些是影响农产品供应链绩效的关键因素。通过显式建模这些不确定性，该框架在波动市场条件下实现了更准确和稳定的决策。其次，该算法强调动态库存策略的优化，认识到其在不确定性缓解和利润最大化中的战略重要性。这种方法允许更细致和有效的库存管理策略，扩展了之前主要关注集中式或非协作范式的研究。")

add_heading("6.2 管理启示", level=2)
add_paragraph("从实践角度，本研究为在不确定和动态环境中运营的供应链管理者提供了几个有价值的见解和工具。")
add_paragraph("• 增强的可扩展性和效率：所提协作A3C-DPPO框架通过实现跨代理的并行化训练来提高计算性能。这使得模型能够随网络规模有效扩展，适应更大或更复杂的农产品供应链，而不影响学习质量。", indent=False)
add_paragraph("• 改进的响应能力和韧性：通过优化补货决策，该算法增强了供应链响应不确定性和在面对意外变化时保持韧性的能力。这种响应能力有助于降低库存成本、提高填充率，并增强对供应中断或季节性变异的韧性。", indent=False)
add_paragraph("• 持续学习和适应：该算法通过允许分销中心和零售商作为智能代理来促进实时学习。这些代理随着环境动态的变化调整其订货策略，实现了能够在高度可变条件下保持性能的自我改进系统。", indent=False)
add_paragraph("• 库存管理者的决策支持：研究结果为库存管理者提供了不确定性下最优库存策略的可操作见解。它帮助管理者在服务水平目标、成本效率和新鲜度损失之间取得平衡，通过响应性策略制定。这最终有助于更好的客户服务、改善的运营韧性和长期竞争力。", indent=False)

add_heading("6.3 与SDGs的一致性", level=2)
add_paragraph("本研究与可持续发展目标12——\"负责任的消费和生产\"密切相关。通过关注不确定性下农产品供应链的有效库存管理，本研究以多种方式促进更可持续的实践：")
add_paragraph("• 减少食物浪费：优化的库存管理可以减少易腐食品的变质和浪费，促进更可持续的消费模式。", indent=False)
add_paragraph("• 资源效率：通过提高需求预测和库存决策的准确性，该算法有助于最小化过度生产和过量库存，从而在整个供应链中更高效地利用资源。", indent=False)
add_paragraph("• 经济可持续性：改进的库存管理带来的盈利能力和韧性增强，有助于农产品行业企业的经济可持续性。", indent=False)
add_paragraph("• 适应能力：算法的持续学习方面支持开发更具适应性和韧性的供应链，这对于面对气候变化和其他全球挑战的长期可持续性至关重要。", indent=False)

# 7. 结论
add_heading("7. 结论", level=1)
add_paragraph("农产品供应链相关的不确定性对供应链决策有重大影响，突出了优化库存决策稳健框架的需求。本研究通过提出一种新型算法来解决这一挑战，该算法设计用于在市场不确定性（包括随机客户需求和可变补货提前期）下优化库存管理。A3C-DPPO算法通过持续适应基础供应链的动态和复杂环境，推荐接近最优的库存补货策略。通过在连续动作空间中选择最优订货量，所提方法有效处理需求和供应双方的不确定性。")

add_paragraph("所提算法的有效性已在随机需求场景下得到研究。计算结果表明，所提库存策略在降低总库存成本和提高农产品供应链整体利润方面，优于基于DQN的库存策略以及经典(s,S)策略。数值结果还表明，基于A3C-DPPO的库存策略即使在高度动态和随机的市场不确定性下也相对稳定，而经典(s,S)策略反映出高度波动的特性。所提模型为食品行业的决策者提供了一个有价值的工具，用于优化库存运营，以提高食品供应链的整体韧性、响应能力和效率。")

add_paragraph("尽管所提研究涵盖了处理农产品供应链不确定性的各个方面，但仍存在一些局限性。首先，为管理系统复杂性而做的简化假设导致对现实世界动态的过度简化，从而可能导致潜在偏差和不准确。其次，本研究未考虑社会和环境可持续性方面，这些对于通过优化运输路线、减少CO₂排放以及通过平衡利润目标与环境影响来强调长期可持续性以改善供应链整体绩效至关重要。")

add_paragraph("在当前研究的基础上，有引人注目的未来探索方向，可以显著提高农产品供应链的效率和有效性。一个主要关注领域是整合食品冷链要求，这对于在供应链运输过程中保持新鲜农产品的质量和安全至关重要。这一扩展可能涉及优化温控物流、评估冷链违规的影响，以及分析先进冷链技术的成本效益。同时，研究越库设施作为关键物流组件的作用是另一个有前景的方向。这些设施有可能通过最小化处理时间和运输费用来优化效率并降低库存持有成本。未来研究可以关注越库设施的最优位置和设计、开发高效的调度算法，以及将越库策略与所提A3C-DPPO算法整合。此外，探索分销中心和零售商之间的多源采购和交叉连接策略可以提高供应链的稳健性和灵活性。通过结合这些方面，研究者可以发现有价值的见解，以创建更具响应性、效率和可持续性的农产品供应链，最终降低成本并提高整体盈利能力。")

# 参考文献
add_heading("参考文献", level=1)
add_paragraph("[1] Handayati, Y., Simatupang, T. M., & Perdana, T. (2015). Agri-food supply chain coordination: the state-of-the-art and recent developments. Logistics Research, 8, 1-15.", indent=False)
add_paragraph("[2] Tamasiga, P., et al. (2023). Forecasting disruptions in global food value chains to tackle food insecurity: The role of AI and big data analytics. Journal of Agriculture and Food Research, 14, 100819.", indent=False)
add_paragraph("[3] Sorooshian, S. (2024). The Sustainable Development Goals of the United Nations: A Comparative Midterm Research Review. Journal of Cleaner Production, 142272.", indent=False)
add_paragraph("[4] Ramos, E., et al. (2021). Measuring agri-food supply chain performance: insights from the Peruvian kiwicha industry. Benchmarking: An International Journal, 29(5), 1484-1512.", indent=False)
add_paragraph("[5] Flanagan, K., Robertson, K., & Hanson, C. (2019). Reducing food loss and waste. Setting the Global Action Agenda. Washington, DC, USA.", indent=False)
add_paragraph("[6] World Health Organization. (2022). WHO global strategy for food safety 2022-2030. World Health Organization.", indent=False)
add_paragraph("[7] Chen, S., et al. (2020). The role of smart packaging system in food supply chain. Journal of Food Science, 85(3), 517-525.", indent=False)
add_paragraph("[8] Chopra, S., & Meindl, P. (2001). Strategy, planning, and operation. Supply Chain Management, 15(5), 71-85.", indent=False)
add_paragraph("[9] Silver, E. A., Pyke, D. F., & Thomas, D. J. (2016). Inventory and production management in supply chains. CRC press.", indent=False)
add_paragraph("[10] Gaudenzi, B., & Christopher, M. (2016). Achieving supply chain 'Leagility' through a project management orientation. International Journal of Logistics Research and Applications, 19(1), 3-18.", indent=False)
add_paragraph("[11] Attaran, M. (2020). Digital technology enablers and their implications for supply chain management. Supply Chain Forum: An International Journal, 21(3).", indent=False)
add_paragraph("[12] Wang, X., & Disney, S. M. (2016). The bullwhip effect: Progress, trends and directions. European Journal of Operational Research, 250(3), 691-701.", indent=False)
add_paragraph("[13]-[48] 其他参考文献原文保留。", indent=False)

# 保存文档
output_path = r"c:\个人资料\申博材料\企业运营与科研管理数据库\Adaptive Inventory Strategies_中文翻译.docx"
doc.save(output_path)
print(f"文档已保存至: {output_path}")
print(f"文件大小: {os.path.getsize(output_path)/1024:.1f} KB")
