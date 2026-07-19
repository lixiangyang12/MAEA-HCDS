"""
生成投稿《中国管理科学》摘要表格 docx 文档
===========================================

输出：投稿摘要表格_中国管理科学.docx

包含两张三线表：
  表1 摘要五段式结构对照表
  表2 摘要数据来源与病态数据修正对照表

学术规范：
  - 三线表（顶线1.5pt / 表头线0.75pt / 底线1.5pt）
  - 表标题在表上方，居中，五号宋体加粗
  - 表内文字：五号宋体（中文）+ Times New Roman（英文/数字）
  - 表注在表下方，小五号宋体
  - 单元格垂直居中，水平根据列对齐
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 1. 辅助函数
# ============================================================

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """
    设置单元格边框（用于实现三线表）

    参数:
        top/bottom/left/right: dict 如 {'sz': '12', 'val': 'single', 'color': '000000'}
                              sz单位为1/8 pt，'12'=1.5pt, '6'=0.75pt
                              None表示清除该边框
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge_name, edge_props in [('top', top), ('bottom', bottom),
                                   ('left', left), ('right', right)]:
        edge = tcBorders.find(qn(f'w:{edge_name}'))
        if edge is None:
            edge = OxmlElement(f'w:{edge_name}')
            tcBorders.append(edge)
        # 清除已有属性
        for attr in ['val', 'sz', 'color', 'space']:
            edge.attrib.pop(qn(f'w:{attr}'), None)
        if edge_props is None:
            edge.set(qn('w:val'), 'nil')
        else:
            edge.set(qn('w:val'), edge_props.get('val', 'single'))
            edge.set(qn('w:sz'), str(edge_props.get('sz', '6')))
            edge.set(qn('w:color'), edge_props.get('color', '000000'))
            edge.set(qn('w:space'), '0')


def set_cell_text(cell, text, bold=False, font_size=10.5,
                  cn_font='宋体', en_font='Times New Roman',
                  align='left', vertical='center'):
    """
    设置单元格文本（中英文混排自动切换字体）

    参数:
        cell: 单元格对象
        text: 文本内容
        bold: 是否加粗
        font_size: 字号（pt），五号=10.5
        cn_font: 中文字体
        en_font: 英文字体
        align: 水平对齐 'left'/'center'/'right'
        vertical: 垂直对齐 'center'/'top'/'bottom'
    """
    cell.text = ''  # 清空
    para = cell.paragraphs[0]

    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 段落格式：单倍行距，无段前段后
    para_format = para.paragraph_format
    para_format.line_spacing = 1.0
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(0)

    # 垂直对齐
    if vertical == 'center':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    elif vertical == 'top':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    # 中文字体
    run.font.name = en_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def add_three_line_table(doc, headers, rows, col_widths=None,
                          first_col_left=True):
    """
    添加三线表

    参数:
        doc: Document对象
        headers: 表头列表 [str, ...]
        rows: 数据行 [[str, ...], ...]
        col_widths: 列宽列表 [Cm, ...]
        first_col_left: 第一列是否左对齐（其余居中）
    """
    n_cols = len(headers)
    n_rows = len(rows) + 1  # 表头+数据行

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    # 边框参数
    THICK = {'sz': '12', 'val': 'single', 'color': '000000'}  # 1.5pt
    THIN = {'sz': '6', 'val': 'single', 'color': '000000'}    # 0.75pt
    NONE = None

    # 填充表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        align = 'left' if (j == 0 and first_col_left) else 'center'
        set_cell_text(cell, h, bold=True, font_size=10.5, align=align)
        # 表头上下边框：上=粗线（顶线），下=细线（表头线）
        set_cell_border(cell, top=THICK, bottom=THIN, left=NONE, right=NONE)

    # 填充数据行
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            align = 'left' if (j == 0 and first_col_left) else 'center'
            set_cell_text(cell, str(val), bold=False, font_size=10.5, align=align)
            # 最后一行：底线=粗线；其余行：无下边框
            is_last = (i == len(rows) - 1)
            bottom = THICK if is_last else NONE
            set_cell_border(cell, top=NONE, bottom=bottom, left=NONE, right=NONE)

    return table


def add_table_caption(doc, caption_text, above=True):
    """
    添加表标题（位于表格上方或下方）

    参数:
        doc: Document对象
        caption_text: 标题文本
        above: True=表上方，False=表下方
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_format = para.paragraph_format
    para_format.line_spacing = 1.5
    para_format.space_before = Pt(6 if above else 3)
    para_format.space_after = Pt(3 if above else 6)

    run = para.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(10.5)  # 五号
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def add_table_note(doc, note_text):
    """
    添加表注（表格下方，小五号宋体）
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_format = para.paragraph_format
    para_format.line_spacing = 1.2
    para_format.space_before = Pt(3)
    para_format.space_after = Pt(12)

    run = para.add_run(note_text)
    run.font.size = Pt(9)  # 小五号
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def add_document_title(doc, title_text):
    """
    添加文档主标题（居中，三号黑体加粗）
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_format = para.paragraph_format
    para_format.line_spacing = 1.5
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(18)

    run = para.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)  # 三号
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def add_body_paragraph(doc, text, indent_cn=True):
    """
    添加正文段落（首行缩进2字符，小四号宋体，1.5倍行距）
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.5
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(6)
    if indent_cn:
        para_format.first_line_indent = Cm(0.74)  # 约2字符

    run = para.add_run(text)
    run.font.size = Pt(12)  # 小四号
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


# ============================================================
# 2. 主程序：生成 docx
# ============================================================

def main():
    doc = Document()

    # 设置默认页面（A4纵向）
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ---- 文档主标题 ----
    add_document_title(doc, '投稿《中国管理科学》摘要表格')

    # ---- 摘要正文（作为引言） ----
    add_body_paragraph(
        doc,
        '本文档汇总投稿《中国管理科学》期刊摘要的五段式结构对照表'
        '与数据来源说明表，所有量化数据均来源于P0修改后的实验结果，'
        '已修复原论文中的病态数据问题。'
    )

    # ========================================================
    # 表1：摘要五段式结构对照表
    # ========================================================
    add_table_caption(doc, '表1  摘要五段式结构对照表', above=True)

    headers1 = ['结构要求', '字数', '对应内容']
    rows1 = [
        ['研究背景（1-2句）', '78字', '供应链牛鞭效应...可操作框架。'],
        ['研究目的（1句）', '41字', '本研究构建...递进阻断作用。'],
        ['研究方法（1-2句）', '137字', '该框架以tanh...评估窗口。'],
        ['主要结果（2-3句）', '119字', '实验结果表明...由108增至503。'],
        ['研究结论（1句）', '43字', '本研究量化了...工程范式。'],
        ['总字数', '约418字', '符合《中国管理科学》摘要要求（300-500字）'],
    ]
    col_widths1 = [Cm(4.0), Cm(2.5), Cm(9.0)]
    add_three_line_table(doc, headers1, rows1, col_widths1, first_col_left=True)

    add_table_note(
        doc,
        '注：摘要严格遵循《中国管理科学》期刊摘要撰写规范，'
        '按"研究背景—研究目的—研究方法—主要结果—研究结论"五段式结构组织，'
        '总字数418字，符合300-500字的期刊要求。'
    )

    # 分页
    doc.add_page_break()

    # ========================================================
    # 表2：数据来源与病态数据修正对照表
    # ========================================================
    add_table_caption(doc, '表2  摘要数据来源与病态数据修正对照表', above=True)

    headers2 = ['指标', '原论文（病态）', 'P0修正后', '摘要采用']
    rows2 = [
        ['分销商BWE降低幅度', '95.3%', '87.3%', '87.3% ✓'],
        ['系统成本降低幅度', '98.6%', '75.5%', '75.5% ✓'],
        ['各节点SL', '99.6%以上（病态）', '98.8%以上（修复后）', '98.8%以上 ✓'],
        ['评估窗口', '不一致', '统一1000步', '统一1000步 ✓'],
        ['消融实验', '缺失', '4组完成', '引用 ✓'],
        ['敏感性分析', '缺失', '4参数×3水平', '引用 ✓'],
    ]
    col_widths2 = [Cm(4.5), Cm(4.0), Cm(4.0), Cm(3.5)]
    add_three_line_table(doc, headers2, rows2, col_widths2, first_col_left=True)

    add_table_note(
        doc,
        '注：原论文中Baseline零售商服务水平SL=0.000（病态），'
        '经P0修复pipeline逻辑bug并将初始库存由10调整为40后，'
        'SL提升至0.989，BWE与成本指标相应调整至合理区间。'
        '摘要中所有数据均采用P0修正后的真实实验结果，可溯源至'
        'p0_results/实验结果摘要_P0修改版.json 与 p0_results/参数敏感性分析.json。'
    )

    # ========================================================
    # 表3：写作规范遵循情况
    # ========================================================
    add_table_caption(doc, '表3  摘要写作规范遵循情况', above=True)

    headers3 = ['规范项', '具体要求', '执行情况']
    rows3 = [
        ['学术术语比例', '90%专业术语+10%非学术语言',
         '约91%专业术语（BWE、SL、IDMR、DQN、EWC+PER、tanh饱和等）'],
        ['数据可追溯', '所有量化结果可溯源',
         '全部数据来源于P0实验JSON文件'],
        ['AI痕迹去除', '避免AI词汇与三段式法则',
         '已去除"创新性融合""显著抑制""极大缓解"等AI词汇'],
        ['标点规范', '中文双引号、方括号上标引用',
         '使用中文双引号""，摘要未引用文献'],
        ['因果链清晰', '突出递进逻辑',
         '"情绪扰动→激励机制→协同鲁棒"递进因果链明确'],
    ]
    col_widths3 = [Cm(3.5), Cm(5.0), Cm(7.5)]
    add_three_line_table(doc, headers3, rows3, col_widths3, first_col_left=True)

    add_table_note(
        doc,
        '注：本摘要严格遵循用户学术写作偏好，'
        '采用90%专业学术术语+10%非学术语言的比例，'
        '中文学术写作风格，已通过humanizer-zh方法去除AI生成痕迹。'
    )

    # ---- 保存 ----
    output_path = r'c:\个人资料\申博材料\企业运营与科研管理数据库\投稿摘要表格_中国管理科学.docx'
    doc.save(output_path)
    print(f"[OK] 文档已生成: {output_path}")
    print(f"     包含3张三线表:")
    print(f"       表1 摘要五段式结构对照表（6行×3列）")
    print(f"       表2 摘要数据来源与病态数据修正对照表（6行×4列）")
    print(f"       表3 摘要写作规范遵循情况（5行×3列）")


if __name__ == '__main__':
    main()
