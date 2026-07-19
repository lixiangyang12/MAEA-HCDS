"""
补充引用脚本
============
功能：为参考文献中未被引用的文献（33篇）在正文中添加引用标注
输入：改进牛鞭效应新途径：人智协同决策系统_引用版.docx
输出：改进牛鞭效应新途径：人智协同决策系统_引用版_补引.docx
"""
import re
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

INPUT_PATH = r'c:\个人资料\申博材料\企业运营与科研管理数据库\改进牛鞭效应新途径：人智协同决策系统\改进牛鞭效应新途径：人智协同决策系统_引用版.docx'
OUTPUT_PATH = r'c:\个人资料\申博材料\企业运营与科研管理数据库\改进牛鞭效应新途径：人智协同决策系统\改进牛鞭效应新途径：人智协同决策系统_引用版_补引.docx'

FONT_CN = '宋体'
FONT_EN = 'Times New Roman'

# ============================================================
# 文本替换规则（为未引用文献添加引用标注）
# 每条规则：(原文片段, 替换后文本)
# ============================================================
TEXT_REPLACEMENTS = [
    # ===== 段落6：引言，牛鞭效应背景 → 添加[12] Towill =====
    ('指供应链中下游需求波动向上游逐级放大的现象',
     '指供应链中下游需求波动向上游逐级放大的现象[12]'),

    # ===== 段落7：理论基础集中引用 =====
    # 行为运营管理 → 添加[34] Bendoly (保留[35])
    ('（Gino F, Pisano G） [35]', '[34][35]'),
    # 六大理论支柱 → 集中引用各领域文献
    ('深度强化学习、多智能体系统、持续学习与情感计算六大理论支柱',
     '深度强化学习[3][17][18][19][21][22][36][37][38]、多智能体系统[6][7][24][25][26][27][28][29]、持续学习[4][5]与情感计算[9][10]六大理论支柱'),

    # ===== 段落8：CTDE范式 → 添加[28] Terry PettingZoo =====
    ('采用CTDE范式训练独立DQN智能体',
     '采用CTDE范式[6][7]训练独立DQN智能体'),

    # ===== 段落11：DQN智慧决策 → 添加[3] Mnih, [18] Watkins =====
    ('李勇等[33]基于DQN设计了',
     '李勇等[33]基于DQN[3][18]设计了'),

    # ===== 段落117：可复现性 → 添加[23] Lopez de Prado =====
    ('确保结果可量化、可复现',
     '确保结果可量化、可复现[23]'),

    # ===== 段落125：损失厌恶 → 添加[14][30][31][32] =====
    ('行为经济学研究[2]表明，人类决策者普遍存在损失厌恶倾向',
     '行为经济学研究[2][14][30][31][32]表明，人类决策者普遍存在损失厌恶倾向'),

    # ===== 段落152：DQN训练 → 添加[39] BatchNorm, [40] Adam =====
    ('训练过程涵盖损失函数收敛、奖励提升、探索率衰减与牛鞭效应控制四个维度',
     '训练过程采用Adam优化器[40]与批归一化[39]技术，涵盖损失函数收敛、奖励提升、探索率衰减与牛鞭效应控制四个维度'),

    # ===== 段落162：EWC+PER → 添加[4] Kirkpatrick, [5] Schaul =====
    ('持续学习机制（EWC+PER+情绪感知噪声）',
     '持续学习机制（EWC[4]+PER[5]+情绪感知噪声）'),

    # ===== 段落173：牛鞭效应四大成因 → 添加[15] Croson, [16] Boute =====
    ('符合Lee等[1]提出的牛鞭效应四大成因',
     '符合Lee等[1]提出的牛鞭效应四大成因[15][16]'),

    # ===== 段落219：大语言模型 → 添加[20] Vaswani =====
    ('大语言模型（LLM）',
     '大语言模型（LLM）[20]'),
]


def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN, size=None, bold=None):
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def has_omml(paragraph):
    p_element = paragraph._element
    ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    return len(p_element.findall('.//m:oMath', ns)) > 0 or \
           len(p_element.findall('.//m:oMathPara', ns)) > 0


def get_paragraph_font_info(paragraph):
    for run in paragraph.runs:
        if run.text and run.text.strip():
            return {
                'size': run.font.size,
                'bold': run.font.bold,
            }
    return {'size': Pt(11), 'bold': None}


def needs_replacement(text):
    for old, new in TEXT_REPLACEMENTS:
        if old in text:
            return True
    return False


def rebuild_paragraph(paragraph, font_info):
    """重建段落：保留段落格式，重建run，将[n]转换为上标"""
    full_text = paragraph.text

    # 应用文本替换
    replaced = []
    for old, new in TEXT_REPLACEMENTS:
        if old in full_text:
            replaced.append(old[:30])
            full_text = full_text.replace(old, new)

    # 清空段落所有run
    p_element = paragraph._element
    for r in p_element.findall(qn('w:r')):
        p_element.remove(r)

    # 按 [n] 模式分割文本
    parts = re.split(r'(\[\d+\])', full_text)

    for part in parts:
        if not part:
            continue
        if re.match(r'^\[\d+\]$', part):
            run = paragraph.add_run(part)
            run.font.superscript = True
            set_run_font(run, size=font_info['size'], bold=font_info['bold'])
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=font_info['size'], bold=font_info['bold'])


def replace_in_runs(paragraph):
    """对含OMML公式的段落，仅在run级别做文本替换"""
    text_runs = []
    for run in paragraph.runs:
        r_element = run._element
        ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        has_math = len(r_element.findall('.//m:oMath', ns)) > 0
        if not has_math and run.text:
            text_runs.append(run)

    if not text_runs:
        return False

    combined = ''.join(run.text for run in text_runs)
    original = combined

    for old, new in TEXT_REPLACEMENTS:
        combined = combined.replace(old, new)

    if combined == original:
        return False

    # 处理[n]上标
    if not re.search(r'\[\d+\]', combined):
        text_runs[0].text = combined
        for run in text_runs[1:]:
            run.text = ''
        return True

    # 包含[n]，需要创建上标run
    parts = re.split(r'(\[\d+\])', combined)

    for run in text_runs:
        run.text = ''

    font_info = get_paragraph_font_info(paragraph)
    first_run_elem = text_runs[0]._element
    parent = first_run_elem.getparent()
    insert_idx = list(parent).index(first_run_elem)

    for run in text_runs:
        parent.remove(run._element)

    for part in parts:
        if not part:
            continue
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_CN)
        rFonts.set(qn('w:ascii'), FONT_EN)
        rFonts.set(qn('w:hAnsi'), FONT_EN)
        rPr.append(rFonts)

        if font_info['size']:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_info['size'].pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_info['size'].pt * 2)))
            rPr.append(szCs)

        if re.match(r'^\[\d+\]$', part):
            vertAlign = OxmlElement('w:vertAlign')
            vertAlign.set(qn('w:val'), 'superscript')
            rPr.append(vertAlign)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = part
        t.set(qn('xml:space'), 'preserve')
        new_run.append(t)

        parent.insert(insert_idx, new_run)
        insert_idx += 1

    return True


def process_document():
    doc = Document(INPUT_PATH)

    # 找到参考文献开始位置
    ref_start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '参考文献':
            ref_start_idx = i
            break

    print(f"[信息] 参考文献起始段落: {ref_start_idx}")
    print(f"[信息] 总段落数: {len(doc.paragraphs)}")

    stats = {'total': 0, 'replaced': 0, 'skipped_ref': 0,
             'skipped_omml': 0, 'skipped_nochange': 0}

    for i, p in enumerate(doc.paragraphs):
        stats['total'] += 1

        if i >= ref_start_idx:
            stats['skipped_ref'] += 1
            continue

        if not p.text.strip():
            stats['skipped_nochange'] += 1
            continue

        if not needs_replacement(p.text):
            stats['skipped_nochange'] += 1
            continue

        if has_omml(p):
            print(f"  [处理-含公式] 段落[{i}]: {p.text[:60]}...")
            if replace_in_runs(p):
                stats['replaced'] += 1
            else:
                stats['skipped_nochange'] += 1
            continue

        font_info = get_paragraph_font_info(p)
        print(f"  [处理] 段落[{i}]: {p.text[:60]}...")
        rebuild_paragraph(p, font_info)
        stats['replaced'] += 1

    # 处理文件占用
    def _try_save(doc, path):
        try:
            if os.path.exists(path):
                os.remove(path)
            doc.save(path)
            return True
        except PermissionError:
            return False

    import datetime
    if not _try_save(doc, OUTPUT_PATH):
        ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        OUTPUT_PATH_TS = OUTPUT_PATH.replace('.docx', f'_{ts}.docx')
        doc.save(OUTPUT_PATH_TS)
        print(f"[提示] 原文件被占用，保存为: {OUTPUT_PATH_TS}")
        OUTPUT_PATH_FINAL = OUTPUT_PATH_TS
    else:
        OUTPUT_PATH_FINAL = OUTPUT_PATH

    print(f"\n[OK] 文档已保存: {OUTPUT_PATH_FINAL}")
    print(f"     大小: {os.path.getsize(OUTPUT_PATH_FINAL) / 1024:.1f} KB")
    print(f"\n[统计]")
    print(f"  总段落: {stats['total']}")
    print(f"  已替换: {stats['replaced']}")
    print(f"  跳过(参考文献): {stats['skipped_ref']}")
    print(f"  跳过(含公式): {stats['skipped_omml']}")
    print(f"  跳过(无需替换): {stats['skipped_nochange']}")

    # 验证引用完整性
    print(f"\n[验证] 引用完整性检查:")
    cited = set()
    for i, p in enumerate(doc.paragraphs):
        if i >= ref_start_idx:
            continue
        matches = re.findall(r'\[(\d+)\]', p.text)
        for m in matches:
            num = int(m)
            if 1 <= num <= 40:
                cited.add(num)

    all_refs = set(range(1, 41))
    uncited = all_refs - cited
    print(f"  已引用序号: {sorted(cited)}")
    print(f"  已引用数量: {len(cited)}/40")
    if uncited:
        print(f"  ⚠ 仍未引用: {sorted(uncited)}")
    else:
        print(f"  ✅ 全部40篇文献均已引用")


if __name__ == '__main__':
    process_document()
